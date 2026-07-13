import pandas as pd
import os
import glob
import io
import numpy as np
import warnings
import re
import datetime
from decimal import Decimal, ROUND_HALF_UP

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import openpyxl
from openpyxl.styles import Alignment, Border, Side, Font, PatternFill

warnings.filterwarnings('ignore')

# ================= 全局参数 =================
TODAY_STR = datetime.datetime.now().strftime("%Y%m%d")
TODAY_MMDD = datetime.datetime.now().strftime("%m%d")

REASON1_ORDER = ['一、客户原因/客户原因为主', '二、我方原因/我方原因为主', '三、既非我方原因也非对方原因']
REASON2_ORDER = {
    '一、客户原因/客户原因为主': ['客户加工计划调整', '客户基差合同点价晚', '客户其他原因（需详细说明）', '客户仓容紧张', '客户资金紧张', '我方正常到货，客户提货能力不足', '客户额度不足', '客户工厂检修', '销价高客户不愿提货'],
    '二、我方原因/我方原因为主': ['一体化协同粮源质量问题', '外采粮源质量问题', '大区责任物流原因', '我方货源不足', '一体化非大区责任物流原因', '外采供应商责任物流原因', '我方其他原因（需详细说明）', '我方到货集中', '我方修路'],
    '三、既非我方原因也非对方原因': ['其他原因（需详细说明）', '天气原因', '政府行为', '自然灾害', '社会异常事件']
}
TIME_ORDER = ['1-10天', '11-20天', '21-30天', '31-60天', '61-90天', '90天以上']

COLOR_BATCH = "E6F7FF"  
COLOR_ONCE = "FFF7E6"   

# ================= 数据加载与清洗 =================
def _normalize_col_name(name):
    """去除列名中的所有空白字符，包括全角空格(U+3000)、不间断空格(U+00A0)等"""
    import re
    s = str(name)
    s = s.replace('　', ' ')   # 全角空格 → 普通空格
    s = s.replace(' ', ' ')   # 不间断空格 → 普通空格
    s = re.sub(r'\s+', '', s)      # 清除所有ASCII空白字符
    return s


def _normalize_brackets(text):
    """
    统一括号格式：全角括号 → 半角括号，中文括号 → 英文括号。
    用于客户名称匹配前的预处理，解决数据源括号不一致导致的匹配失败问题。
    例如：（中粮）→ (中粮)，(中粮）→ (中粮)，（中粮) → (中粮)
    """
    if not isinstance(text, str):
        return str(text) if text is not None else ''
    # 全角括号 → 半角
    text = text.replace('（', '(').replace('）', ')')
    # 中文括号（与全角相同，此处做兜底）统一为半角
    text = text.replace('【', '[').replace('】', ']')
    text = text.replace('《', '<').replace('》', '>')
    return text.strip()

def locate_header_and_read(file_stream, key_columns):
    try:
        file_stream.seek(0)
        df_raw = pd.read_excel(file_stream, header=None, nrows=200)
        header_row_index = -1
        
        for i, row in df_raw.iterrows():
            row_values = [_normalize_col_name(x) for x in row.values if pd.notna(x)]
            match_count = sum(1 for key in key_columns if key in row_values)
            if match_count >= len(key_columns) - 1:
                header_row_index = i
                break
        
        if header_row_index == -1:
            return None

        file_stream.seek(0)
        df = pd.read_excel(file_stream, header=header_row_index)
        
        df.columns = [_normalize_col_name(c) for c in df.columns]
        
        if '大区' in df.columns:
            col_idx = df.columns.get_loc('大区')
            if isinstance(col_idx, np.ndarray):
                start = np.where(col_idx)[0][0]
            elif isinstance(col_idx, slice):
                start = col_idx.start
            else:
                start = col_idx
            df = df.iloc[:, start:]
            
        df = df.loc[:, ~df.columns.duplicated()]
        
        df.dropna(how='all', inplace=True)
        return df
    except Exception:
        return None

def process_basic_columns(df, date_cols, float_cols, int_cols=None):
    for col in date_cols:
        if col in df.columns:
            temp_dates = pd.to_datetime(df[col], errors='coerce')
            mask_invalid = (temp_dates.dt.year < 2025) | (temp_dates.dt.year > 2030) | (temp_dates.isna())
            if mask_invalid.any():
                raw_values = df.loc[mask_invalid, col]
                clean_raw = raw_values.astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
                corrected = pd.to_datetime(clean_raw, format='%Y%m%d', errors='coerce')
                temp_dates.loc[mask_invalid] = corrected
            df[col] = temp_dates

    for col in float_cols:
        if col in df.columns:
            if df[col].dtype == object:
                df[col] = df[col].astype(str).str.replace(',', '', regex=False)
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
    
    if int_cols:
        for col in int_cols:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
    return df

def merge_detail_variety_columns(df):
    source_cols = ['明细品种', '细分品种']
    existing_cols = [c for c in source_cols if c in df.columns]
    if not existing_cols:
        df['明细品种'] = ""
        return df
    if '明细品种' in df.columns and '细分品种' in df.columns:
        df['明细品种'] = df['明细品种'].fillna(df['细分品种'])
        df['明细品种'] = df['明细品种'].astype(str).replace('nan', '')
    elif '细分品种' in df.columns:
        df.rename(columns={'细分品种': '明细品种'}, inplace=True)
    if '明细品种' in df.columns:
        df['明细品种'] = df['明细品种'].astype(str).replace('nan', '')
    return df

def process_variety_logic(df):
    if '品种' not in df.columns: return df
    if '明细品种' not in df.columns: df['明细品种'] = ""

    df['品种'] = df['品种'].astype(str).str.strip()
    df['明细品种'] = df['明细品种'].astype(str).str.strip()

    rice_condition = df['明细品种'].str.contains('稻谷|中晚籼', regex=True, na=False)
    df.loc[rice_condition, '品种'] = '稻谷'

    whitelist = ['大豆', '稻谷', '小麦']
    not_in_whitelist = ~df['品种'].isin(whitelist)
    detail_not_empty = df['明细品种'] != ''
    mask = not_in_whitelist & detail_not_empty
    df.loc[mask, '品种'] = df.loc[mask, '明细品种']

    return df

def get_customer_mappings():
    base_dir = os.getcwd()
    search_pattern = os.path.join(base_dir, "**", "*.xlsx")
    files = glob.glob(search_pattern, recursive=True)
    mapping_file = None
    for f in files:
        filename = os.path.basename(f)
        if "客户关系清单" in filename and not filename.startswith("~$"):
            mapping_file = f
            break
            
    if not mapping_file:
        return {}, {}, {}, {}

    try:
        df_total = pd.read_excel(mapping_file, sheet_name='总')
        df_total.columns = df_total.columns.astype(str).str.strip().str.replace('\n', '').str.replace('\r', '')
        df_total.dropna(subset=['客户名称'], inplace=True)
        # 原始映射
        group_map = dict(zip(df_total['客户名称'].astype(str).str.strip(), df_total['客户所属集团']))
        # 括号归一化映射：键名统一转为半角英文括号，提升匹配鲁棒性
        group_map_normalized = {
            _normalize_brackets(k): v
            for k, v in group_map.items()
        }

        df_internal = pd.read_excel(mapping_file, sheet_name='内部')
        df_internal.columns = df_internal.columns.astype(str).str.strip().str.replace('\n', '').str.replace('\r', '')
        df_internal.dropna(subset=['客户名称'], inplace=True)
        internal_company_map = dict(zip(df_internal['客户名称'].astype(str).str.strip(), df_internal['所属专业化公司']))
        # 括号归一化映射
        internal_company_map_normalized = {
            _normalize_brackets(k): v
            for k, v in internal_company_map.items()
        }

        return group_map, internal_company_map, group_map_normalized, internal_company_map_normalized
    except Exception:
        return {}, {}, {}, {}

def map_reason1(x):
    if pd.isna(x): return x
    if "客户原因" in str(x): return "一、客户原因/客户原因为主"
    if "我方原因" in str(x): return "二、我方原因/我方原因为主"
    if "既非我方" in str(x) or "非对方" in str(x): return "三、既非我方原因也非对方原因"
    return str(x)

# ================= 格式化与催收提醒 =================
def format_num(val, dec=2, is_int=False, is_percent=False):
    if pd.isna(val) or val == "": return ""
    try: d_val = Decimal(str(round(float(val), 6)))
    except: return val
    
    if is_percent:
        if d_val == 0: return "0%"
        elif abs(d_val) >= 1:
            rounded = d_val.quantize(Decimal('1'), rounding=ROUND_HALF_UP)
            return f"{int(rounded)}%"
        else:
            s_val = f"{float(abs(d_val)):.6f}"
            if '.' in s_val:
                dec_part = s_val.split('.')[1]
                non_zero_idx = next((i for i, digit in enumerate(dec_part) if digit != '0'), -1)
                if non_zero_idx != -1:
                    q_str = '0.' + '0' * non_zero_idx + '1'
                    rounded = d_val.quantize(Decimal(q_str), rounding=ROUND_HALF_UP)
                    return f"{float(rounded)}".rstrip('0').rstrip('.') + "%"
            return f"{float(d_val)}%"
            
    if is_int:
        rounded_int = d_val.quantize(Decimal('1'), rounding=ROUND_HALF_UP)
        return f"{int(rounded_int):,}"
    
    q = Decimal('1.' + '0' * dec) if dec > 0 else Decimal('1')
    rounded = d_val.quantize(q, rounding=ROUND_HALF_UP)
    res = f"{float(rounded):,.{dec}f}"
    if '.' in res:
        res = res.rstrip('0').rstrip('.')
    if res == "": res = "0"
    return res

def format_qty(val):
    if pd.isna(val) or val == "": return ""
    try: d_val = Decimal(str(round(float(val), 6)))
    except: return val
    if d_val == 0: return "0"
    
    if abs(d_val) >= 1:
        return format_num(val, dec=2) 
        
    q2 = Decimal('1.00')
    rounded2 = d_val.quantize(q2, rounding=ROUND_HALF_UP)
    if rounded2 != 0:
        res = f"{float(rounded2):.2f}".rstrip('0').rstrip('.')
        return res if res else "0"
    else:
        s_val = f"{float(abs(d_val)):.10f}"
        if '.' in s_val:
            dec_part = s_val.split('.')[1]
            non_zero_idx = next((i for i, digit in enumerate(dec_part) if digit != '0'), -1)
            if non_zero_idx != -1:
                q_str = '0.' + '0' * non_zero_idx + '1'
                rounded = d_val.quantize(Decimal(q_str), rounding=ROUND_HALF_UP)
                return f"{float(rounded)}".rstrip('0').rstrip('.')
    return f"{float(d_val)}".rstrip('0').rstrip('.')

def generate_collection_reminder(df_unique):
    yesterday = datetime.datetime.now() - datetime.timedelta(days=1)
    date_str = f"{yesterday.month}月{yesterday.day}日"
    lines = []
    
    regions = df_unique['大区'].dropna().unique() if '大区' in df_unique.columns else []
    
    if len(regions) >= 2 or len(regions) == 0:
        lines.append("中粮贸易：")
        tot_cnt = len(df_unique)
        tot_qty = df_unique['逾期数量（万吨）'].sum()
        tot_amt = df_unique['逾期金额（万元）'].sum()
        safe_tot_qty = tot_qty if tot_qty > 0 else 1e-9
        avg_days_val = (df_unique['逾期数量（万吨）'] * df_unique['逾期天数']).sum() / safe_tot_qty if tot_qty > 0 else 0
        avg_days = int(Decimal(str(round(avg_days_val, 6))).quantize(Decimal('1'), rounding=ROUND_HALF_UP))
        base_str = f"截至{date_str}，中粮贸易逾期销售提货合同合计{tot_cnt}笔，逾期数量{format_qty(tot_qty)}万吨，逾期金额{format_num(tot_amt, 0, True)}万元，平均逾期{avg_days}天。"
        
        v_stats = df_unique.groupby('品种').agg({'合同编号': 'count', '逾期数量（万吨）': 'sum', '逾期金额（万元）': 'sum'}).sort_values(by='逾期金额（万元）', ascending=False)
        v_parts = []
        safe_tot_amt = tot_amt if tot_amt > 0 else 1e-9
        for v in v_stats.index:
            v_amt = v_stats.loc[v, '逾期金额（万元）']
            v_ratio = format_num(v_amt / safe_tot_amt * 100, is_percent=True)
            v_parts.append(f"{v}{v_stats.loc[v, '合同编号']}笔，逾期数量{format_qty(v_stats.loc[v, '逾期数量（万吨）'])}万吨，逾期金额{format_num(v_amt, 0, True)}万元（{v_ratio}）")
        
        if v_parts:
            base_str += "其中，" + "；".join(v_parts) + "。"
        lines.append(base_str)
        
        lines.append("分大区看：")
        if '大区' in df_unique.columns:
            r_stats = df_unique.groupby('大区').agg({'合同编号': 'count', '逾期数量（万吨）': 'sum', '逾期金额（万元）': 'sum'}).sort_values(by='逾期金额（万元）', ascending=False)
            for i, r in enumerate(r_stats.index, 1):
                lines.append(f"{i}、{r}，逾期销售提货合同合计{r_stats.loc[r, '合同编号']}笔，逾期数量{format_qty(r_stats.loc[r, '逾期数量（万吨）'])}万吨，逾期金额{format_num(r_stats.loc[r, '逾期金额（万元）'], 0, True)}万元。")
    
    elif len(regions) == 1:
        region = regions[0]
        r_df = df_unique[df_unique['大区'] == region]
        r_cnt = len(r_df)
        r_qty = r_df['逾期数量（万吨）'].sum()
        r_amt = r_df['逾期金额（万元）'].sum()
        r_safe_qty = r_qty if r_qty > 0 else 1e-9
        r_avg_days_val = (r_df['逾期数量（万吨）'] * r_df['逾期天数']).sum() / r_safe_qty if r_qty > 0 else 0
        r_avg_days = int(Decimal(str(round(r_avg_days_val, 6))).quantize(Decimal('1'), rounding=ROUND_HALF_UP))
        
        lines.append(f"{region}：")
        lines.append(f"截至{date_str}，{region}逾期销售提货合同合计{r_cnt}笔，逾期数量{format_qty(r_qty)}万吨，逾期金额{format_num(r_amt, 0, True)}万元，平均逾期{r_avg_days}天。")
        lines.append("\n分经营部看：")
        
        if '经营部' in r_df.columns:
            d_stats = r_df.groupby('经营部').agg({'合同编号': 'count', '逾期数量（万吨）': 'sum', '逾期金额（万元）': 'sum'}).sort_values(by='逾期金额（万元）', ascending=False)
            for i, d in enumerate(d_stats.index, 1):
                lines.append(f"{i}、{d}，逾期销售提货合同合计{d_stats.loc[d, '合同编号']}笔，逾期数量{format_qty(d_stats.loc[d, '逾期数量（万吨）'])}万吨，逾期金额{format_num(d_stats.loc[d, '逾期金额（万元）'], 0, True)}万元。")
                
                d_df = r_df[r_df['经营部'] == d].copy()
                has_focus = '是否重点关注' in d_df.columns
                has_severe = '是否严重逾期' in d_df.columns
                
                def get_label(row):
                    if pd.to_numeric(row.get('逾期天数', 0), errors='coerce') >= 60: return "逾期60天以上"
                    if has_severe and '严重逾期' in str(row.get('是否严重逾期', '')): return "严重逾期"
                    if has_focus and '重点关注' in str(row.get('是否重点关注', '')): return "重点关注"
                    return ""
                
                d_df['特殊标签'] = d_df.apply(get_label, axis=1)
                spec_df = d_df[d_df['特殊标签'] != ""]
                spec_df = spec_df.sort_values(by='逾期数量（万吨）', ascending=False)
                
                for _, row in spec_df.iterrows():
                    c_name = row.get('客户名称', '')
                    l_tag = row['特殊标签']
                    s_qty = format_qty(row.get('逾期数量（万吨）', 0))
                    s_amt = format_num(row.get('逾期金额（万元）', 0), 0, True)
                    s_days = format_num(row.get('逾期天数', 0), 0, True)
                    lines.append(f"• {c_name}，{l_tag}，逾期数量{s_qty}万吨，逾期金额{s_amt}万元，逾期{s_days}天。")
        
        lines.append("\n分品种看：")
        rv_stats = r_df.groupby('品种').agg({'合同编号': 'count', '逾期数量（万吨）': 'sum', '逾期金额（万元）': 'sum'}).sort_values(by='逾期金额（万元）', ascending=False)
        r_safe_amt = r_amt if r_amt > 0 else 1e-9
        for i, v in enumerate(rv_stats.index, 1):
            v_amt = rv_stats.loc[v, '逾期金额（万元）']
            v_ratio = format_num(v_amt / r_safe_amt * 100, is_percent=True)
            lines.append(f"{i}、{v}{rv_stats.loc[v, '合同编号']}笔，逾期数量{format_qty(rv_stats.loc[v, '逾期数量（万吨）'])}万吨，逾期金额{format_num(v_amt, 0, True)}万元（{v_ratio}）。")

    return "\n".join(lines)

def set_font_mixed(run_or_style, size_pt, bold=False, east_asia='仿宋_GB2312', ascii_font='Times New Roman'):
    run_or_style.font.size = Pt(size_pt)
    run_or_style.font.name = ascii_font
    run_or_style._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    run_or_style._element.rPr.rFonts.set(qn('w:ascii'), ascii_font)
    run_or_style._element.rPr.rFonts.set(qn('w:hAnsi'), ascii_font)
    run_or_style.font.bold = bold

def init_styles(doc):
    style_ch = doc.styles.add_style('ChapterTitle', 1)
    set_font_mixed(style_ch, 14.0, False, '黑体', 'Times New Roman')
    style_ch.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_ch.paragraph_format.line_spacing = Pt(28)
    style_ch.paragraph_format.space_before = Pt(0)
    style_ch.paragraph_format.space_after = Pt(0)
    style_ch.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_ch.paragraph_format.first_line_indent = Pt(28)

    style_tb = doc.styles.add_style('TableTitle', 1)
    set_font_mixed(style_tb, 12.0, False, '微软雅黑', 'Times New Roman')
    style_tb.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_tb.paragraph_format.line_spacing = Pt(22)
    style_tb.paragraph_format.space_before = Pt(0)
    style_tb.paragraph_format.space_after = Pt(0)
    style_tb.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_tb.paragraph_format.first_line_indent = Pt(0)

    style_nm = doc.styles.add_style('NormalContent', 1)
    set_font_mixed(style_nm, 14.0, False, '仿宋_GB2312', 'Times New Roman')
    style_nm.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_nm.paragraph_format.line_spacing = Pt(22)
    style_nm.paragraph_format.space_before = Pt(0)
    style_nm.paragraph_format.space_after = Pt(0)
    style_nm.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_nm.paragraph_format.first_line_indent = Pt(28) 

    style_app = doc.styles.add_style('AppendixTitle', 1)
    set_font_mixed(style_app, 15.0, False, '黑体', 'Times New Roman')
    style_app.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_app.paragraph_format.line_spacing = Pt(28)
    style_app.paragraph_format.space_before = Pt(0)
    style_app.paragraph_format.space_after = Pt(0)
    style_app.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_app.paragraph_format.first_line_indent = Pt(0)

def set_page_margins(doc):
    section = doc.sections[0]
    section.page_width = Cm(266710.5 / 12700.0)
    section.page_height = Cm(377194.0 / 12700.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(19063.55 / 12700.0)
    section.footer_distance = Cm(22222.14 / 12700.0)

def build_cell_text(cell, text, align='center', bold=False, is_max=False, is_appendix=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    if is_appendix:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(12)
        font_sz = 7.5
        en_font_sz = 7.5
    else:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(10)
        font_sz = 9.0
        en_font_sz = 10.0

    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    text_str = str(text) if not pd.isna(text) else ""
    parts = re.split(r'([a-zA-Z0-9.,%+-]+)', text_str)
    for part in parts:
        if not part: continue
        run = p.add_run(part)
        if is_max:
            run.font.color.rgb = RGBColor(255, 0, 0)
        
        if re.match(r'^[a-zA-Z0-9.,%+-]+$', part):
            set_font_mixed(run, en_font_sz, bold, '微软雅黑', 'Times New Roman')
        else:
            set_font_mixed(run, font_sz, bold, '微软雅黑', 'Times New Roman')
            
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_cell_background(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def apply_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') 
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    tblCellMar = OxmlElement('w:tblCellMar')
    for m in [('top', '28'), ('bottom', '28'), ('left', '57'), ('right', '57')]:
        node = OxmlElement(f'w:{m[0]}')
        node.set(qn('w:w'), m[1])
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)
    tblPr.append(tblCellMar)

def set_fixed_col_widths(table, widths, is_cm=False):
    table.autofit = False
    table.allow_autofit = False
    for i, w_val in enumerate(widths):
        w = Cm(w_val) if is_cm else Cm(w_val / 12700.0)
        table.columns[i].width = w
        for cell in table.columns[i].cells:
            cell.width = w

def set_table_row_height(row, height_pt):
    row.height = Pt(height_pt)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

def generate_report(df, df_unique):
    total_amount = df_unique['逾期金额（万元）'].sum()
    safe_total = total_amount if total_amount > 0 else 1e-9
    total_amount_str = format_num(total_amount, 0, True)
    total_qty = df_unique['逾期数量（万吨）'].sum()
    safe_total_qty = total_qty if total_qty > 0 else 1e-9

    doc = Document()
    set_page_margins(doc)
    init_styles(doc)

    regions = df_unique['大区'].dropna().unique() if '大区' in df_unique.columns else []
    is_hq = (len(regions) >= 2 or len(regions) == 0)

    def gen_tianshu(title):
        doc.add_paragraph(title, style='ChapterTitle')

        avg_days_val = (df_unique['逾期数量（万吨）'] * df_unique['逾期天数']).sum() / safe_total_qty if total_qty > 0 else 0
        avg_days = int(Decimal(str(round(avg_days_val, 6))).quantize(Decimal('1'), rounding=ROUND_HALF_UP))

        time_stats = df_unique.groupby('逾期天数分类').agg({'逾期金额（万元）': 'sum', '合同编号': 'count', '逾期数量（万吨）': 'sum'}).reindex(TIME_ORDER).fillna(0)

        p1 = doc.add_paragraph(style='NormalContent')
        run_avg1 = p1.add_run(f"平均逾期{avg_days}天，")
        set_font_mixed(run_avg1, 14.0, bold=True)
        run_avg2 = p1.add_run("周环比无数据")
        set_font_mixed(run_avg2, 14.0, bold=True)
        run_avg2.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run_avg3 = p1.add_run("。")
        set_font_mixed(run_avg3, 14.0, bold=True)

        max_qty_cat = time_stats['逾期数量（万吨）'].idxmax()
        valid_times = [t for t in TIME_ORDER if time_stats.loc[t, '逾期金额（万元）'] > 0]
        for i, t in enumerate(valid_times):
            amt = time_stats.loc[t, '逾期金额（万元）']
            qty = time_stats.loc[t, '逾期数量（万吨）']
            qty_ratio = qty / safe_total_qty * 100
            is_last = (i == len(valid_times) - 1)
            punctuation = "。" if is_last else "；"

            text_part = f"逾期天数在{t.replace('天', '')}天的，逾期数量{format_qty(qty)}万吨，占比{format_num(qty_ratio, is_percent=True)}，逾期金额{format_num(amt, 0, True)}万元"
            run_part = p1.add_run(text_part + punctuation)

            is_max = (t == max_qty_cat)
            set_font_mixed(run_part, 14.0, bold=is_max)
            if is_max:
                run_part.font.color.rgb = RGBColor(255, 0, 0)

        p1.paragraph_format.space_after = Pt(0)

        doc.add_paragraph('逾期销售提货分时间情况表', style='TableTitle')

        table1 = doc.add_table(rows=1, cols=5)
        table1.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths1 = [2.72, 2.95, 2.95, 2.95, 2.95]
        set_fixed_col_widths(table1, widths1, is_cm=True)

        headers1 = ['逾期时间', '逾期数量\n（万吨）', '逾期数量\n占比', '合同个数\n（笔）', '逾期金额\n（万元）']
        for i, h in enumerate(headers1):
            build_cell_text(table1.cell(0, i), h, bold=True)
            set_cell_background(table1.cell(0, i), 'D9D9D9')
        set_table_row_height(table1.rows[0], Cm(0.62).pt)
        set_repeat_table_header(table1.rows[0])

        for t in TIME_ORDER:
            amt = time_stats.loc[t, '逾期金额（万元）']
            if amt > 0:
                row_cells = table1.add_row().cells
                is_max = (t == max_qty_cat)
                qty_val = time_stats.loc[t, '逾期数量（万吨）']
                build_cell_text(row_cells[0], t, bold=is_max, is_max=is_max)
                build_cell_text(row_cells[1], format_qty(qty_val), bold=is_max, is_max=is_max)
                build_cell_text(row_cells[2], format_num(qty_val / safe_total_qty * 100, is_percent=True), bold=is_max, is_max=is_max)
                build_cell_text(row_cells[3], format_num(time_stats.loc[t, '合同编号'], 0, True), bold=is_max, is_max=is_max)
                build_cell_text(row_cells[4], format_num(amt, 0, True), bold=is_max, is_max=is_max)
                set_table_row_height(table1.rows[-1], Cm(0.48).pt)

        tot_cells1 = table1.add_row().cells
        build_cell_text(tot_cells1[0], '总计', bold=True)
        build_cell_text(tot_cells1[1], format_qty(time_stats['逾期数量（万吨）'].sum()), bold=True)
        build_cell_text(tot_cells1[2], '100%', bold=True)
        build_cell_text(tot_cells1[3], format_num(time_stats['合同编号'].sum(), 0, True), bold=True)
        build_cell_text(tot_cells1[4], total_amount_str, bold=True)
        set_cell_background(tot_cells1[0], 'D9E1F4')
        for cell in tot_cells1[1:]:
            set_cell_background(cell, 'DEEBF6')
        set_table_row_height(table1.rows[-1], Cm(0.48).pt)
        apply_table_borders(table1)

    def gen_yuanyin(title):
        doc.add_paragraph(title, style='ChapterTitle')
        r1_stats = df_unique.groupby('标准原因分类1').agg({'逾期金额（万元）': 'sum', '逾期数量（万吨）': 'sum'})
        r1_stats['逾期数量占比'] = r1_stats['逾期数量（万吨）'] / safe_total_qty * 100
        max_r1 = r1_stats['逾期金额（万元）'].idxmax() if not r1_stats.empty else None

        p2 = doc.add_paragraph(style='NormalContent')
        r1_texts_count = sum(1 for r1 in REASON1_ORDER if r1 in r1_stats.index and r1_stats.loc[r1, '逾期金额（万元）'] > 0)
        current_idx = 0

        for r1 in REASON1_ORDER:
            amt = r1_stats.loc[r1, '逾期金额（万元）'] if r1 in r1_stats.index else 0
            qty_r1 = r1_stats.loc[r1, '逾期数量（万吨）'] if r1 in r1_stats.index else 0
            qty_ratio_r1 = r1_stats.loc[r1, '逾期数量占比'] if r1 in r1_stats.index else 0
            if amt > 0:
                current_idx += 1
                if r1 == '一、客户原因/客户原因为主':
                    prefix = "主要由客户原因造成的"
                elif r1 == '二、我方原因/我方原因为主':
                    prefix = "主要由我方原因造成的"
                elif r1 == '三、既非我方原因也非对方原因':
                    prefix = "既非我方原因也非对方原因造成的"
                else:
                    prefix = ""

                if prefix:
                    text_part = f"{prefix}逾期数量{format_qty(qty_r1)}万吨，占比{format_num(qty_ratio_r1, is_percent=True)}，逾期金额{format_num(amt, 0, True)}万元"
                else:
                    text_part = f"其他原因逾期数量{format_qty(qty_r1)}万吨，占比{format_num(qty_ratio_r1, is_percent=True)}，逾期金额{format_num(amt, 0, True)}万元"

                is_last = (current_idx == r1_texts_count)
                punctuation = "；详情如下：" if is_last else "；"
                run_part = p2.add_run(text_part + punctuation)
                is_max = (r1 == max_r1)
                set_font_mixed(run_part, 14.0, bold=is_max)
                if is_max:
                    run_part.font.color.rgb = RGBColor(255, 0, 0)

        p2.paragraph_format.space_after = Pt(0)

        doc.add_paragraph('逾期销售提货分原因情况表', style='TableTitle')
        table2 = doc.add_table(rows=1, cols=5)
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths2 = [7.05, 1.99, 2.01, 1.76, 1.99]
        set_fixed_col_widths(table2, widths2, is_cm=True)

        headers2 = ['逾期原因分类', '逾期数量\n（万吨）', '逾期数量\n占比', '合同笔数\n（笔）', '逾期金额\n（万元）']
        for i, h in enumerate(headers2):
            build_cell_text(table2.cell(0, i), h, bold=True)
            set_cell_background(table2.cell(0, i), 'D9D9D9')
        set_table_row_height(table2.rows[0], Cm(0.3).pt)
        set_repeat_table_header(table2.rows[0])

        for r1 in REASON1_ORDER:
            r1_df = df_unique[df_unique['标准原因分类1'] == r1]
            r1_amt = r1_df['逾期金额（万元）'].sum()
            r1_qty = r1_df['逾期数量（万吨）'].sum()
            row_cells = table2.add_row().cells
            build_cell_text(row_cells[0], r1, align='left', bold=True)
            build_cell_text(row_cells[1], format_qty(r1_qty) if r1_amt>0 else '', align='right', bold=True)
            build_cell_text(row_cells[2], format_num(r1_qty/safe_total_qty*100, is_percent=True) if r1_amt>0 else '', align='right', bold=True)
            build_cell_text(row_cells[3], format_num(len(r1_df), 0, True) if r1_amt>0 else '', align='right', bold=True)
            build_cell_text(row_cells[4], format_num(r1_amt, 0, True) if r1_amt>0 else '', align='right', bold=True)
            for cell in row_cells:
                set_cell_background(cell, 'D9E1F4')
            set_table_row_height(table2.rows[-1], Cm(0.46).pt)

            r2_stats_amt = r1_df.groupby('原因分类2')['逾期金额（万元）'].sum()
            r2_stats_qty = r1_df.groupby('原因分类2')['逾期数量（万吨）'].sum()
            r2_list = REASON2_ORDER.get(r1, [])
            r2_val_map = {r2: r2_stats_amt.get(r2, 0) for r2 in r2_list}
            sorted_r2 = sorted(r2_list, key=lambda x: r2_val_map[x], reverse=True)

            for r2 in sorted_r2:
                r2_df = r1_df[r1_df['原因分类2'] == r2]
                r2_amt = r2_val_map[r2]
                r2_qty = r2_stats_qty.get(r2, 0)
                row_cells = table2.add_row().cells
                build_cell_text(row_cells[0], r2, align='left')
                if r2_amt > 0:
                    build_cell_text(row_cells[1], format_qty(r2_qty), align='right')
                    build_cell_text(row_cells[2], format_num(r2_qty/safe_total_qty*100, is_percent=True), align='right')
                    build_cell_text(row_cells[3], format_num(len(r2_df), 0, True), align='right')
                    build_cell_text(row_cells[4], format_num(r2_amt, 0, True), align='right')
                else:
                    build_cell_text(row_cells[1], '', align='right')
                    build_cell_text(row_cells[2], '', align='right')
                    build_cell_text(row_cells[3], '', align='right')
                    build_cell_text(row_cells[4], '', align='right')
                set_table_row_height(table2.rows[-1], Cm(0.46).pt)

        tot_cells2 = table2.add_row().cells
        tot_cells2[0].merge(tot_cells2[0])
        build_cell_text(tot_cells2[0], '总计', bold=True)
        build_cell_text(tot_cells2[1], format_qty(df_unique['逾期数量（万吨）'].sum()), bold=True)
        build_cell_text(tot_cells2[2], '100%', bold=True)
        build_cell_text(tot_cells2[3], format_num(len(df_unique), 0, True), bold=True)
        build_cell_text(tot_cells2[4], total_amount_str, bold=True)
        for cell in tot_cells2:
            set_cell_background(cell, 'D9E1F4')
        set_table_row_height(table2.rows[-1], Cm(0.46).pt)
        apply_table_borders(table2)

    def gen_pinzhong(title):
        doc.add_paragraph(title, style='ChapterTitle')
        variety_stats = df_unique.groupby('品种').agg({'逾期金额（万元）': 'sum', '合同编号': 'count', '逾期数量（万吨）': 'sum'}).sort_values(by='逾期金额（万元）', ascending=False)
        variety_stats['逾期数量占比'] = variety_stats['逾期数量（万吨）'] / safe_total_qty * 100
        max_v = variety_stats['逾期金额（万元）'].idxmax() if not variety_stats.empty else None

        p3 = doc.add_paragraph(style='NormalContent')
        v_count = len(variety_stats)
        for i, v in enumerate(variety_stats.index):
            v_amt = variety_stats.loc[v, '逾期金额（万元）']
            v_qty = variety_stats.loc[v, '逾期数量（万吨）']
            v_qty_ratio = variety_stats.loc[v, '逾期数量占比']
            text_part = f"{v}逾期数量{format_qty(v_qty)}万吨，占比{format_num(v_qty_ratio, is_percent=True)}，逾期金额{format_num(v_amt, 0, True)}万元"
            is_last = (i == v_count - 1)
            punctuation = "。详情如下：" if is_last else "；"
            run_part = p3.add_run(text_part + punctuation)
            is_max = (v == max_v)
            set_font_mixed(run_part, 14.0, bold=is_max)
            if is_max:
                run_part.font.color.rgb = RGBColor(255, 0, 0)

        p3.paragraph_format.space_after = Pt(0)

        doc.add_paragraph('逾期销售提货分品种情况表', style='TableTitle')
        table3 = doc.add_table(rows=1, cols=5)
        table3.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths3 = [2.72, 2.95, 2.95, 2.95, 2.95]
        set_fixed_col_widths(table3, widths3, is_cm=True)

        headers3 = ['品种', '逾期数量\n（万吨）', '逾期数量\n占比', '合同笔数\n（笔）', '逾期金额\n（万元）']
        for i, h in enumerate(headers3):
            build_cell_text(table3.cell(0, i), h, bold=True)
            set_cell_background(table3.cell(0, i), 'D9D9D9')
        set_table_row_height(table3.rows[0], Cm(0.62).pt)
        set_repeat_table_header(table3.rows[0])

        for v in variety_stats.index:
            row_cells = table3.add_row().cells
            v_amt = variety_stats.loc[v, '逾期金额（万元）']
            v_qty = variety_stats.loc[v, '逾期数量（万吨）']
            v_qty_ratio = variety_stats.loc[v, '逾期数量占比']
            is_max = (v == max_v)
            build_cell_text(row_cells[0], v, bold=is_max, is_max=is_max)
            build_cell_text(row_cells[1], format_qty(v_qty), bold=is_max, is_max=is_max)
            build_cell_text(row_cells[2], format_num(v_qty_ratio, is_percent=True), bold=is_max, is_max=is_max)
            build_cell_text(row_cells[3], format_num(variety_stats.loc[v, '合同编号'], 0, True), bold=is_max, is_max=is_max)
            build_cell_text(row_cells[4], format_num(v_amt, 0, True), bold=is_max, is_max=is_max)
            set_table_row_height(table3.rows[-1], Cm(0.48).pt)

        tot_cells3 = table3.add_row().cells
        tot_cells3[0].merge(tot_cells3[0])
        build_cell_text(tot_cells3[0], '总计', bold=True)
        build_cell_text(tot_cells3[1], format_qty(df_unique['逾期数量（万吨）'].sum()), bold=True)
        build_cell_text(tot_cells3[2], '100%', bold=True)
        build_cell_text(tot_cells3[3], format_num(len(df_unique), 0, True), bold=True)
        build_cell_text(tot_cells3[4], total_amount_str, bold=True)
        set_cell_background(tot_cells3[0], 'D9E1F4')
        for cell in tot_cells3[1:]:
            set_cell_background(cell, 'DEEBF6')
        set_table_row_height(table3.rows[-1], Cm(0.48).pt)
        apply_table_borders(table3)

    def gen_kehu(title, is_hq_mode=False):
        total_qty = df_unique['逾期数量（万吨）'].sum()
        safe_total_qty = total_qty if total_qty > 0 else 1e-9

        def get_cust_type(row):
            grp = str(row.get('所属集团', '')).strip()
            intr = str(row.get('集团内部客户', '')).strip()
            if grp and '中粮集团' not in grp and grp != 'nan': return '战略大客户', grp
            if intr and intr != 'nan': return '集团内部客户', intr
            return '中小客户', str(row.get('客户名称', '')).strip()

        df_unique['客户大类'], df_unique['展示客户名'] = zip(*df_unique.apply(get_cust_type, axis=1))
        c_stats = df_unique.groupby('客户大类')['逾期数量（万吨）'].sum().fillna(0)
        strat_total = c_stats.get('战略大客户', 0)
        mid_total = c_stats.get('中小客户', 0)
        int_total = c_stats.get('集团内部客户', 0)
        strat_cnt = df_unique[df_unique['客户大类'] == '战略大客户']['展示客户名'].nunique()
        mid_cnt = df_unique[df_unique['客户大类'] == '中小客户']['展示客户名'].nunique()
        int_cnt = df_unique[df_unique['客户大类'] == '集团内部客户']['展示客户名'].nunique()
        total_customers = strat_cnt + mid_cnt + int_cnt

        doc.add_paragraph(title, style='ChapterTitle')
        max_c_type = c_stats.idxmax() if not c_stats.empty else None

        p4 = doc.add_paragraph(style='NormalContent')
        run_base1 = p4.add_run(f"逾期提货客户共{total_customers}家，")
        set_font_mixed(run_base1, 14.0, bold=True)
        run_base2 = p4.add_run("周环比无数据")
        set_font_mixed(run_base2, 14.0, bold=True)
        run_base2.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run_base3 = p4.add_run("。包括")
        set_font_mixed(run_base3, 14.0, bold=True)

        c_parts = [
            ('战略大客户', f"{strat_cnt}家战略大客户共逾期{format_num(strat_total, 2)}万吨", "，"),
            ('中小客户', f"{mid_cnt}家中小客户共逾期{format_num(mid_total, 2)}万吨", "，"),
            ('集团内部客户', f"{int_cnt}家集团内部客户共逾期{format_num(int_total, 2)}万吨", "。具体情况如下表：")
        ]
        for c_type, text_part, punct in c_parts:
            run_part = p4.add_run(text_part + punct)
            is_max = (c_type == max_c_type and c_stats.get(c_type, 0) > 0)
            set_font_mixed(run_part, 14.0, bold=is_max)
            if is_max:
                run_part.font.color.rgb = RGBColor(255, 0, 0)

        p4.paragraph_format.space_after = Pt(0)

        # 中粮贸易模式：构建客户→逾期原因映射（按数量降序）
        customer_reasons = {}
        if is_hq_mode:
            reason_agg = df_unique.groupby(['展示客户名', '原因分类2'])['逾期数量（万吨）'].sum().reset_index()
            for name, group in reason_agg.groupby('展示客户名'):
                sorted_reasons = group.sort_values('逾期数量（万吨）', ascending=False)
                lines = []
                for _, rr in sorted_reasons.iterrows():
                    r_text = str(rr['原因分类2']) if pd.notna(rr['原因分类2']) else ''
                    r_qty = rr['逾期数量（万吨）']
                    if r_text:
                        lines.append(f"{r_text}{format_qty(r_qty)}万吨")
                customer_reasons[name] = '；'.join(lines) if lines else ''

        doc.add_paragraph('逾期销售提货分客户明细表', style='TableTitle')

        p_unit = doc.add_paragraph()
        p_unit.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_unit.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p_unit.paragraph_format.line_spacing = Pt(12)
        p_unit.paragraph_format.space_before = Pt(0)
        p_unit.paragraph_format.space_after = Pt(0)
        run_unit = p_unit.add_run('单位：万吨')
        set_font_mixed(run_unit, 9.0, bold=False, east_asia='微软雅黑', ascii_font='Times New Roman')

        if is_hq_mode:
            # ===== 中粮贸易：7列表格 + 逾期原因列 =====
            table4 = doc.add_table(rows=1, cols=7)
            table4.alignment = WD_TABLE_ALIGNMENT.CENTER
            widths4 = [0.77, 2.58, 1.57, 0.91, 1.16, 1.26, 7.1]
            set_fixed_col_widths(table4, widths4, is_cm=True)

            headers4 = ['序号', '客户名称/所属集团', '品种', '最长逾期天数', '逾期数量', '逾期数量占比', '逾期原因']
            for i, h in enumerate(headers4):
                build_cell_text(table4.cell(0, i), h, bold=True)
                set_cell_background(table4.cell(0, i), 'D9D9D9')
            set_table_row_height(table4.rows[0], Cm(0.73).pt)
            set_repeat_table_header(table4.rows[0])

            RED_KEYWORDS = ['我方其他原因（需详细说明）', '客户其他原因（需详细说明）']

            def write_reason_cell(cell, reason_text):
                cell.text = ''
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                p.paragraph_format.line_spacing = Pt(10)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if not reason_text:
                    return
                pattern = '(' + '|'.join(re.escape(k) for k in RED_KEYWORDS) + ')'
                parts = re.split(pattern, reason_text)
                for part in parts:
                    if not part:
                        continue
                    run = p.add_run(part)
                    is_red = part in RED_KEYWORDS
                    set_font_mixed(run, 9.0, bold=False, east_asia='微软雅黑', ascii_font='Times New Roman')
                    if is_red:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            def add_cust_rows_hq(c_type, subtotal_name, start_idx):
                sub_df = df_unique[df_unique['客户大类'] == c_type]
                if sub_df.empty: return start_idx
                agg_df = sub_df.groupby('展示客户名').agg({'逾期数量（万吨）': 'sum', '品种': lambda x: '、'.join(x.dropna().astype(str).unique()), '逾期天数': 'max'}).sort_values(by='逾期数量（万吨）', ascending=False).reset_index()
                for _, row in agg_df.iterrows():
                    cells = table4.add_row().cells
                    cust_name = row['展示客户名']
                    build_cell_text(cells[0], start_idx)
                    build_cell_text(cells[1], cust_name)
                    build_cell_text(cells[2], row.get('品种', ''))
                    build_cell_text(cells[3], format_num(row.get('逾期天数', 0), 0, True))
                    build_cell_text(cells[4], format_qty(row['逾期数量（万吨）']))
                    ratio = (row['逾期数量（万吨）'] / safe_total_qty * 100) if safe_total_qty > 0 else 0
                    build_cell_text(cells[5], format_num(ratio, is_percent=True))
                    reason_text = customer_reasons.get(cust_name, '')
                    write_reason_cell(cells[6], reason_text)
                    set_table_row_height(table4.rows[-1], Cm(0.44).pt)
                    start_idx += 1
                sub_cells = table4.add_row().cells
                sub_cells[0].merge(sub_cells[3])
                build_cell_text(sub_cells[0], subtotal_name, bold=True)
                build_cell_text(sub_cells[4], format_qty(agg_df['逾期数量（万吨）'].sum()), bold=True)
                sub_ratio = (agg_df['逾期数量（万吨）'].sum() / safe_total_qty * 100) if safe_total_qty > 0 else 0
                build_cell_text(sub_cells[5], format_num(sub_ratio, is_percent=True), bold=True)
                build_cell_text(sub_cells[6], '', bold=True)
                set_cell_background(sub_cells[6], 'D9D9D9')
                for c in [sub_cells[0], sub_cells[4], sub_cells[5]]: set_cell_background(c, 'D9D9D9')
                set_table_row_height(table4.rows[-1], Cm(0.44).pt)
                return start_idx

            idx = 1
            idx = add_cust_rows_hq('战略大客户', '战略客户小计', idx)
            idx = add_cust_rows_hq('中小客户', '中小客户小计', idx)
            idx = add_cust_rows_hq('集团内部客户', '集团内部客户小计', idx)

            tot_cells4 = table4.add_row().cells
            tot_cells4[0].merge(tot_cells4[3])
            build_cell_text(tot_cells4[0], '汇总', bold=True)
            build_cell_text(tot_cells4[4], format_qty(df_unique['逾期数量（万吨）'].sum()), bold=True)
            build_cell_text(tot_cells4[5], '100%', bold=True)
            build_cell_text(tot_cells4[6], '', bold=True)
            set_cell_background(tot_cells4[6], 'DEEBF6')
            for c in [tot_cells4[0], tot_cells4[4], tot_cells4[5]]: set_cell_background(c, 'DEEBF6')
            set_table_row_height(table4.rows[-1], Cm(0.44).pt)
            apply_table_borders(table4)
            # 强制锁定所有单元格列宽
            set_fixed_col_widths(table4, widths4, is_cm=True)
        else:
            # ===== 大区：原有6列表格（保持不变） =====
            table4 = doc.add_table(rows=1, cols=6)
            table4.alignment = WD_TABLE_ALIGNMENT.CENTER
            widths4 = [1.32, 6.14, 2.88, 1.77, 1.84, 1.73]
            set_fixed_col_widths(table4, widths4, is_cm=True)

            headers4 = ['序号', '客户名称/所属集团', '品种', '最长逾期天数', '逾期数量', '逾期数量占比']
            for i, h in enumerate(headers4):
                build_cell_text(table4.cell(0, i), h, bold=True)
                set_cell_background(table4.cell(0, i), 'D9D9D9')
            set_table_row_height(table4.rows[0], Cm(0.73).pt)
            set_repeat_table_header(table4.rows[0])

            def add_cust_rows(c_type, subtotal_name, start_idx):
                sub_df = df_unique[df_unique['客户大类'] == c_type]
                if sub_df.empty: return start_idx
                agg_df = sub_df.groupby('展示客户名').agg({'逾期数量（万吨）': 'sum', '品种': lambda x: '、'.join(x.dropna().astype(str).unique()), '逾期天数': 'max'}).sort_values(by='逾期数量（万吨）', ascending=False).reset_index()
                for _, row in agg_df.iterrows():
                    cells = table4.add_row().cells
                    build_cell_text(cells[0], start_idx)
                    build_cell_text(cells[1], row['展示客户名'])
                    build_cell_text(cells[2], row.get('品种', ''))
                    build_cell_text(cells[3], format_num(row.get('逾期天数', 0), 0, True))
                    build_cell_text(cells[4], format_qty(row['逾期数量（万吨）']))
                    ratio = (row['逾期数量（万吨）'] / safe_total_qty * 100) if safe_total_qty > 0 else 0
                    build_cell_text(cells[5], format_num(ratio, is_percent=True))
                    set_table_row_height(table4.rows[-1], Cm(0.44).pt)
                    start_idx += 1
                sub_cells = table4.add_row().cells
                sub_cells[0].merge(sub_cells[3])
                build_cell_text(sub_cells[0], subtotal_name, bold=True)
                build_cell_text(sub_cells[4], format_qty(agg_df['逾期数量（万吨）'].sum()), bold=True)
                sub_ratio = (agg_df['逾期数量（万吨）'].sum() / safe_total_qty * 100) if safe_total_qty > 0 else 0
                build_cell_text(sub_cells[5], format_num(sub_ratio, is_percent=True), bold=True)
                for c in [sub_cells[0], sub_cells[4], sub_cells[5]]: set_cell_background(c, 'D9D9D9')
                set_table_row_height(table4.rows[-1], Cm(0.44).pt)
                return start_idx

            idx = 1
            idx = add_cust_rows('战略大客户', '战略客户小计', idx)
            idx = add_cust_rows('中小客户', '中小客户小计', idx)
            idx = add_cust_rows('集团内部客户', '集团内部客户小计', idx)

            tot_cells4 = table4.add_row().cells
            tot_cells4[0].merge(tot_cells4[3])
            build_cell_text(tot_cells4[0], '汇总', bold=True)
            build_cell_text(tot_cells4[4], format_qty(df_unique['逾期数量（万吨）'].sum()), bold=True)
            build_cell_text(tot_cells4[5], '100%', bold=True)
            for c in [tot_cells4[0], tot_cells4[4], tot_cells4[5]]: set_cell_background(c, 'DEEBF6')
            set_table_row_height(table4.rows[-1], Cm(0.44).pt)
            apply_table_borders(table4)

    def gen_empty_section(title):
        doc.add_paragraph(title, style='ChapterTitle')
        doc.add_paragraph('', style='NormalContent')

    def gen_appendices(is_hq_mode):
        # ----- 附表 -----
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = Cm(377194.0 / 12700.0)
        new_section.page_height = Cm(266710.5 / 12700.0)
        new_section.left_margin = Cm(40322.44 / 12700.0)
        new_section.right_margin = Cm(40322.44 / 12700.0)
        new_section.top_margin = Cm(32257.95 / 12700.0)
        new_section.bottom_margin = Cm(32257.95 / 12700.0)
        new_section.header_distance = Cm(19063.55 / 12700.0)
        new_section.footer_distance = Cm(22222.14 / 12700.0)

        def create_appendix(title, df_subset, table_idx):
            doc.add_paragraph(title, style='AppendixTitle')
            p_app_unit = doc.add_paragraph()
            p_app_unit.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_app_unit.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p_app_unit.paragraph_format.line_spacing = Pt(28)
            p_app_unit.paragraph_format.space_before = Pt(0)
            p_app_unit.paragraph_format.space_after = Pt(0)
            run_app_unit = p_app_unit.add_run('单位：万吨、元/吨、万元')
            set_font_mixed(run_app_unit, 14.0, bold=False, east_asia='仿宋_GB2312', ascii_font='Times New Roman')

            table = doc.add_table(rows=1, cols=15)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            if table_idx == 5: widths_app = [9386.17, 15792.96, 25313.53, 22154.94, 12656.77, 15792.96, 12656.77, 12656.77, 18055.49, 18077.89, 117674.32, 39224.77, 12656.77, 13194.4, 14157.66]
            elif table_idx == 6: widths_app = [9430.97, 15860.16, 25403.14, 19959.61, 12544.76, 16128.98, 13127.19, 12253.54, 17921.08, 17921.08, 118771.98, 38104.7, 12701.57, 12947.98, 12947.98]
            else: widths_app = [11021.47, 16128.98, 24775.9, 20430.04, 11962.32, 15815.36, 13440.81, 11962.32, 18212.3, 18212.3, 111021.11, 45385.14, 13127.19, 12253.54, 14516.08]
            set_fixed_col_widths(table, widths_app)

            headers = ['序号', '经营部', '客户名称', '交货结束日期', '合同\n数量', '合同\n单价', '逾期\n天数', '逾期数量', '原因\n分类1', '原因\n分类2', '具体逾期原因', '解决方案及解决时间', '责任人', '上级领导', '是否为赊销合同']
            for i, h in enumerate(headers):
                cell = table.cell(0, i)
                build_cell_text(cell, h, bold=True, is_appendix=True)
                set_cell_background(cell, 'D9D9D9')
            set_table_row_height(table.rows[0], Cm(0.9).pt)
            set_repeat_table_header(table.rows[0])

            if not df_subset.empty:
                df_subset = df_subset.sort_values(by='逾期天数', ascending=False).reset_index(drop=True)
                leader_map = {'珠三角': '黄旭东', '福建': '肖灿', '广西': '丁峰', '海南': '宋永伍', '粤西': '张文韬'}
                for i, row in df_subset.iterrows():
                    cells = table.add_row().cells
                    dept_val = str(row.get('经营部', '')) if pd.notna(row.get('经营部')) else ""
                    dept_clean = dept_val.replace('经营部', '')
                    leader_val = leader_map.get(dept_clean, str(row.get('上级领导', '')) if pd.notna(row.get('上级领导')) else "")

                    reason_detail = str(row.get('具体逾期原因', ''))
                    match = re.search(r'(?:业务|责任|负责|联系)(?:人|人员|员)?(?:[:：，,])?([一-龥]{2,3})(?:[\d\s,。、，.!?！？]|$)', reason_detail)
                    person_val = match.group(1) if match else ''

                    build_cell_text(cells[0], i+1, is_appendix=True)
                    build_cell_text(cells[1], dept_clean, align='center', is_appendix=True)
                    build_cell_text(cells[2], row.get('客户名称', ''), align='center', is_appendix=True)
                    dt = row.get('交货结束日期', '')
                    build_cell_text(cells[3], str(dt)[:10] if pd.notna(dt) else "", is_appendix=True)

                    # ====== 核心修复点 1：抓取键更换为"合同数量(万吨)" ======
                    build_cell_text(cells[4], format_num(row.get('合同数量(万吨)', ''), 2), is_appendix=True)

                    build_cell_text(cells[5], format_num(row.get('合同单价', ''), 0, True), is_appendix=True)
                    build_cell_text(cells[6], format_num(row.get('逾期天数', ''), 0, True), is_appendix=True)
                    build_cell_text(cells[7], format_qty(row.get('逾期数量（万吨）', '')), is_appendix=True)
                    build_cell_text(cells[8], row.get('原因分类1', ''), align='left', is_appendix=True)
                    build_cell_text(cells[9], row.get('原因分类2', ''), align='left', is_appendix=True)
                    build_cell_text(cells[10], reason_detail, align='left', is_appendix=True)
                    sol = str(row.get('解决方案', '')) if pd.notna(row.get('解决方案')) else ""
                    sol_dt = str(row.get('预计完成日期', ''))[:10] if pd.notna(row.get('预计完成日期')) else ""
                    build_cell_text(cells[11], f"{sol}{sol_dt}", align='left', is_appendix=True)
                    build_cell_text(cells[12], person_val, is_appendix=True)
                    build_cell_text(cells[13], leader_val, is_appendix=True)

                    credit_val = "是" if pd.notna(row.get('销售类型')) and str(row.get('销售类型')).strip() == '赊销' else ""
                    build_cell_text(cells[14], credit_val, is_appendix=True)
                    set_table_row_height(table.rows[-1], Cm(0.6).pt)
            else:
                empty_row = table.add_row().cells
                for cell in empty_row: build_cell_text(cell, '', is_appendix=True)
                set_table_row_height(table.rows[-1], Cm(0.6).pt)
            apply_table_borders(table)

        df_app1 = df[df['逾期天数'] >= 50]
        title1 = '附表1：逾期60天以上的销售合同情况' if is_hq_mode else '逾期60天以上的销售合同情况'
        create_appendix(title1, df_app1, 5)
        if '是否重点关注' in df.columns: df_app2 = df[df['是否重点关注'].astype(str).str.contains('重点关注', na=False)]
        else: df_app2 = pd.DataFrame()
        title2 = '附表2：其他需要重点关注的销售合同情况（尤其是有潜在风险的）' if is_hq_mode else '其他需要重点关注的销售合同情况（尤其是有潜在风险的）'
        create_appendix(title2, df_app2, 6)
        if '是否严重逾期' in df.columns: df_app3 = df[df['是否严重逾期'].astype(str).str.contains('严重逾期', na=False)]
        else: df_app3 = pd.DataFrame()
        title3 = '附表3：严重逾期销售合同情况' if is_hq_mode else '严重逾期销售合同情况'
        create_appendix(title3, df_app3, 7)

    def _gen_overdue_summary(is_hq_mode):
        """生成逾期销售总结概述段落（纯新增功能）"""
        _today = datetime.datetime.now()
        _days_to_wed = (_today.weekday() - 2) % 7
        if _days_to_wed == 0:
            _days_to_wed = 7
        _last_wed = _today - datetime.timedelta(days=_days_to_wed)
        _cutoff_date_str = f"{_last_wed.month}月{_last_wed.day}日"

        _total_qty = df_unique['逾期数量（万吨）'].sum()
        _safe_total = _total_qty if _total_qty > 0 else 1e-9

        if is_hq_mode:
            _group_col = '大区'
        else:
            _group_col = '经营部'

        if _group_col not in df_unique.columns or df_unique[_group_col].dropna().empty:
            return

        _group_stats = df_unique.groupby(_group_col)['逾期数量（万吨）'].sum().sort_values(ascending=False)
        if len(_group_stats) == 0:
            return
        _top1_name = str(_group_stats.index[0])
        _top1_qty = _group_stats.iloc[0]
        _top1_pct = int(round(_top1_qty / _safe_total * 100)) if _safe_total > 0 else 0

        _top1_df = df_unique[df_unique[_group_col] == _top1_name]
        if '标准原因分类1' not in _top1_df.columns:
            return
        _r1_stats = _top1_df.groupby('标准原因分类1')['逾期数量（万吨）'].sum().sort_values(ascending=False)
        if len(_r1_stats) == 0:
            return
        _top1_r1_name = str(_r1_stats.index[0])
        _top1_r1_qty = _r1_stats.iloc[0]
        _top1_r1_pct = int(round(_top1_r1_qty / _top1_qty * 100)) if _top1_qty > 0 else 0

        _top1_r1_df = _top1_df[_top1_df['标准原因分类1'] == _top1_r1_name]
        _r2_stats = _top1_r1_df.groupby('原因分类2')['逾期数量（万吨）'].sum().sort_values(ascending=False)

        # 原因分类1名称简化映射（仅展示用，不改源数据）
        def _simplify_r1_name(name):
            _s = str(name)
            if '一、客户原因' in _s or '客户原因为主' in _s:
                return '客户原因'
            if '二、我方原因' in _s or '我方原因为主' in _s:
                return '我方原因'
            if '三、既非我方原因也非对方原因' in _s:
                return '其他原因'
            return _s

        _simplified_r1 = _simplify_r1_name(_top1_r1_name)
        _RED_R2_KEYWORDS = ['我方其他原因（需详细说明）', '客户其他原因（需详细说明）']

        # 总结概述段落（连续一整段，默认黑色不加粗，仅特殊标记有格式）
        _p = doc.add_paragraph(style='NormalContent')

        _run = _p.add_run(f"截至{_cutoff_date_str}，中粮贸易逾期销售提货数量共")
        set_font_mixed(_run, 14.0, bold=False)

        # 总吨位+单位（标红加粗 ①）
        _run = _p.add_run(f"{format_qty(_total_qty)}万吨")
        set_font_mixed(_run, 14.0, bold=True)
        _run.font.color.rgb = RGBColor(255, 0, 0)

        _run = _p.add_run("，")
        set_font_mixed(_run, 14.0, bold=False)

        # 环比增量描述（标红加粗 ② + 亮黄色底纹）
        _run = _p.add_run("环比无数据")
        set_font_mixed(_run, 14.0, bold=True)
        _run.font.color.rgb = RGBColor(255, 0, 0)
        _run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        _run = _p.add_run("。本周，")
        set_font_mixed(_run, 14.0, bold=False)

        # Top1主体逾期提货量完整描述（标红加粗 ③）
        _run = _p.add_run(f"{_top1_name}逾期提货数量{format_qty(_top1_qty)}万吨")
        set_font_mixed(_run, 14.0, bold=True)
        _run.font.color.rgb = RGBColor(255, 0, 0)

        _run = _p.add_run("，占中粮贸易整体")
        set_font_mixed(_run, 14.0, bold=False)

        _run = _p.add_run(f"{format_qty(_total_qty)}万吨")
        set_font_mixed(_run, 14.0, bold=False)

        _run = _p.add_run("的")
        set_font_mixed(_run, 14.0, bold=False)

        # 占比百分比（标红加粗）
        _run = _p.add_run(f"{_top1_pct}%")
        set_font_mixed(_run, 14.0, bold=True)
        _run.font.color.rgb = RGBColor(255, 0, 0)

        _run = _p.add_run("。")
        set_font_mixed(_run, 14.0, bold=False)

        _run = _p.add_run(f"{_top1_name}{format_qty(_top1_qty)}万吨")
        set_font_mixed(_run, 14.0, bold=False)

        _run = _p.add_run("逾期提货量中，")
        set_font_mixed(_run, 14.0, bold=False)

        # 原因分类1吨位+原因完整分句（仅加粗 ①）
        _run = _p.add_run(f"{format_qty(_top1_r1_qty)}万吨逾期是因{_simplified_r1}")
        set_font_mixed(_run, 14.0, bold=True)

        _run = _p.add_run("，占比约")
        set_font_mixed(_run, 14.0, bold=False)

        # 原因分类1占比百分比数值（仅加粗 ②）
        _run = _p.add_run(f"{_top1_r1_pct}%")
        set_font_mixed(_run, 14.0, bold=True)

        _run = _p.add_run("，包括：")
        set_font_mixed(_run, 14.0, bold=False)

        # 原因分类2逐项拼接
        _r2_items = list(_r2_stats.items())
        for _idx, (_r2_name, _r2_qty) in enumerate(_r2_items):
            _r2_name_str = str(_r2_name) if pd.notna(_r2_name) else ''

            # 数值+万吨（仅加粗 ③）
            _run = _p.add_run(f"{format_qty(_r2_qty)}万吨")
            set_font_mixed(_run, 14.0, bold=True)

            _run = _p.add_run("是因")
            set_font_mixed(_run, 14.0, bold=False)

            # 原因名称：匹配关键词则仅标红不加粗，否则黑色不加粗
            _run = _p.add_run(_r2_name_str)
            if _r2_name_str in _RED_R2_KEYWORDS:
                set_font_mixed(_run, 14.0, bold=False)
                _run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                set_font_mixed(_run, 14.0, bold=False)

            # 分隔符
            _sep = "。" if _idx == len(_r2_items) - 1 else "；"
            _run = _p.add_run(_sep)
            set_font_mixed(_run, 14.0, bold=False)

        _run = _p.add_run("具体情况如下：")
        set_font_mixed(_run, 14.0, bold=False)

        _p.paragraph_format.space_after = Pt(0)

    if is_hq:
        # 中粮贸易专属报告顺序
        _gen_overdue_summary(is_hq_mode=True)
        gen_empty_section('（一）逾期销售数量')
        gen_kehu('（二）逾期销售分客户', True)
        gen_yuanyin('（三）逾期销售原因')
        gen_tianshu('（四）逾期销售天数')
        gen_pinzhong('（五）逾期销售分品种')
        gen_empty_section('（六）逾期销售金额')
        gen_appendices(True)
    else:
        # 大区报告顺序与修改后的标题
        _gen_overdue_summary(is_hq_mode=False)
        gen_tianshu('3、逾期销售天数')
        gen_appendices(False)

        # 核心恢复：由于附表是横向的，大区报告还需要继续写竖向内容，必须添加竖向分节符并恢复原有页面边距
        new_section_port = doc.add_section(WD_SECTION.NEW_PAGE)
        new_section_port.orientation = WD_ORIENT.PORTRAIT
        new_section_port.page_width = Cm(266710.5 / 12700.0)
        new_section_port.page_height = Cm(377194.0 / 12700.0)
        new_section_port.left_margin = Cm(3.0)
        new_section_port.right_margin = Cm(3.0)
        new_section_port.top_margin = Cm(2.54)
        new_section_port.bottom_margin = Cm(2.54)
        new_section_port.header_distance = Cm(19063.55 / 12700.0)
        new_section_port.footer_distance = Cm(22222.14 / 12700.0)

        gen_yuanyin('4、逾期销售原因')
        gen_pinzhong('5、逾期销售分品种')
        gen_kehu('6、逾期销售分客户')

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def beautify_excel_for_io(wb):
    ws = wb.active
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    header_font = Font(name='微软雅黑', bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    body_font = Font(name='微软雅黑', size=10)
    fill_batch = PatternFill(start_color=COLOR_BATCH, end_color=COLOR_BATCH, fill_type="solid")
    fill_once = PatternFill(start_color=COLOR_ONCE, end_color=COLOR_ONCE, fill_type="solid")

    source_col_idx = None
    for cell in ws[1]:
        if cell.value == "_Data_Source":
            source_col_idx = cell.column
            break

    for row in ws.iter_rows():
        current_fill = None
        if source_col_idx and row[0].row > 1:
            source_val = row[source_col_idx - 1].value
            if source_val == 'batch': current_fill = fill_batch
            elif source_val == 'once': current_fill = fill_once
        
        ws.row_dimensions[row[0].row].height = 20
        for cell in row:
            cell.border = thin_border
            cell.font = body_font
            cell.alignment = Alignment(vertical='center', wrap_text=False)
            if cell.row > 1 and current_fill: cell.fill = current_fill
            if cell.row == 1:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center', vertical='center')

    if source_col_idx: ws.delete_cols(source_col_idx)

    wide_cols = ["合同编号", "客户名称", "具体逾期原因", "最新进展", "解决方案", "集团内部客户"]
    for col in ws.columns:
        try:
            header_val = str(col[0].value).strip()
            width = 14
            if any(k in header_val for k in wide_cols):
                width = 30
                for cell in col:
                    if cell.row > 1: cell.alignment = Alignment(vertical='center', wrap_text=True)
            elif "日期" in header_val or "品种" in header_val: width = 12
            elif "重点关注" in header_val: width = 15
            ws.column_dimensions[col[0].column_letter].width = width
        except: pass
    return wb

def process_overdue_sales(uploaded_files, need_report=False):
    """处理逾期销售数据，内部自动识别分批次/一次性文件"""
    logs = []
    header_keywords = ["大区", "经营部", "合同编号", "客户名称"]
    date_columns = ["合同签订日期", "交货开始日期", "交货结束日期", "预计完成日期"]
    all_numeric_columns = ["合同数量", "合同单价", "合同金额", "调整后逾期销售金额", "逾期天数I", "逾期天数II", "逾期天数III", "逾期天数IV", "逾期天数V", "逾期天数VI"]
    special_int_columns = ["逾期天数I", "逾期天数II", "逾期天数III", "逾期天数IV", "逾期天数V", "逾期天数VI"]
    # 分批次的标志性列名（用于表头校验）
    batch_indicator_columns = [f"逾期天数{roman}" for roman in ['I', 'II', 'III', 'IV', 'V', 'VI']]

    # ==========================================
    # 智能分类：将上传文件自动分为 batch_files 和 once_files
    # ==========================================
    batch_files = []
    once_files = []

    for f in uploaded_files:
        # Step 1: 文件名初筛 —— 检查文件名中是否包含"分批次"或"一次性"关键字
        fname = f.name if hasattr(f, 'name') else ''
        filename_is_batch = '分批次' in fname
        filename_is_once = '一次性' in fname

        # Step 2: 表头校验定性 —— 读取标题行，检查是否包含分批次的特定列名
        # （逾期天数I, 逾期天数II, ... 逾期天数VI）
        header_is_batch = False
        try:
            f.seek(0)
            df_head = pd.read_excel(f, header=None, nrows=5)
            # 遍历前5行，查找包含分批标志列的标题行
            for i, row in df_head.iterrows():
                row_values = [_normalize_col_name(x) for x in row.values if pd.notna(x)]
                row_text = ' '.join(row_values)
                # 检查是否同时包含多个分批标志列（至少命中2个即判定为分批次）
                match_count = sum(1 for col in batch_indicator_columns if col in row_text)
                if match_count >= 2:
                    header_is_batch = True
                    break
        except Exception:
            pass

        # 优先级原则：表头内容 > 文件名关键字
        # 如果表头包含分批标志列，则绝对锁定为分批次文件
        if header_is_batch:
            batch_files.append(f)
            logs.append(f"📎 {fname} → 识别为【分批次】文件（表头含逾期天数I-VI列）")
        # 如果文件名明确标注"分批次"，归入分批次
        elif filename_is_batch:
            batch_files.append(f)
            logs.append(f"📎 {fname} → 识别为【分批次】文件（文件名匹配）")
        # 如果文件名明确标注"一次性"，归入一次性
        elif filename_is_once:
            once_files.append(f)
            logs.append(f"📎 {fname} → 识别为【一次性】文件（文件名匹配）")
        # 默认归入一次性（不含分批次标志列且文件名无明确标注）
        else:
            once_files.append(f)
            logs.append(f"📎 {fname} → 识别为【一次性】文件（默认归类）")

    # ==========================================
    # 以下为原有的分批次/一次性分别处理逻辑（保持不变）
    # ==========================================
    df_batch = pd.DataFrame()
    if batch_files:
        df_list = []
        for f in batch_files:
            temp = locate_header_and_read(f, header_keywords)
            if temp is not None: df_list.append(temp)
        if df_list:
            temp_combined = pd.concat(df_list, ignore_index=True).drop_duplicates()
            df_batch = process_basic_columns(temp_combined, date_columns, all_numeric_columns, special_int_columns)
            col_map = {_normalize_col_name(c): c for c in df_batch.columns}
            calc_cols = [col_map[_normalize_col_name(c)] for c in special_int_columns if _normalize_col_name(c) in col_map]
            df_batch['逾期天数'] = df_batch[calc_cols].max(axis=1).fillna(0) if calc_cols else 0
            df_batch['_Data_Source'] = 'batch'

    df_once = pd.DataFrame()
    if once_files:
        df_list = []
        for f in once_files:
            temp = locate_header_and_read(f, header_keywords)
            if temp is not None: df_list.append(temp)
        if df_list:
            temp_combined = pd.concat(df_list, ignore_index=True).drop_duplicates()
            df_once = process_basic_columns(temp_combined, date_columns, all_numeric_columns)
            if '逾期天数' not in df_once.columns: df_once['逾期天数'] = 0
            df_once['_Data_Source'] = 'once'

    if df_batch.empty and df_once.empty:
        return None, None, None, ["❌ 未能从上传的文件中读取到有效数据，请检查文件格式。"]

    df_merged = pd.concat([df_batch, df_once], ignore_index=True)
    df_merged = merge_detail_variety_columns(df_merged)
    df_merged = process_variety_logic(df_merged)

    target_col = "逾期分类（业绩考核角度）"
    if target_col in df_merged.columns:
        s = df_merged[target_col].astype(str)
        cond = s.str.contains("A", na=False) & s.str.contains("超过交货结束日期", na=False) & s.str.contains("未完成交提货的", na=False)
        df_final = df_merged[cond].copy()
    else:
        df_final = df_merged.copy()

    group_map, internal_map, group_map_norm, internal_map_norm = get_customer_mappings()

    cols_to_fill = ['调整后逾期销售金额', '合同单价', '合同金额', '交货结束日期', '交货开始日期', '合同数量']
    for col in cols_to_fill:
        if col not in df_final.columns:
            df_final[col] = pd.NaT if '日期' in col else 0

    if '客户名称' in df_final.columns:
        df_final['客户名称_clean'] = df_final['客户名称'].astype(str).str.strip()
        # 括号归一化后的键名，用于兜底匹配（解决全角/半角、中文/英文括号不一致问题）
        df_final['客户名称_norm'] = df_final['客户名称_clean'].apply(_normalize_brackets)
        # 优先原始匹配，失败时回退到归一化匹配
        df_final['所属集团'] = df_final['客户名称_clean'].map(group_map).fillna(
            df_final['客户名称_norm'].map(group_map_norm)
        ).fillna("")
        df_final['集团内部客户'] = df_final.apply(
            lambda row: (
                internal_map.get(row['客户名称_clean'], "")
                or internal_map_norm.get(row['客户名称_norm'], "")
            ) if row['所属集团'] == '中粮集团' else "",
            axis=1
        )
        df_final.drop(columns=['客户名称_clean', '客户名称_norm'], inplace=True)
    else:
        df_final['所属集团'] = ""
        df_final['集团内部客户'] = ""

    bins = [-float('inf'), 10, 20, 30, 60, 90, float('inf')]
    labels = ["1-10天", "11-20天", "21-30天", "31-60天", "61-90天", "90天以上"]
    if '逾期天数' in df_final.columns:
        df_final['逾期天数分类'] = pd.cut(df_final['逾期天数'], bins=bins, labels=labels)

    df_final['合同单价_safe'] = df_final['合同单价'].replace(0, np.nan)
    df_final['逾期数量'] = df_final['调整后逾期销售金额'] / df_final['合同单价_safe'] / 10000
    df_final['逾期数量'] = df_final['逾期数量'].fillna(0)

    df_final['合同执行期(天数)'] = (df_final['交货结束日期'] - df_final['合同签订日期']).dt.days
    df_final['合同执行期(天数)'] = df_final['合同执行期(天数)'].fillna(1).replace(0, 1)
    
    ratio_days = (df_final['逾期天数'] / df_final['合同执行期(天数)']).fillna(0)
    ratio_amt = (df_final['调整后逾期销售金额'] / df_final['合同金额'].replace(0, np.nan)).fillna(0)
    df_final['是否严重逾期'] = np.where((ratio_days >= 0.5) & (ratio_amt >= 0.5), "严重逾期", "")

    df_final['合同数量'] = (df_final['合同数量'] / 10000).round(4)
    df_final['合同金额'] = (df_final['合同金额'] / 10000).round(2)
    df_final['调整后逾期销售金额'] = (df_final['调整后逾期销售金额'] / 10000).round(2)
    df_final['逾期数量'] = df_final['逾期数量'].round(4)

    cond_focus = (df_final['调整后逾期销售金额'] >= 500) & (df_final['逾期天数'] >= 10)
    df_final['是否重点关注'] = np.where(cond_focus, "重点关注", "")

    if '销售类型' not in df_final.columns: df_final['销售类型'] = ""

    df_final.sort_values(
        by=['逾期数量', '逾期天数', '是否严重逾期', '是否重点关注'],
        ascending=[False, False, False, False],
        inplace=True
    )

    col_rename_map = {
        "合同签订日期": "签订日期",
        "合同数量": "合同数量(万吨)",
        "合同金额": "合同金额（万元）",
        "逾期数量": "逾期数量（万吨）",
        "调整后逾期销售金额": "逾期金额（万元）",
        "逾期原因分类1（责任划分角度）": "原因分类1",
        "逾期原因分类2（责任划分角度）": "原因分类2",
        "当日最新进展": "最新进展"
    }
    df_final.rename(columns=col_rename_map, inplace=True)

    for col in ["签订日期", "交货结束日期", "预计完成日期"]:
        if col in df_final.columns:
            df_final[col] = df_final[col].dt.strftime('%Y-%m-%d')

    final_columns = [
        "大区", "经营部", "合同编号", "客户名称", "签订日期", "交货结束日期", "品种", 
        "合同数量(万吨)", "合同单价", "合同金额（万元）", "逾期天数", "逾期数量（万吨）", 
        "逾期金额（万元）", "原因分类1", "原因分类2", "具体逾期原因", "预计完成日期", 
        "解决方案", "最新进展", "是否严重逾期", "逾期天数分类", "所属集团", 
        "集团内部客户", "是否重点关注", "销售类型"
    ]
    
    for col in final_columns:
        if col not in df_final.columns: df_final[col] = ""

    output_cols_with_source = final_columns + ["_Data_Source"]
    df_output = df_final[output_cols_with_source]
    
    excel_io = io.BytesIO()
    with pd.ExcelWriter(excel_io, engine='openpyxl') as writer:
        df_output.to_excel(writer, index=False, sheet_name='逾期监控明细')
    excel_io.seek(0)
    
    wb = openpyxl.load_workbook(excel_io)
    wb = beautify_excel_for_io(wb)
    final_excel_io = io.BytesIO()
    wb.save(final_excel_io)
    final_excel_io.seek(0)
    
    logs.append("✅ 逾期销售监控明细表生成完成。")
    
    word_io = None
    collection_text = ""
    
    df_unique = df_output.drop_duplicates(subset=['合同编号']).copy()
    if '原因分类1' in df_unique.columns:
        df_unique['标准原因分类1'] = df_unique['原因分类1'].apply(map_reason1)

    try:
        collection_text = generate_collection_reminder(df_unique)
        logs.append("✅ 催收提醒内容生成成功。")
    except Exception as e:
        logs.append(f"⚠️ 催收提醒生成失败：{str(e)}")

    if need_report:
        try:
            word_io = generate_report(df_output, df_unique)
            logs.append("✅ 周报 Word 文档生成完成。")
        except Exception as e:
            logs.append(f"⚠️ 周报生成失败：{str(e)}")
            
    return final_excel_io, word_io, collection_text, logs
