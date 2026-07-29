# -*- coding: utf-8 -*-
"""
逾期赊销周报自动化处理 — CTMR Web 模块
从 RP.PY 移植并适配 BytesIO 内存流处理
"""
import io
import re
from datetime import date, datetime, timedelta

import openpyxl
from openpyxl.styles import Font
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 工具函数（原封保留）
# ============================================================

def safe_float(value):
    if value is None:
        return None
    try:
        return float(value)
    except (ValueError, TypeError):
        return None


def parse_number(text):
    if text is None:
        return 0.0
    cleaned = str(text).replace(",", "").replace("，", "").strip()
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def get_last_wednesday(today_date):
    days_since_wed = (today_date.weekday() - 2) % 7
    if days_since_wed == 0:
        days_since_wed = 7
    return today_date - timedelta(days=days_since_wed)


def find_sheet_by_keyword(wb, keywords):
    for sn in wb.sheetnames:
        for kw in keywords:
            if kw in sn:
                return wb[sn]
    return None


# ============================================================
# Word 样式工具函数（原封保留）
# ============================================================

def set_font_mixed(run_or_style, size_pt, bold=False, east_asia='仿宋_GB2312', ascii_font='Times New Roman'):
    run_or_style.font.size = Pt(size_pt)
    run_or_style.font.name = ascii_font
    rPr = run_or_style._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    run_or_style.font.bold = bold


def set_page_margins(doc):
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)


def build_cell_text(cell, text, align='center', bold=False, cn_size=9.0, en_size=10.0):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(10)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    text_str = str(text) if text is not None else ""
    parts = re.split(r'([a-zA-Z0-9.,%+\-()/]+)', text_str)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part)
        if re.match(r'^[a-zA-Z0-9.,%+\-()/]+$', part):
            set_font_mixed(run, en_size, bold, '微软雅黑', 'Times New Roman')
        else:
            set_font_mixed(run, cn_size, bold, '微软雅黑', 'Times New Roman')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_background(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


def apply_table_borders(table):
    """为表格添加全部边框 + 固定布局（WPS 兼容：防重复元素）"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # 固定列宽布局（先清理旧元素防重复）
    old_layout = tblPr.find(qn('w:tblLayout'))
    if old_layout is not None:
        tblPr.remove(old_layout)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    # 边框（先清理旧元素防重复）
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def enforce_fixed_table_layout(table):
    """WPS 兼容：显式写入 tblLayout type='fixed'（防重复）"""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblLayout'))
    if old is not None:
        tblPr.remove(old)
    tl = OxmlElement('w:tblLayout')
    tl.set(qn('w:type'), 'fixed')
    tblPr.append(tl)


def set_fixed_col_widths(table, widths_cm):
    """锁定表格列宽（WPS 深度兼容版）— tblGrid + 逐单元格 <w:tcW> 双重锁定"""
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl

    # 1. tblGrid + gridCol
    old_grid = tbl.find(qn('w:tblGrid'))
    if old_grid is not None:
        tbl.remove(old_grid)
    tblGrid = OxmlElement('w:tblGrid')
    for w_cm in widths_cm:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w_cm * 567)))
        tblGrid.append(gridCol)
    tbl.insert(0, tblGrid)

    # 2. 逐单元格设置宽度
    for i, w_cm in enumerate(widths_cm):
        if i < len(table.columns):
            for cell in table.columns[i].cells:
                cell.width = Cm(w_cm)

    # 3. WPS 关键：为每一行的每一个 <w:tc> 注入 <w:tcW>（WPS 严格依赖此标签）
    twips = [int(w * 567) for w in widths_cm]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(twips):
                break
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            # 移除旧 tcW
            old_tcw = tcPr.find(qn('w:tcW'))
            if old_tcw is not None:
                tcPr.remove(old_tcw)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(twips[i]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)


def set_table_row_height(row, height_cm):
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


# ============================================================
# 数据解析工具（原封保留）
# ============================================================

def parse_overdue_amounts_days(remark_text):
    result = {'actual_amount': 0.0, 'actual_days': 0.0,
              'past_amount': 0.0, 'past_days': 0.0,
              'is_actual': False, 'is_past': False}
    if not remark_text:
        return result
    text = str(remark_text)
    actual_amt_match = re.search(r'实际逾期[^\d]{0,15}(\d[\d,]*\.?\d*)\s*元', text)
    if actual_amt_match:
        result['actual_amount'] = parse_number(actual_amt_match.group(1))
    actual_day_match = re.search(r'实际逾期[^\d]{0,10}(\d+\.?\d*)\s*天', text)
    if actual_day_match:
        result['actual_days'] = parse_number(actual_day_match.group(1))
    if '实际逾期' in text:
        result['is_actual'] = True
    past_amt_match = re.search(r'曾逾期[^\d]{0,15}(\d[\d,]*\.?\d*)\s*元', text)
    if past_amt_match:
        result['past_amount'] = parse_number(past_amt_match.group(1))
    past_day_match = re.search(r'曾逾期[^\d]{0,10}(\d+\.?\d*)\s*天', text)
    if past_day_match:
        result['past_days'] = parse_number(past_day_match.group(1))
    if '曾逾期' in text:
        result['is_past'] = True
    if not result['is_actual'] and not result['is_past']:
        remaining_match = re.search(r'剩余逾期[^\d]{0,10}(\d[\d,]*\.?\d*)\s*元', text)
        if remaining_match:
            result['actual_amount'] = parse_number(remaining_match.group(1))
            result['is_actual'] = True
    return result


def format_amount_wan(yuan_val):
    if yuan_val is None:
        return 0
    return round(float(yuan_val) / 10000.0)


def format_qty_auto(ton_val):
    tons = float(ton_val) if ton_val else 0
    if 1 <= tons <= 49:
        return f"{int(round(tons))}吨"
    else:
        wan = tons / 10000.0
        return f"{wan:.2f}万吨"


def extract_past_settlement(remark_text):
    if not remark_text:
        return ''
    text = str(remark_text)
    patterns = [
        (r'(\d{1,2})\s*月\s*(\d{1,2})\s*日\s*(已回清)', None),
        (r'(\d{1,2})\s*月\s*(\d{1,2})\s*日\s*(已回款)', None),
        (r'(?<!\d)(\d{1,2})\s*[-/]\s*(\d{1,2})\s*(已回清)', None),
        (r'(?<!\d)(\d{1,2})\s*[-/]\s*(\d{1,2})\s*(已回款)', None),
        (r'(已回清)\D{0,8}(\d{1,2})\s*[-/月]\s*(\d{1,2})', 'reversed'),
        (r'(已回款)\D{0,8}(\d{1,2})\s*[-/月]\s*(\d{1,2})', 'reversed'),
    ]
    for pat, mode in patterns:
        m = re.search(pat, text)
        if m:
            if mode == 'reversed':
                status = m.group(1)
                month = int(m.group(2))
                day = int(m.group(3))
            else:
                month = int(m.group(1))
                day = int(m.group(2))
                status = m.group(3)
            return f'{month}月{day}日{status}'
    return ''


def calc_pct_change(cur, prev):
    diff = cur - prev
    if prev == 0:
        return 100 if diff > 0 else 0
    return round(diff / prev * 100)


# ============================================================
# 常量
# ============================================================

REFERENCE_COLUMNS = [
    "客户名称", "授信类型", "授信模式", "批复赊销额度",
    "批复赊销账期", "期末赊销余额", "剩余可用授信额度",
    "逾期总额", "逾期积数"
]

FONT_RED_BOLD = Font(bold=True, color="FF0000")

# ============================================================
# 内部构件：生成 Word 报告
# ============================================================

def _generate_weekly_credit_report(daily_bytes_io, overdue_bytes_io,
                                    updated_summary_io, g_value_map, all_k_values,
                                    existing_l_values, ext_data, target_row_balance,
                                    overdue_contracts, actual_only_contracts,
                                    actual_count, actual_amount_wan):
    """在内存中生成完整的逾期赊销周报 Word 文档，返回 BytesIO。
    overdue_contracts / actual_count 等由主函数步骤1传入，避免重复解析Excel。"""
    today = date.today()

    # 截至日期统一为距今天最近的上一个周三（1、2、3 段落用）
    last_wed_local = get_last_wednesday(today)
    date1_str = f"{last_wed_local.month}月{last_wed_local.day}日"
    date2 = last_wed_local + timedelta(days=2)
    date2_str = f"{date2.month}月{date2.day}日"
    mmdd_str = today.strftime('%m%d')
    last_wed = last_wed_local  # 周报行号推算复用同一周三

    # 沿海大区摘要行的截止日期：距今天最近的上一个周五（若今天周五则取今天）
    days_since_fri = (today.weekday() - 4) % 7
    last_fri = today - timedelta(days=days_since_fri)
    date_fri_str = f"{last_fri.month}月{last_fri.day}日"

    # --- 读取 周逾期赊销对比 数据（从更新后的汇总文件）---
    updated_summary_io.seek(0)
    wb_summary_r = openpyxl.load_workbook(io.BytesIO(updated_summary_io.read()), data_only=True)
    ws_overdue_b = wb_summary_r['周逾期赊销对比']

    base_date_overdue = None
    i5_raw = ws_overdue_b.cell(row=5, column=9).value
    if isinstance(i5_raw, datetime):
        base_date_overdue = i5_raw.date()
    elif isinstance(i5_raw, date):
        base_date_overdue = i5_raw
    else:
        try:
            base_date_overdue = (datetime(1899, 12, 30) + timedelta(days=int(float(str(i5_raw))))).date()
        except Exception:
            pass
    if base_date_overdue is None:
        wb_summary_r.close()
        raise ValueError("无法读取'周逾期赊销对比'I5 单元格的基准日期")

    weeks_diff = (last_wed - base_date_overdue).days // 7
    target_row_b = 5 + weeks_diff
    target_row_b = min(target_row_b, ws_overdue_b.max_row)

    def _read_from_memory_or_excel(row, col, memory_dict, col_hint=""):
        if row in memory_dict and memory_dict[row] is not None:
            return float(memory_dict[row])
        excel_val = safe_float(ws_overdue_b.cell(row=row, column=col).value)
        if excel_val is not None:
            return excel_val
        return 0.0

    def _scan_upward_for_prev(start_row, col, memory_dict, label):
        for r in range(start_row, 4, -1):
            if r in memory_dict and memory_dict[r] is not None:
                val = safe_float(memory_dict[r])
                if val is not None and val > 0:
                    return val
            excel_val = safe_float(ws_overdue_b.cell(row=r, column=col).value)
            if excel_val is not None and excel_val > 0:
                return excel_val
        return 0.0

    k_latest = _read_from_memory_or_excel(target_row_b, 11, all_k_values)
    k_prev = _scan_upward_for_prev(target_row_b - 1, 11, all_k_values, "K上期")
    g_latest = _read_from_memory_or_excel(target_row_b, 7, g_value_map)
    l_latest = _read_from_memory_or_excel(target_row_b, 12, existing_l_values)
    l_prev = _scan_upward_for_prev(target_row_b - 1, 12, existing_l_values, "L上期")

    n_latest = safe_float(ws_overdue_b.cell(row=target_row_b, column=14).value)
    if n_latest is None:
        n_latest = 0.0

    q4_val = safe_float(ws_overdue_b.cell(row=4, column=17).value) or 0.0
    q5_val = safe_float(ws_overdue_b.cell(row=5, column=17).value) or 0.0

    g_values = []
    for r in range(5, target_row_b + 1):
        gv = safe_float(ws_overdue_b.cell(row=r, column=7).value)
        if gv is not None:
            g_values.append(gv)
    total_rows_g = target_row_b - 5 + 1
    g_avg = sum(g_values) / total_rows_g if total_rows_g > 0 else 0.0

    wb_summary_r.close()

    # --- 计算环比/同比 ---
    k_change = k_latest - k_prev
    k_pct_change = calc_pct_change(k_latest, k_prev)
    g_change = k_latest - g_latest
    g_pct_change = calc_pct_change(k_latest, g_latest)
    l_change_ppt = (l_latest - l_prev) * 100
    q_change_ppt = (q4_val - q5_val) * 100
    n_rounded = int(round(n_latest))
    g_avg_rounded = int(round(g_avg))
    n_change = n_latest - g_avg
    n_pct_change = calc_pct_change(n_latest, g_avg)

    # overdue_contracts / actual_count / actual_amount_wan 由主函数步骤1传入，零重复解析

    # ==================== 构建 Word 文档 ====================
    doc = Document()
    set_page_margins(doc)

    style_nm = doc.styles.add_style('NormalContent', 1)
    set_font_mixed(style_nm, 14.0, False, '仿宋_GB2312', 'Times New Roman')
    style_nm.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_nm.paragraph_format.line_spacing = Pt(28)
    style_nm.paragraph_format.space_before = Pt(0)
    style_nm.paragraph_format.space_after = Pt(0)
    style_nm.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_nm.paragraph_format.first_line_indent = Pt(28)

    style_tb = doc.styles.add_style('TableTitle', 1)
    set_font_mixed(style_tb, 12.0, False, '微软雅黑', 'Times New Roman')
    style_tb.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_tb.paragraph_format.line_spacing = Pt(22)
    style_tb.paragraph_format.space_before = Pt(0)
    style_tb.paragraph_format.space_after = Pt(0)
    style_tb.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_tb.paragraph_format.first_line_indent = Pt(0)

    # -- 大标题 --
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(0)
    title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title_para.paragraph_format.line_spacing = Pt(28)
    title_para.paragraph_format.first_line_indent = 0
    run_title = title_para.add_run('（一）逾期赊销')
    set_font_mixed(run_title, 14.0, True, '黑体', 'Times New Roman')

    # -- 1、整体情况 --
    p1 = doc.add_paragraph(style='NormalContent')
    run_p1_1 = p1.add_run(f'1、整体情况：截至{date1_str}，逾期赊销余额 {k_latest:.2f} 万元')
    set_font_mixed(run_p1_1, 14.0, False)
    run_hb = p1.add_run('，环比')
    set_font_mixed(run_hb, 14.0, False)
    run_hb.font.color.rgb = RGBColor(255, 0, 0)
    run_hb_val = p1.add_run(f'上期的 {k_prev:.2f} 万元')
    set_font_mixed(run_hb_val, 14.0, False)
    if k_change >= 0:
        run_hb_pct = p1.add_run(f'增加 {k_pct_change:.0f}% 或 {abs(k_change):.2f} 万元')
    else:
        run_hb_pct = p1.add_run(f'减少 {abs(k_pct_change):.0f}% 或 {abs(k_change):.2f} 万元')
    set_font_mixed(run_hb_pct, 14.0, False)
    run_hb_pct.font.color.rgb = RGBColor(255, 0, 0)
    run_tb = p1.add_run('；同比')
    set_font_mixed(run_tb, 14.0, False)
    run_tb.font.color.rgb = RGBColor(255, 0, 0)
    run_tb_val = p1.add_run(f'上年同期的 {g_latest:.2f} 万元')
    set_font_mixed(run_tb_val, 14.0, False)
    if g_change >= 0:
        run_tb_pct = p1.add_run(f'增加 {g_pct_change:.0f}% 或 {abs(g_change):.2f} 万元')
    else:
        run_tb_pct = p1.add_run(f'减少 {abs(g_pct_change):.0f}% 或 {abs(g_change):.2f} 万元')
    set_font_mixed(run_tb_pct, 14.0, False)
    run_tb_pct.font.color.rgb = RGBColor(255, 0, 0)
    p1.add_run('。')

    # -- 2、逾期赊销率 --
    p2 = doc.add_paragraph(style='NormalContent')
    run_p2_1 = p2.add_run(f'2、逾期赊销率：截止{date1_str}，逾期率 {l_latest:.2%}')
    set_font_mixed(run_p2_1, 14.0, False)
    p2.add_run(f'，上期逾期率 {l_prev:.2%}')
    if l_change_ppt >= 0:
        p2.add_run(f' 增加 {abs(l_change_ppt):.1f} 个PPT')
    else:
        p2.add_run(f' 减少 {abs(l_change_ppt):.1f} 个PPT')
    run_tb2 = p2.add_run(f'；本年累计逾期率 {q4_val:.2%}，同比')
    set_font_mixed(run_tb2, 14.0, False)
    run_tb2.font.color.rgb = RGBColor(255, 0, 0)
    run_tb2_val = p2.add_run(f'上年同期的 {q5_val:.2%}')
    set_font_mixed(run_tb2_val, 14.0, False)
    if q_change_ppt >= 0:
        run_q_pct = p2.add_run(f' 增加 {abs(q_change_ppt):.1f} 个PPT')
    else:
        run_q_pct = p2.add_run(f' 减少 {abs(q_change_ppt):.1f} 个PPT')
    set_font_mixed(run_q_pct, 14.0, False)
    run_q_pct.font.color.rgb = RGBColor(255, 0, 0)
    p2.add_run('。')

    # -- 3、周均逾期额 --
    p3 = doc.add_paragraph(style='NormalContent')
    run_p3_1 = p3.add_run(f'3、周均逾期额：2026年周均逾期额 {n_rounded} 万元')
    set_font_mixed(run_p3_1, 14.0, False)
    run_tb3 = p3.add_run(f'，同比')
    set_font_mixed(run_tb3, 14.0, False)
    run_tb3.font.color.rgb = RGBColor(255, 0, 0)
    run_tb3_val = p3.add_run(f'上年同期的 {g_avg_rounded} 万元')
    set_font_mixed(run_tb3_val, 14.0, False)
    if n_change >= 0:
        run_n_pct = p3.add_run(f' 增加 {n_pct_change:.0f}% 或 {abs(n_change):.0f} 万元')
    else:
        run_n_pct = p3.add_run(f' 减少 {abs(n_pct_change):.0f}% 或 {abs(n_change):.0f} 万元')
    set_font_mixed(run_n_pct, 14.0, False)
    run_n_pct.font.color.rgb = RGBColor(255, 0, 0)
    p3.add_run('。')
    p3.paragraph_format.space_after = Pt(6)

    # -- 沿海大区逾期赊销款剩余情况 --
    p4 = doc.add_paragraph(style='NormalContent')
    run_p4_1 = p4.add_run(f'截至{date_fri_str}中午，沿海大区逾期赊销款剩余 {actual_count} 笔未回清，共 {actual_amount_wan} 万元，具体情况如下：')
    set_font_mixed(run_p4_1, 14.0, True)

    # ==================== Table 1: 逾期赊销明细表 ====================
    doc.add_paragraph('逾期赊销明细表（万元）', style='TableTitle')
    COL_WIDTHS_T1 = [0.75, 1.38, 3.08, 1.05, 1.05, 1.28, 1.28, 1.28, 1.28, 1.28, 4.8]
    total_rows_t1 = 3 + len(overdue_contracts)
    table1 = doc.add_table(rows=total_rows_t1, cols=11)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_col_widths(table1, COL_WIDTHS_T1)

    row0_merge_cols = [0, 1, 2, 3, 4, 5]
    for c in row0_merge_cols:
        table1.cell(0, c).merge(table1.cell(1, c))
    table1.cell(0, 6).merge(table1.cell(0, 7))
    table1.cell(0, 8).merge(table1.cell(0, 9))
    table1.cell(0, 10).merge(table1.cell(1, 10))

    headers_row0 = ['序号', '业务单位', '客户名称', '品种', '交货方式', '合同数量',
                    f'截至{date1_str}', '', f'截至{date2_str}', '', '备注']
    for c, h in enumerate(headers_row0):
        if h:
            build_cell_text(table1.cell(0, c), h, bold=True, cn_size=9.0)
    headers_row1 = ['', '', '', '', '', '', '逾期天数', '逾期金额', '逾期天数', '逾期金额', '']
    for c, h in enumerate(headers_row1):
        if h:
            cell = table1.cell(1, c)
            build_cell_text(cell, h, bold=True, cn_size=8.5)
            set_cell_background(cell, 'D9D9D9')

    sum_col7_amt = 0
    sum_col9_amt = 0
    for contract in overdue_contracts:
        sum_col7_amt += format_amount_wan(contract['total_amount'])
        sum_col9_amt += format_amount_wan(contract['actual_amount'])

    table1.cell(2, 0).merge(table1.cell(2, 2))
    build_cell_text(table1.cell(2, 0), '沿海大区合计', bold=True, cn_size=9.0)
    for c in range(1, 11):
        if c == 7:
            build_cell_text(table1.cell(2, c), str(sum_col7_amt), bold=True, cn_size=9.0)
        elif c == 9:
            build_cell_text(table1.cell(2, c), str(sum_col9_amt), bold=True, cn_size=9.0)
        else:
            build_cell_text(table1.cell(2, c), '', bold=True, cn_size=9.0)

    for i, contract in enumerate(overdue_contracts):
        row_idx = 3 + i
        build_cell_text(table1.cell(row_idx, 0), str(i + 1), cn_size=9.0)
        build_cell_text(table1.cell(row_idx, 1), contract['dept'], align='left', cn_size=9.0)
        build_cell_text(table1.cell(row_idx, 2), contract['customer'], align='left', cn_size=9.0)
        build_cell_text(table1.cell(row_idx, 3), contract['variety'], cn_size=9.0)
        build_cell_text(table1.cell(row_idx, 4), contract['delivery'], cn_size=9.0)
        
        # 修改1：仅在表格填充时，通过字符串替换去除“万吨”和“吨”后缀，不影响其他地方
        qty_display = format_qty_auto(contract['qty_tons']).replace('万吨', '').replace('吨', '')
        build_cell_text(table1.cell(row_idx, 5), qty_display, cn_size=9.0)
        
        days1 = int(round(contract['total_days'])) if contract['total_days'] > 0 else '-'
        build_cell_text(table1.cell(row_idx, 6), str(days1) if days1 != '-' else days1, cn_size=9.0)
        amt1 = format_amount_wan(contract['total_amount'])
        build_cell_text(table1.cell(row_idx, 7), str(amt1) if amt1 > 0 else '0', cn_size=9.0)
        
        if contract['is_actual']:
            # 修改2：将条件成立时的逾期天数 + 2
            days2 = int(round(contract['actual_days'])) + 2 if contract['actual_days'] > 0 else '-'
            build_cell_text(table1.cell(row_idx, 8), str(days2) if days2 != '-' else days2, cn_size=9.0)
        else:
            build_cell_text(table1.cell(row_idx, 9), '-', cn_size=9.0)
        if contract['is_past']:
            settlement = extract_past_settlement(contract['remark'])
            remark_display = settlement if settlement else contract['remark']
        else:
            remark_display = contract['remark']
        build_cell_text(table1.cell(row_idx, 10), remark_display, align='left', cn_size=7.5)

    for i, row in enumerate(table1.rows):
        set_table_row_height(row, 0.66)
        for j, cell in enumerate(row.cells):
            if i < 2:
                set_cell_background(cell, 'D9D9D9')
    set_repeat_table_header(table1.rows[0])
    apply_table_borders(table1)
    # WPS 二次重锁：确保合并单元格后的新行携带宽度标签
    enforce_fixed_table_layout(table1)
    set_fixed_col_widths(table1, COL_WIDTHS_T1)

    # ==================== 4、逾期分客户 ====================
    p5 = doc.add_paragraph(style='NormalContent')
    run_p5_1 = p5.add_run('4、逾期分客户。')
    set_font_mixed(run_p5_1, 14.0, False)

    credit_stats_yuan = {'信保业务': 0.0, '集团内部企业': 0.0, '政策性业务': 0.0}
    for contract in overdue_contracts:
        ct = str(contract.get('credit_type', '')).strip()
        amt_yuan = float(contract.get('total_amount', 0.0))
        if '信用' in ct:
            credit_stats_yuan['信保业务'] += amt_yuan
        elif '中粮' in ct:
            credit_stats_yuan['集团内部企业'] += amt_yuan
        elif '政策' in ct:
            credit_stats_yuan['政策性业务'] += amt_yuan

    internal_amt = format_amount_wan(credit_stats_yuan['集团内部企业'])
    policy_amt = format_amount_wan(credit_stats_yuan['政策性业务'])
    insurance_amt = format_amount_wan(credit_stats_yuan['信保业务'])
    p5.add_run(f'集团内部企业逾期 {internal_amt} 万元；政策性业务逾期 {policy_amt} 万元；信保业务逾期 {insurance_amt} 万元。')

    # ==================== Table 2: 2026年十大逾期客户 ====================
    REF_COLS_T2 = ["周数", "客户所属集团", "客户名称", "销售合同编号",
                   "有关情况说明", "实际逾期金额", "回款时间"]
    overdue_bytes_io.seek(0)
    wb_overdue_t2 = openpyxl.load_workbook(io.BytesIO(overdue_bytes_io.read()), data_only=True)
    ws_2026 = wb_overdue_t2['2026年汇总']

    header_row_t2 = None
    col_map_t2 = {}
    for r in range(1, min(ws_2026.max_row, 20) + 1):
        score = 0
        temp_map = {}
        for c in range(1, ws_2026.max_column + 1):
            v = ws_2026.cell(row=r, column=c).value
            if v is not None and isinstance(v, str):
                stripped = v.strip()
                if stripped in REF_COLS_T2:
                    score += 1
                    temp_map[stripped] = c
        if score >= 5:
            header_row_t2 = r
            col_map_t2 = temp_map
            break

    col_group_idx = col_map_t2.get('客户所属集团', None)
    col_cust_idx = col_map_t2.get('客户名称', None)
    col_amount_idx = col_map_t2.get('实际逾期金额', None)
    col_days_indices = []
    for c in range(1, ws_2026.max_column + 1):
        v = ws_2026.cell(row=header_row_t2, column=c).value if header_row_t2 else None
        if v is not None and isinstance(v, str) and '分批实际逾期天数' in v.strip():
            col_days_indices.append(c)

    group_data = {}
    if header_row_t2 is not None and col_amount_idx is not None:
        for r in range(header_row_t2 + 1, ws_2026.max_row + 1):
            group_val = ws_2026.cell(row=r, column=col_group_idx).value if col_group_idx else None
            cust_val = ws_2026.cell(row=r, column=col_cust_idx).value if col_cust_idx else None
            key_raw = group_val if (group_val is not None and str(group_val).strip() != '') else cust_val
            if key_raw is None or str(key_raw).strip() == '':
                continue
            key = str(key_raw).strip()
            if '储备管理' in key:
                continue
            amt_raw = ws_2026.cell(row=r, column=col_amount_idx).value
            amt = safe_float(amt_raw)
            if amt is None or amt <= 0:
                continue
            max_days = 0.0
            for dc in col_days_indices:
                day_raw = ws_2026.cell(row=r, column=dc).value
                if day_raw is None:
                    continue
                day_str = str(day_raw).strip()
                if day_str in ('', '-', '—', '无', 'N/A', 'NA'):
                    continue
                dv = safe_float(day_raw)
                if dv is not None and dv > max_days:
                    max_days = dv
            if key not in group_data:
                group_data[key] = {'total_amount': 0.0, 'row_count': 0, 'max_days': 0.0}
            group_data[key]['total_amount'] += amt
            group_data[key]['row_count'] += 1
            if max_days > group_data[key]['max_days']:
                group_data[key]['max_days'] = max_days

    group_list = []
    for k, v in group_data.items():
        group_list.append({'name': k, 'total_amount': v['total_amount'],
                           'row_count': v['row_count'], 'max_days': v['max_days']})
    group_list.sort(key=lambda x: x['total_amount'], reverse=True)
    top10 = group_list[:10]
    wb_overdue_t2.close()

    doc.add_paragraph('2026年十大逾期客户（万元）', style='TableTitle')
    COL_WIDTHS_T2 = [1.08, 2.6, 3.15, 9.5]
    data_rows_t2 = max(len(top10), 1)
    table2 = doc.add_table(rows=1 + data_rows_t2, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_col_widths(table2, COL_WIDTHS_T2)
    headers_t2 = ['序号', '客户名称', '本年累计逾期发生额', f'客户逾期情况说明（截止{date1_str}）']
    for c, h in enumerate(headers_t2):
        build_cell_text(table2.cell(0, c), h, bold=True, cn_size=9.0)
        set_cell_background(table2.cell(0, c), 'D9D9D9')
    for i, item in enumerate(top10):
        row_idx = i + 1
        build_cell_text(table2.cell(row_idx, 0), str(i + 1), cn_size=9.0)
        build_cell_text(table2.cell(row_idx, 1), item['name'], align='left', cn_size=9.0)
        amt_wan = round(item['total_amount'] / 10000.0)
        build_cell_text(table2.cell(row_idx, 2), str(amt_wan), cn_size=9.0)
        x_val = int(item['row_count'])
        y_val = int(round(item['max_days']))
        desc_text = f"本年共{x_val}笔合同出现逾期、最长逾期天数{y_val}天"
        build_cell_text(table2.cell(row_idx, 3), desc_text, align='left', cn_size=9.0)
    if len(top10) == 0:
        for c in range(4):
            build_cell_text(table2.cell(1, c), '', cn_size=9.0)
    for i, row in enumerate(table2.rows):
        set_table_row_height(row, 0.66)
    set_repeat_table_header(table2.rows[0])
    apply_table_borders(table2)
    enforce_fixed_table_layout(table2)
    set_fixed_col_widths(table2, COL_WIDTHS_T2)

    # ================================================================
    # （一）赊销业务 — 新增章节
    # ================================================================
    doc.add_page_break()
    title2_para = doc.add_paragraph()
    title2_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title2_para.paragraph_format.space_before = Pt(0)
    title2_para.paragraph_format.space_after = Pt(0)
    title2_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title2_para.paragraph_format.line_spacing = Pt(28)
    title2_para.paragraph_format.first_line_indent = 0
    run_title2 = title2_para.add_run('（一）赊销业务')
    set_font_mixed(run_title2, 14.0, True, '黑体', 'Times New Roman')

    # val1 = 周赊销余额对比 最新写入行 G列（优先内存字典 g_value_map）
    val1 = float(g_value_map.get(target_row_balance, 0.0))
    val1_prev = 0.0
    for r in range(target_row_balance - 1, 2, -1):
        pv = g_value_map.get(r)
        if pv is not None and pv > 0:
            val1_prev = pv
            break
    diff1 = val1 - val1_prev
    pct1 = calc_pct_change(val1, val1_prev)

    val2 = float(ext_data.get('M4', 0.0))
    val3 = float(ext_data.get('M5', 0.0))
    val4 = float(ext_data.get('M6', 0.0))
    rate = float(ext_data.get('N4', 0.0))
    target_d = int(ext_data.get('target_row_d', 0))

    # prev_val2 从更新后的赊销外部余额读取（updated_summary_io 含 Step2C 写入）
    updated_summary_io.seek(0)
    wb_ext2 = openpyxl.load_workbook(io.BytesIO(updated_summary_io.read()), data_only=True)
    ws_ext_d = wb_ext2['赊销外部余额']
    prev_val2 = 0.0
    if target_d > 3:
        for r in range(target_d - 1, 2, -1):
            pv2 = safe_float(ws_ext_d.cell(row=r, column=7).value)
            if pv2 is not None and pv2 > 0:
                prev_val2 = pv2
                break
    diff2 = val2 - prev_val2
    pct2 = calc_pct_change(val2, prev_val2)
    wb_ext2.close()

    v1 = int(round(val1))
    d1 = int(round(abs(diff1)))
    v2 = int(round(val2))
    pv2_i = int(round(prev_val2))
    d2 = int(round(abs(diff2)))
    v3 = int(round(val3))
    v4 = int(round(val4))
    rt = int(round(rate * 100))

    if diff1 >= 0:
        hb1_str = f'增加 {pct1}% 或 {d1} 万元'
    else:
        hb1_str = f'减少 {abs(pct1)}% 或 {d1} 万元'
    if diff2 >= 0:
        hb2_str = f'增加 {pct2}% 或 {d2} 万元'
    else:
        hb2_str = f'减少 {abs(pct2)}% 或 {d2} 万元'

    full_text = (
        f'1、整体情况：截至{date1_str}，沿海大区赊销余额共计 {v1} 万元，'
        f'环比{hb1_str}，'
        f'其中：外部赊销合 {v2} 万元，环比上期 {pv2_i} 万元{hb2_str}，'
        f'外部授权赊销限额使用率 {rt}%。'
        f'政策性 {v3} 万元，集团内赊销余额 {v4} 万元。'
    )

    p_sale = doc.add_paragraph(style='NormalContent')
    run_sale = p_sale.add_run(full_text)
    set_font_mixed(run_sale, 14.0, False)

    # ==================== 经营部赊销结构表 ====================
    doc.add_paragraph()
    p_tbl_title = doc.add_paragraph('经营部赊销结构表', style='TableTitle')
    daily_bytes_io.seek(0)
    wb_dept = openpyxl.load_workbook(io.BytesIO(daily_bytes_io.read()), data_only=True)
    ws_dept = wb_dept['经营部']
    dept_data = []
    for r in range(2, 8):
        row_vals = []
        for c in range(1, 11):
            row_vals.append(ws_dept.cell(row=r, column=c).value)
        dept_data.append(row_vals)
    wb_dept.close()

    COL_WIDTHS_T3 = [1.2, 1.6, 2, 1.5, 2, 1.5, 1.5, 2, 1.5, 1.5]
    table3 = doc.add_table(rows=7, cols=10)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_col_widths(table3, COL_WIDTHS_T3)
    headers_t3 = ['序号', '经营部', '本周赊销余额', '占比', '上周赊销余额',
                  '占比', '环比', '2025年同期余额', '占比', '同比']
    for c, h in enumerate(headers_t3):
        build_cell_text(table3.cell(0, c), h, bold=True, cn_size=9.0)
        set_cell_background(table3.cell(0, c), 'D9D9D9')

    for i, row_vals in enumerate(dept_data):
        row_idx = i + 1
        is_last = (i == 5)
        for c, raw_val in enumerate(row_vals):
            display_val = raw_val
            if c == 0 and raw_val is not None:
                fv = safe_float(raw_val)
                display_val = str(int(fv)) if fv is not None else str(raw_val)
            elif raw_val is not None and c in (3, 5, 6, 8, 9):
                fv = safe_float(raw_val)
                if fv is not None:
                    display_val = f'{fv*100:.0f}%' if abs(fv) < 10 else f'{fv:.1f}%'
            elif raw_val is not None and isinstance(raw_val, (int, float)):
                display_val = f'{raw_val:,.2f}'
            elif raw_val is None:
                display_val = ''
            else:
                display_val = str(raw_val)
            align = 'left' if c == 1 else 'center'
            build_cell_text(table3.cell(row_idx, c), display_val,
                           align=align, bold=is_last, cn_size=9.0)

    for i, row in enumerate(table3.rows):
        set_table_row_height(row, 0.8)
    set_repeat_table_header(table3.rows[0])
    apply_table_borders(table3)
    enforce_fixed_table_layout(table3)
    set_fixed_col_widths(table3, COL_WIDTHS_T3)

    # ==================== 保存到 BytesIO ====================
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes


# ============================================================
# 主处理函数
# ============================================================

def process_credit_sale(uploaded_files):
    """
    逾期赊销周报处理入口。
    uploaded_files: list of BytesIO from st.file_uploader
    Returns: (excel_bytes, docx_bytes, logs)
    """
    logs = []

    # --- 文件自动识别 ---
    file_daily = None
    file_overdue = None
    file_summary = None

    for f in uploaded_files:
        fname = f.name if hasattr(f, 'name') else ''
        f.seek(0)
        data = f.read()
        if '赊销日报表' in fname:
            file_daily = (fname, data)
        elif '逾期明细汇总表' in fname:
            file_overdue = (fname, data)
        elif '赊销数据汇总' in fname:
            file_summary = (fname, data)

    missing = []
    if file_daily is None:
        missing.append('赊销日报表')
    if file_overdue is None:
        missing.append('逾期明细汇总表')
    if file_summary is None:
        missing.append('赊销数据汇总')
    if missing:
        logs.append(f"❌ 缺少必要文件: {', '.join(missing)}。请上传包含对应关键词的3个Excel文件。")
        return None, None, logs

    try:
        daily_bytes = file_daily[1]
        overdue_bytes = file_overdue[1]
        summary_bytes_raw = file_summary[1]

        # ============================================================
        # 步骤 1：处理赊销明细
        # ============================================================
        wb_daily = openpyxl.load_workbook(io.BytesIO(daily_bytes), read_only=True, data_only=True)
        ws_detail = wb_daily['赊销明细']

        header_row_idx = None
        header_col_map = {}
        header_search_row = 1
        for row in ws_detail.iter_rows(min_row=1, max_row=min(ws_detail.max_row, 20), values_only=False):
            score = 0
            temp_map = {}
            for cell in row:
                if cell.value is not None and isinstance(cell.value, str):
                    val_stripped = cell.value.strip()
                    if val_stripped in REFERENCE_COLUMNS:
                        score += 1
                        temp_map[val_stripped] = cell.column - 1
            if score >= 5:
                header_row_idx = header_search_row
                header_col_map = temp_map
                break
            header_search_row += 1

        if header_row_idx is None:
            wb_daily.close()
            logs.append("❌ 未能在'赊销明细'表单中找到标题行。")
            return None, None, logs

        additional_cols = {}
        for row in ws_detail.iter_rows(min_row=header_row_idx, max_row=header_row_idx, values_only=False):
            for cell in row:
                if cell.value is not None and isinstance(cell.value, str):
                    v = cell.value.strip()
                    if "销售合同编号" in v or "销售合同" in v:
                        additional_cols["销售合同编号"] = cell.column - 1
                    elif v == "客户名称":
                        additional_cols["客户名称"] = cell.column - 1
                    elif "逾期分类" in v:
                        additional_cols["逾期分类"] = cell.column - 1
                    elif "有关情况说明" in v or "情况说明" in v:
                        additional_cols["有关情况说明"] = cell.column - 1
                    elif v == "期末赊销余额":
                        additional_cols["期末赊销余额"] = cell.column - 1

        all_col_map = {}
        all_col_map.update(additional_cols)
        all_col_map.update(header_col_map)

        col_overdue_class = all_col_map.get("逾期分类")
        col_remark = all_col_map.get("有关情况说明")
        col_contract = all_col_map.get("销售合同编号")
        col_customer_idx = all_col_map.get("客户名称")
        col_balance = all_col_map.get("期末赊销余额")

        if col_overdue_class is None:
            wb_daily.close()
            logs.append("❌ 未找到'逾期分类'列。")
            return None, None, logs

        AMOUNT_RE_PRIMARY = re.compile(
            r'(?:实际已逾期|实际逾期|曾逾期)[^\d]{0,15}(\d[\d,]*\.?\d*)\s*元')
        AMOUNT_RE_FALLBACK = re.compile(r'剩余逾期[^\d]{0,10}(\d[\d,]*\.?\d*)\s*元')
        DAYS_RE = re.compile(r'(?:实际已逾期|实际逾期|曾逾期)[^\d]{0,10}(\d+\.?\d*)\s*天')

        total_overdue_amount = 0.0
        total_balance_sum = 0.0
        balance_count = 0

        for row in ws_detail.iter_rows(min_row=header_row_idx + 1, max_row=ws_detail.max_row, values_only=False):
            h_val = row[col_customer_idx].value if col_customer_idx is not None else None
            j_val = row[col_contract].value if col_contract is not None else None
            v_val = row[col_balance].value if col_balance is not None else None

            if (h_val is not None and str(h_val).strip() != "" and
                j_val is not None and str(j_val).strip() != ""):
                b = safe_float(v_val)
                if b is not None:
                    total_balance_sum += b
                    balance_count += 1

            al_val = row[col_overdue_class].value if col_overdue_class is not None else None
            if al_val is None:
                continue
            al_str = str(al_val).strip()
            if "A" not in al_str or "实际已逾期" not in al_str:
                continue

            ak_val = row[col_remark].value if col_remark is not None else None
            ak_str = str(ak_val) if ak_val is not None else ""
            amounts = AMOUNT_RE_PRIMARY.findall(ak_str)
            if not amounts:
                amounts = AMOUNT_RE_FALLBACK.findall(ak_str)
            row_amount = sum(parse_number(a) for a in amounts)
            days = DAYS_RE.findall(ak_str)
            if not days:
                days = re.findall(r'逾期[^\d]{0,5}(\d+\.?\d*)\s*天', ak_str)
            row_days = sum(parse_number(d) for d in days)
            total_overdue_amount += row_amount

        wb_daily.close()

        # --- 提取逾期合同明细（复用步骤1已定位的标题行与列索引，避免 Word 函数重复解析）---
        wb_daily_detail = openpyxl.load_workbook(io.BytesIO(daily_bytes), data_only=True)
        ws_detail2 = wb_daily_detail['赊销明细']

        # 复用步骤1的 all_col_map 定位辅助列（与步骤1完全一致的标题行）
        _col_customer_d = all_col_map.get('客户名称', None)
        _col_contract_d = all_col_map.get('销售合同编号', None)
        _col_balance_d  = all_col_map.get('期末赊销余额', None)
        _col_overdue_d  = all_col_map.get('逾期分类', None)

        # 扫描辅助列（品种/合同数量/经营部/交货方式/授信模式/有关情况说明）
        _col_variety_d   = None
        _col_quantity_d  = None
        _col_dept_d      = None
        _col_delivery_d  = None
        _col_credit_d    = None
        _col_remark_d    = None

        SCAN_MAX_D = max(ws_detail2.max_column + 30, 80)
        for _r in range(header_row_idx, min(header_row_idx + 2, ws_detail2.max_row + 1)):
            for _c in range(1, SCAN_MAX_D + 1):
                _v = ws_detail2.cell(row=_r, column=_c).value
                if _v is None:
                    continue
                _vs = str(_v).strip()
                if ('品种' == _vs or _vs == '品种') and _col_variety_d is None:
                    _col_variety_d = _c - 1
                elif '合同数量' in _vs and _col_quantity_d is None:
                    _col_quantity_d = _c - 1
                elif '经营部' in _vs and _col_dept_d is None:
                    _col_dept_d = _c - 1
                elif '交货方式' in _vs and _col_delivery_d is None:
                    _col_delivery_d = _c - 1
                elif '授信模式' in _vs and _col_credit_d is None:
                    _col_credit_d = _c - 1
                elif ('有关情况说明' in _vs or '情况说明' in _vs) and _col_remark_d is None:
                    _col_remark_d = _c - 1

        overdue_contracts_all = []
        actual_only_contracts_all = []

        for _r in range(header_row_idx + 1, ws_detail2.max_row + 1):
            al_v = ws_detail2.cell(row=_r, column=_col_overdue_d + 1).value if _col_overdue_d is not None else None
            if al_v is None:
                continue
            al_s = str(al_v).strip()
            if "A" not in al_s or "实际已逾期" not in al_s:
                continue

            remark_r = ws_detail2.cell(row=_r, column=_col_remark_d + 1).value if _col_remark_d is not None else ""
            remark_t = str(remark_r) if remark_r else ""
            parsed = parse_overdue_amounts_days(remark_t)

            variety_r = ws_detail2.cell(row=_r, column=_col_variety_d + 1).value if _col_variety_d is not None else ""
            variety = str(variety_r).strip() if variety_r else ""
            qty_r = ws_detail2.cell(row=_r, column=_col_quantity_d + 1).value if _col_quantity_d is not None else 0
            qty_tons = safe_float(qty_r) or 0
            dept_r = ws_detail2.cell(row=_r, column=_col_dept_d + 1).value if _col_dept_d is not None else ""
            dept = str(dept_r).strip() if dept_r else ""
            cust_r = ws_detail2.cell(row=_r, column=_col_customer_d + 1).value if _col_customer_d is not None else ""
            cust = str(cust_r).strip() if cust_r else ""
            delivery_r = ws_detail2.cell(row=_r, column=_col_delivery_d + 1).value if _col_delivery_d is not None else ""
            delivery = str(delivery_r).strip() if delivery_r else ""
            credit_r = ws_detail2.cell(row=_r, column=_col_credit_d + 1).value if _col_credit_d is not None else ""
            credit_type = str(credit_r).strip() if credit_r is not None else ""
            contract_r = ws_detail2.cell(row=_r, column=_col_contract_d + 1).value if _col_contract_d is not None else ""
            contract_no = str(contract_r).strip() if contract_r else ""

            total_amt = parsed['actual_amount'] + parsed['past_amount']
            total_day = parsed['actual_days'] + parsed['past_days']

            info = {
                'dept': dept, 'customer': cust, 'variety': variety,
                'delivery': delivery, 'qty_tons': qty_tons,
                'total_amount': total_amt, 'total_days': total_day,
                'actual_amount': parsed['actual_amount'],
                'actual_days': parsed['actual_days'],
                'past_amount': parsed['past_amount'],
                'past_days': parsed['past_days'],
                'is_actual': parsed['is_actual'], 'is_past': parsed['is_past'],
                'remark': remark_t, 'credit_type': credit_type,
                'contract_no': contract_no,
            }
            overdue_contracts_all.append(info)
            if parsed['is_actual'] and not parsed['is_past']:
                actual_only_contracts_all.append(info)

        _actual_count = len(actual_only_contracts_all)
        _actual_amount_wan = sum(format_amount_wan(c['actual_amount']) for c in overdue_contracts_all)
        wb_daily_detail.close()

        # ============================================================
        # 步骤 2：写入赊销数据汇总
        # ============================================================
        TODAY = date.today()
        wb_summary_read = openpyxl.load_workbook(io.BytesIO(summary_bytes_raw), data_only=True)

        ws_balance_read = find_sheet_by_keyword(wb_summary_read, ["周赊销余额对比", "赊销余额对比"])
        ws_overdue_read = find_sheet_by_keyword(wb_summary_read, ["周逾期赊销对比", "逾期赊销对比"])
        if ws_balance_read is None:
            ws_balance_read = wb_summary_read[wb_summary_read.sheetnames[0]]
        if ws_overdue_read is None:
            ws_overdue_read = wb_summary_read[wb_summary_read.sheetnames[1]]

        wb_summary = openpyxl.load_workbook(io.BytesIO(summary_bytes_raw))
        ws_balance = wb_summary[ws_balance_read.title]
        ws_overdue = wb_summary[ws_overdue_read.title]

        # 2A. 周赊销余额对比
        base_date_balance = None
        e3_val = ws_balance_read.cell(row=3, column=5).value
        if isinstance(e3_val, datetime):
            base_date_balance = e3_val.date()
        elif isinstance(e3_val, date):
            base_date_balance = e3_val
        else:
            e3_raw = ws_balance.cell(row=3, column=5).value
            if isinstance(e3_raw, datetime):
                base_date_balance = e3_raw.date()
            elif isinstance(e3_raw, date):
                base_date_balance = e3_raw

        if base_date_balance is None:
            wb_summary_read.close()
            wb_summary.close()
            logs.append("❌ 无法读取'周赊销余额对比'E3 基准日期。")
            return None, None, logs

        target_row_balance = None
        best_row = None
        best_date = None
        for r in range(3, ws_balance_read.max_row + 1):
            d = None
            e_val = ws_balance_read.cell(row=r, column=5).value
            if isinstance(e_val, datetime):
                d = e_val.date()
            elif isinstance(e_val, date):
                d = e_val
            else:
                d = base_date_balance + timedelta(days=(r - 3) * 7)
            if d == TODAY:
                target_row_balance = r
                break
            if d < TODAY:
                if best_date is None or d > best_date:
                    best_date = d
                    best_row = r
        if target_row_balance is None:
            target_row_balance = best_row
        if target_row_balance is None:
            weeks_diff = (TODAY - base_date_balance).days // 7
            target_row_balance = 3 + weeks_diff

        # G列空行填充
        nearest_g_row = None
        nearest_g_value = None
        for r in range(target_row_balance - 1, 3, -1):
            g_val = ws_balance.cell(row=r, column=7).value
            is_empty = (g_val is None or (isinstance(g_val, str) and g_val.strip() == ""))
            if not is_empty:
                nearest_g_row = r
                nearest_g_value = g_val
                break
        filled_g_rows = []
        if nearest_g_row is not None and nearest_g_value is not None:
            for r in range(nearest_g_row + 1, target_row_balance):
                cell_g = ws_balance.cell(row=r, column=7)
                g_val = cell_g.value
                is_empty = (g_val is None or (isinstance(g_val, str) and g_val.strip() == ""))
                if is_empty:
                    cell_g.value = nearest_g_value
                    cell_g.font = FONT_RED_BOLD
                    filled_g_rows.append(r)

        balance_wan = total_balance_sum / 10000.0
        cell_g_target = ws_balance.cell(row=target_row_balance, column=7)
        cell_g_target.value = balance_wan
        cell_g_target.font = FONT_RED_BOLD

        g_value_map = {}
        for r in range(3, target_row_balance + 1):
            if r == target_row_balance:
                g_value_map[r] = balance_wan
            elif r in filled_g_rows:
                g_value_map[r] = safe_float(ws_balance.cell(row=r, column=7).value)
            else:
                gv = safe_float(ws_balance_read.cell(row=r, column=7).value)
                if gv is not None:
                    g_value_map[r] = gv

        # 2B. 周逾期赊销对比
        base_date_overdue = None
        i5_val = ws_overdue_read.cell(row=5, column=9).value
        if isinstance(i5_val, datetime):
            base_date_overdue = i5_val.date()
        elif isinstance(i5_val, date):
            base_date_overdue = i5_val
        else:
            i5_raw = ws_overdue.cell(row=5, column=9).value
            if isinstance(i5_raw, datetime):
                base_date_overdue = i5_raw.date()
            elif isinstance(i5_raw, date):
                base_date_overdue = i5_raw

        if base_date_overdue is None:
            wb_summary_read.close()
            wb_summary.close()
            logs.append("❌ 无法读取'周逾期赊销对比'I5 基准日期。")
            return None, None, logs

        target_row_overdue = None
        best_row_o = None
        best_date_o = None
        for r in range(5, ws_overdue_read.max_row + 1):
            d = None
            i_val = ws_overdue_read.cell(row=r, column=9).value
            if isinstance(i_val, datetime):
                d = i_val.date()
            elif isinstance(i_val, date):
                d = i_val
            else:
                d = base_date_overdue + timedelta(days=(r - 5) * 7)
            if d == TODAY:
                target_row_overdue = r
                break
            if d < TODAY:
                if best_date_o is None or d > best_date_o:
                    best_date_o = d
                    best_row_o = r
        if target_row_overdue is None:
            target_row_overdue = best_row_o
        if target_row_overdue is None:
            weeks_diff_o = (TODAY - base_date_overdue).days // 7
            target_row_overdue = 5 + weeks_diff_o

        overdue_amount_wan = total_overdue_amount / 10000.0

        def is_k_cell_truly_empty(r):
            raw_k_read = ws_overdue_read.cell(row=r, column=11).value
            raw_k_real = ws_overdue.cell(row=r, column=11).value
            is_cached_empty = (raw_k_read is None or (isinstance(raw_k_read, str) and raw_k_read.strip() == ""))
            has_formula = isinstance(raw_k_real, str) and raw_k_real.strip().startswith("=")
            is_real_empty = (raw_k_real is None or (isinstance(raw_k_real, str) and raw_k_real.strip() == ""))
            return is_cached_empty and not has_formula and is_real_empty

        rows_to_process = []
        nearest_k_row = None
        nearest_k_value = None
        for r in range(target_row_overdue - 1, 4, -1):
            if is_k_cell_truly_empty(r):
                continue
            fv = safe_float(ws_overdue_read.cell(row=r, column=11).value)
            if fv is not None:
                nearest_k_row = r
                nearest_k_value = fv
                break
        if nearest_k_row is not None and nearest_k_value is not None:
            for r in range(nearest_k_row + 1, target_row_overdue):
                if is_k_cell_truly_empty(r):
                    cell_k = ws_overdue.cell(row=r, column=11)
                    cell_k.value = nearest_k_value
                    cell_k.font = FONT_RED_BOLD
                    rows_to_process.append((r, True))

        cell_k_target = ws_overdue.cell(row=target_row_overdue, column=11)
        cell_k_target.value = overdue_amount_wan
        cell_k_target.font = FONT_RED_BOLD
        rows_to_process.append((target_row_overdue, False))

        all_k_values = {}
        for r in range(5, target_row_overdue + 1):
            val = ws_overdue_read.cell(row=r, column=11).value
            fv = safe_float(val)
            if fv is not None:
                all_k_values[r] = fv
        for r, _ in rows_to_process:
            cell = ws_overdue.cell(row=r, column=11)
            if cell.value is not None:
                fv = safe_float(cell.value)
                if fv is not None:
                    all_k_values[r] = fv

        existing_l_values = {}
        for r in range(5, target_row_overdue + 1):
            l_val = ws_overdue_read.cell(row=r, column=12).value
            fv = safe_float(l_val)
            if fv is not None:
                existing_l_values[r] = fv
            else:
                raw_l = ws_overdue.cell(row=r, column=12).value
                if isinstance(raw_l, str) and raw_l.startswith("="):
                    pass
                elif raw_l is not None:
                    fv = safe_float(raw_l)
                    if fv is not None:
                        existing_l_values[r] = fv

        rows_to_process.sort(key=lambda x: x[0])
        for r, is_filled in rows_to_process:
            k_val = all_k_values.get(r)
            if k_val is None:
                continue
            g_row = r - 2
            g_val = g_value_map.get(g_row)
            if g_val is not None and g_val != 0:
                l_val = k_val / g_val
            else:
                l_val = 0.0
            cell_l = ws_overdue.cell(row=r, column=12)
            cell_l.value = l_val
            cell_l.number_format = '0.00%'
            cell_l.font = FONT_RED_BOLD
            prev_l = None
            for pr in range(r - 1, 4, -1):
                prev_l = existing_l_values.get(pr)
                if prev_l is not None:
                    break
                for processed_r, _ in rows_to_process:
                    if processed_r == pr:
                        pl_cell = ws_overdue.cell(row=pr, column=12)
                        if pl_cell.value is not None:
                            prev_l = safe_float(pl_cell.value)
                        break
                if prev_l is not None:
                    break
            m_val = l_val - prev_l if prev_l is not None else 0.0
            cell_m = ws_overdue.cell(row=r, column=13)
            cell_m.value = m_val
            cell_m.font = FONT_RED_BOLD
            k_list = []
            for kr in range(5, r + 1):
                kv = all_k_values.get(kr)
                if kv is not None:
                    k_list.append(kv)
            n_val = sum(k_list) / len(k_list) if k_list else 0.0
            n_val = round(n_val, 2)
            cell_n = ws_overdue.cell(row=r, column=14)
            cell_n.value = n_val
            cell_n.number_format = '0.00'
            cell_n.font = FONT_RED_BOLD
            existing_l_values[r] = l_val

        # Q4/Q5
        l_list_q4 = []
        for r in range(5, target_row_overdue + 1):
            lv = existing_l_values.get(r)
            if lv is None:
                lv = safe_float(ws_overdue_read.cell(row=r, column=12).value)
            if lv is not None:
                l_list_q4.append(lv)
        if l_list_q4:
            q4_val = sum(l_list_q4) / len(l_list_q4)
            ws_overdue.cell(row=4, column=17).value = q4_val
            ws_overdue.cell(row=4, column=17).font = FONT_RED_BOLD

        h_list_q5 = []
        for r in range(5, target_row_overdue + 1):
            hv = safe_float(ws_overdue_read.cell(row=r, column=8).value)
            if hv is not None:
                h_list_q5.append(hv)
        total_h_rows = target_row_overdue - 5 + 1
        if len(h_list_q5) >= total_h_rows * 0.5 and h_list_q5:
            q5_val = sum(h_list_q5) / len(h_list_q5)
            ws_overdue.cell(row=5, column=17).value = q5_val
            ws_overdue.cell(row=5, column=17).font = FONT_RED_BOLD

        # 2C. 赊销外部余额
        daily_bytes_io = io.BytesIO(daily_bytes)
        wb_daily_2c = openpyxl.load_workbook(daily_bytes_io, data_only=True)
        ws_quota = wb_daily_2c['额度']
        val_M4 = safe_float(ws_quota.cell(row=4, column=13).value) or 0.0
        val_M6 = safe_float(ws_quota.cell(row=6, column=13).value) or 0.0
        val_M5 = safe_float(ws_quota.cell(row=5, column=13).value) or 0.0
        val_N4 = safe_float(ws_quota.cell(row=4, column=14).value) or 0.0

        haida_sum_wan = 0.0
        if '海大' in wb_daily_2c.sheetnames:
            ws_haida = wb_daily_2c['海大']
            HAIDA_REF = ["客户名称", "合同数量", "期末赊销余额", "赊销天数", "授信类型"]
            haida_header = None
            haida_balance_col = None
            for r in range(1, min(ws_haida.max_row, 20) + 1):
                score = 0
                for c in range(1, ws_haida.max_column + 1):
                    v = ws_haida.cell(row=r, column=c).value
                    if v is not None and isinstance(v, str):
                        stripped = v.strip()
                        if stripped in HAIDA_REF:
                            score += 1
                        if '期末赊销余额' in stripped:
                            haida_balance_col = c
                if score >= 3:
                    haida_header = r
                    break
            haida_sum = 0.0
            if haida_header is not None and haida_balance_col is not None:
                for r in range(haida_header + 1, ws_haida.max_row + 1):
                    bv = safe_float(ws_haida.cell(row=r, column=haida_balance_col).value)
                    if bv is not None:
                        haida_sum += bv
            haida_sum_wan = haida_sum / 10000.0
        wb_daily_2c.close()

        ws_external = wb_summary['赊销外部余额']
        target_row_d = None
        for r in range(3, ws_external.max_row + 2):
            g_val = ws_external.cell(row=r, column=7).value
            if g_val is None or (isinstance(g_val, str) and g_val.strip() == ''):
                target_row_d = r
                break
        if target_row_d is None:
            target_row_d = ws_external.max_row + 1

        ws_external.cell(row=target_row_d, column=7).value = val_M4
        ws_external.cell(row=target_row_d, column=7).font = FONT_RED_BOLD
        ws_external.cell(row=target_row_d, column=8).value = val_M6
        ws_external.cell(row=target_row_d, column=8).font = FONT_RED_BOLD
        ws_external.cell(row=target_row_d, column=9).value = val_M5
        ws_external.cell(row=target_row_d, column=9).font = FONT_RED_BOLD
        ws_external.cell(row=target_row_d, column=10).value = haida_sum_wan
        ws_external.cell(row=target_row_d, column=10).font = FONT_RED_BOLD

        wb_summary_read.close()

        # 构建 ext_data
        ext_data = {
            'M4': val_M4, 'M6': val_M6, 'M5': val_M5, 'N4': val_N4,
            'target_row_d': target_row_d,
        }

        # ============================================================
        # 保存 Excel 到 BytesIO
        # ============================================================
        excel_bytes = io.BytesIO()
        wb_summary.save(excel_bytes)
        excel_bytes.seek(0)
        wb_summary.close()

        # ============================================================
        # 步骤 3：生成 Word 报告 — 传入更新后的 excel_bytes（含Step2写入）
        # ============================================================
        updated_summary_io = io.BytesIO(excel_bytes.getvalue())  # 更新后的汇总

        docx_bytes = _generate_weekly_credit_report(
            io.BytesIO(daily_bytes), io.BytesIO(overdue_bytes),
            updated_summary_io, g_value_map, all_k_values,
            existing_l_values, ext_data, target_row_balance,
            overdue_contracts_all, actual_only_contracts_all,
            _actual_count, _actual_amount_wan)

        logs.append("✅ 逾期赊销周报处理完成。")
        return excel_bytes, docx_bytes, logs

    except Exception as e:
        import traceback
        logs.append(f"❌ 处理失败: {str(e)}")
        logs.append(traceback.format_exc())
        return None, None, logs
