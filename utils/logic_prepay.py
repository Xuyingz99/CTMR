# ==========================================
# 预付款业务情况周报处理逻辑 (从 prepay.py 移植并适配 BytesIO)
# ==========================================
import io
import datetime
import re
from copy import copy
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, Color
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- 配置参数 ---
SHEET_NAME = "预付款业务情况周报表"
KEYWORDS = [
    "客商", "供应商", "平台名称", "收款方名称", "合同号", "预付款数量",
    "合同金额", "拟交易金额", "预付款性质", "品种", "本年累计预付金额",
    "期末预付金额", "合同约定货物交付时间", "实际货物交付时间",
    "业务所属区域", "经营部"
]


def get_last_wednesday():
    """获取最近一个周三日期对象（含今日）"""
    today = datetime.date.today()
    offset = (today.weekday() - 2) % 7
    return today - datetime.timedelta(days=offset)


def score_and_find_header(sheet):
    """智能定位标题行及标题行所占的行数"""
    header_row_start = -1
    header_row_end = -1

    for row_idx in range(1, 11):
        score = 0
        cell_values = [str(sheet.cell(row=row_idx, column=col).value or "").strip()
                       for col in range(1, sheet.max_column + 1)]

        for kw in KEYWORDS:
            if any(kw in val for val in cell_values):
                score += 1

        if score >= 4:
            header_row_start = row_idx
            header_row_end = row_idx

            if row_idx < 10:
                next_row_values = [str(sheet.cell(row=row_idx + 1, column=col).value or "").strip()
                                   for col in range(1, sheet.max_column + 1)]
                sub_keywords = ["1-10天", "11-30天", "31天以上", "逾期且财务已计提坏账金额"]
                if any(sub_kw in val for val in next_row_values for sub_kw in sub_keywords):
                    header_row_end = row_idx + 1
            break

    return header_row_start, header_row_end


def copy_cell_style(source_cell, target_cell):
    """克隆单元格的全部样式"""
    target_cell.font = copy(source_cell.font)
    target_cell.border = copy(source_cell.border)
    target_cell.fill = copy(source_cell.fill)
    target_cell.alignment = copy(source_cell.alignment)
    target_cell.number_format = source_cell.number_format
    if target_cell.value and any(kw in str(target_cell.value) for kw in ["合同约定货物交付时间", "实际货物交付时间", "预付时间"]):
        target_cell.number_format = 'yyyy/mm/dd'


def apply_new_sheet_style(ws, sum_col_idx):
    """为新生成的表单（保证金、货款）统一设置列宽、底纹格式与边框"""
    fill_blue = PatternFill(fill_type="solid", start_color="DEEBF6", end_color="DEEBF6")
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                         top=Side(style='thin'), bottom=Side(style='thin'))

    # 标题行底纹、字体与边框
    ws.row_dimensions[1].height = 22.5
    for col_idx in range(1, ws.max_column + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = fill_blue
        cell.font = Font(bold=True, name="宋体")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border

    # 合计行底纹与边框
    last_row = ws.max_row
    if last_row > 1:
        ws.row_dimensions[last_row].height = 22.5
        for col_idx in range(1, ws.max_column + 1):
            cell = ws.cell(row=last_row, column=col_idx)
            cell.fill = fill_blue
            cell.border = thin_border
            cell.font = Font(bold=True, name="宋体")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            if col_idx == sum_col_idx:
                cell.number_format = '#,##0.00'

    # 动态适应列宽和日期格式
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if isinstance(cell.value, datetime.datetime):
                cell.number_format = 'yyyy/mm/dd'

            if cell.row > 1 and cell.row < last_row:
                cell.border = thin_border
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.font = Font(name="宋体")
                if cell.column == sum_col_idx:
                    cell.number_format = '#,##0.00'

            val_str = str(cell.value) if cell.value is not None else ""
            str_len = len(val_str.encode('gbk', errors='ignore')) if val_str else 0
            if str_len > max_len:
                max_len = str_len

        adjusted_width = min(max_len * 1.2, 40)
        adjusted_width = max(adjusted_width, 14)
        ws.column_dimensions[col_letter].width = adjusted_width


def process_prepay(uploaded_files):
    """
    处理预付款业务周报文件。

    Args:
        uploaded_files: list of BytesIO file objects from st.file_uploader
            文件名含"汇总"的会被自动识别为周度汇总文件；
            其余 .xlsx 文件为数据文件（需含 Sheet "预付款业务情况周报表"）。

    Returns:
        (excel_bytes, docx_bytes, summary_bytes, logs)
    """
    # ==========================================
    # 0. 文件分类
    # ==========================================
    data_files = []
    summary_file = None
    for f in uploaded_files:
        fname = f.name if hasattr(f, 'name') else ''
        if '汇总' in fname:
            summary_file = f
        else:
            data_files.append(f)

    logs = []
    if not data_files:
        logs.append("⚠️ 未找到预付款数据文件，请上传包含\"预付款\"及品种关键词（粮谷/玉米/大豆）的 Excel 文件。")
        return None, None, None, logs

    # ==========================================
    # 1. 数据合并（原 process_and_merge 核心逻辑）
    # ==========================================
    out_wb = Workbook()
    out_ws = out_wb.active
    out_ws.title = SHEET_NAME

    all_data_rows = []
    header_copied = False
    master_header_end = 1

    col_widths = {}

    # 列索引追踪变量
    sort_col_1_idx = -1
    sort_col_2_idx = -1
    serial_col_idx = -1
    subtotal_cols_idx = []
    descriptive_cols_idx = []
    summary_match_cols_idx = []
    actual_max_col = 1

    for f in data_files:
        fname = getattr(f, 'name', 'unknown')
        logs.append(f"正在读取文件: {fname}")

        f.seek(0)
        file_bytes = f.read()
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True)

        if SHEET_NAME not in wb.sheetnames:
            logs.append(f"警告：文件 {fname} 中未找到表单「{SHEET_NAME}」，跳过。")
            wb.close()
            continue

        ws = wb[SHEET_NAME]
        h_start, h_end = score_and_find_header(ws)

        if h_start == -1:
            logs.append(f"警告：文件 {fname} 前10行未能识别出标题行，跳过。")
            wb.close()
            continue

        max_col = ws.max_column

        # --- 首份有效文件：复制标题行 ---
        if not header_copied:
            for col_idx in range(1, max_col + 1):
                col_letter = get_column_letter(col_idx)
                if ws.column_dimensions[col_letter].width:
                    col_widths[col_letter] = ws.column_dimensions[col_letter].width

            actual_max_col = 1
            for c_idx in range(1, max_col + 1):
                if any(ws.cell(row=r_idx, column=c_idx).value not in (None, "")
                       for r_idx in range(h_start, h_end + 1)):
                    actual_max_col = c_idx

            for r in range(1, h_end + 1):
                out_ws.row_dimensions[r].height = ws.row_dimensions[r].height
                for c in range(1, max_col + 1):
                    src_cell = ws.cell(row=r, column=c)
                    tgt_cell = out_ws.cell(row=r, column=c)
                    tgt_cell.value = src_cell.value
                    copy_cell_style(src_cell, tgt_cell)

            for merged_cell_range in ws.merged_cells.ranges:
                if merged_cell_range.min_row <= h_end:
                    out_ws.merge_cells(str(merged_cell_range))

            header_copied = True
            master_header_end = h_end

            # 特殊列（合同数量、预付时间等）标蓝色底纹
            special_fill_cols = []
            for c_idx in range(1, actual_max_col + 1):
                for r_idx in range(1, h_end + 1):
                    h_val = str(ws.cell(row=r_idx, column=c_idx).value or "").strip()
                    if any(kw in h_val for kw in ["合同数量", "预付时间", "付款天数", "逾期天数", "备注"]):
                        if c_idx not in special_fill_cols:
                            special_fill_cols.append(c_idx)
                        break

            for r in range(h_start, h_end + 1):
                for c in range(1, actual_max_col + 1):
                    cell = out_ws.cell(row=r, column=c)
                    if c in special_fill_cols:
                        cell.fill = PatternFill(fill_type="solid", start_color="D9E1F4", end_color="D9E1F4")
                    else:
                        cell.fill = PatternFill(fill_type="solid", start_color="DEEBF6", end_color="DEEBF6")
                    header_font = copy(cell.font)
                    header_font.color = Color(rgb="00000000")
                    cell.font = header_font

            # 更新截止日期
            last_wed = get_last_wednesday()
            formatted_date = f"{last_wed.year}年{last_wed.month}月{last_wed.day}日"
            for r in range(1, h_end + 1):
                for c in range(1, max_col + 1):
                    cell = out_ws.cell(row=r, column=c)
                    if isinstance(cell.value, str) and "截止日期" in cell.value:
                        cell.value = re.sub(
                            r'截止日期[:：]\s*\d{4}年\d{1,2}月\d{1,2}日',
                            f'截止日期：{formatted_date}',
                            cell.value
                        )

            # 检测各业务列索引
            for r_idx in range(h_start, h_end + 1):
                for c_idx in range(1, max_col + 1):
                    h_name = str(ws.cell(row=r_idx, column=c_idx).value or "").replace('\n', '').strip()
                    if not h_name:
                        continue

                    if "期末预付金额" in h_name:
                        sort_col_1_idx = c_idx - 1
                    if "合同约定货物交付时间" in h_name:
                        sort_col_2_idx = c_idx - 1
                    if "序号" in h_name:
                        serial_col_idx = c_idx - 1

                    if any(kw in h_name for kw in ["预付款数量", "合同金额", "拟交易金额", "本年累计预付金额", "期末预付金额"]):
                        if c_idx not in subtotal_cols_idx:
                            subtotal_cols_idx.append(c_idx)
                        if (c_idx - 1) not in summary_match_cols_idx:
                            summary_match_cols_idx.append(c_idx - 1)

                    if any(kw in h_name for kw in ["客商", "供应商", "平台名称", "收款方名称", "预付款性质"]):
                        if (c_idx - 1) not in descriptive_cols_idx:
                            descriptive_cols_idx.append(c_idx - 1)

        # --- 提取数据行 ---
        for r in range(h_end + 1, ws.max_row + 1):
            row_values = [ws.cell(row=r, column=c).value for c in range(1, max_col + 1)]

            row_str = "".join([str(v).strip() for v in row_values if v is not None])
            if not row_str or any(keyword in row_str for keyword in ["合计", "汇总", "总计", "小计"]):
                continue

            # 跳过末尾汇总行
            if descriptive_cols_idx and summary_match_cols_idx:
                descriptive_empty = all(
                    ws.cell(row=r, column=idx + 1).value in (None, "")
                    for idx in descriptive_cols_idx
                )
                summary_has_value = any(
                    isinstance(ws.cell(row=r, column=idx + 1).value, (int, float))
                    for idx in summary_match_cols_idx
                )
                tail_rows_blank = all(
                    ws.cell(row=rr, column=c).value in (None, "")
                    for rr in range(r + 1, ws.max_row + 1)
                    for c in range(1, max_col + 1)
                )
                if descriptive_empty and summary_has_value and tail_rows_blank:
                    continue

            row_data = []
            for c in range(1, max_col + 1):
                src_cell = ws.cell(row=r, column=c)
                cell_info = {
                    "value": src_cell.value,
                    "font": copy(src_cell.font),
                    "border": copy(src_cell.border),
                    "fill": copy(src_cell.fill),
                    "alignment": copy(src_cell.alignment),
                    "number_format": src_cell.number_format,
                    "height": ws.row_dimensions[r].height
                }
                row_data.append(cell_info)
            all_data_rows.append(row_data)

        wb.close()

    if not all_data_rows:
        logs.append("⚠️ 所有文件中均未提取到有效数据行。")
        return None, None, None, logs

    # --- 多条件排序 ---
    def sort_key(row):
        val1 = row[sort_col_1_idx]['value'] if sort_col_1_idx != -1 else float('inf')
        val2 = row[sort_col_2_idx]['value'] if sort_col_2_idx != -1 else datetime.datetime.max

        if val1 is None:
            val1 = float('inf')
        elif not isinstance(val1, (int, float)):
            try:
                val1 = float(val1)
            except Exception:
                val1 = float('inf')

        if val2 is None:
            val2 = datetime.datetime.max
        elif not isinstance(val2, datetime.datetime):
            val2 = datetime.datetime.max

        return (val1, val2)

    all_data_rows.sort(key=sort_key)

    # --- 写入合并数据到输出 sheet ---
    start_row = master_header_end + 1
    current_out_row = start_row

    for row_data in all_data_rows:
        if row_data[0]['height']:
            out_ws.row_dimensions[current_out_row].height = row_data[0]['height']

        for col_idx, cell_info in enumerate(row_data, start=1):
            tgt_cell = out_ws.cell(row=current_out_row, column=col_idx)

            if (col_idx - 1) == serial_col_idx:
                tgt_cell.value = f"=SUBTOTAL(103, $D${start_row}:D{current_out_row})"
            else:
                tgt_cell.value = cell_info['value']

            tgt_cell.font = cell_info['font']
            tgt_cell.border = cell_info['border']
            tgt_cell.fill = PatternFill()
            tgt_cell.alignment = cell_info['alignment']
            tgt_cell.number_format = cell_info['number_format']

        current_out_row += 1

    # --- 合计行 ---
    if all_data_rows:
        summary_row = current_out_row
        out_ws.row_dimensions[summary_row].height = out_ws.row_dimensions[summary_row - 1].height

        for col_idx in range(1, actual_max_col + 1):
            prev_cell = out_ws.cell(row=summary_row - 1, column=col_idx)
            tgt_cell = out_ws.cell(row=summary_row, column=col_idx)
            copy_cell_style(prev_cell, tgt_cell)
            tgt_cell.value = None

            if col_idx == 1:
                tgt_cell.value = "合计"
            elif col_idx in subtotal_cols_idx:
                col_letter = get_column_letter(col_idx)
                tgt_cell.value = f"=SUBTOTAL(9, {col_letter}{start_row}:{col_letter}{summary_row - 1})"
                old_font = tgt_cell.font
                tgt_cell.font = Font(
                    name=old_font.name, sz=old_font.sz, b=True, i=old_font.i, strike=old_font.strike,
                    color=old_font.color, vertAlign=old_font.vertAlign, underline=old_font.underline,
                    outline=old_font.outline, shadow=old_font.shadow, condense=old_font.condense,
                    extend=old_font.extend, family=old_font.family, charset=old_font.charset, scheme=old_font.scheme,
                )
            tgt_cell.fill = PatternFill(fill_type="solid", start_color="DEEBF6", end_color="DEEBF6")

    # --- 冻结窗格 ---
    freeze_row = master_header_end + 1
    out_ws.freeze_panes = f"A{freeze_row}"
    out_ws.sheet_view.selection[0].activeCell = 'A1'
    out_ws.sheet_view.selection[0].sqref = 'A1'

    # --- 恢复列宽 ---
    for col_letter, width in col_widths.items():
        out_ws.column_dimensions[col_letter].width = width

    # ==========================================
    # 2. 生成保证金、货款子表
    # ==========================================
    logs.append("正在生成保证金、货款表单…")

    # 构建列映射
    col_map = {}
    for c in range(1, actual_max_col + 1):
        header_text = "".join([str(out_ws.cell(row=r, column=c).value or "").strip()
                               for r in range(1, master_header_end + 1)])
        if "客商" in header_text or "供应商" in header_text:
            col_map['supplier'] = c - 1
        if "区域" in header_text or "经营部" in header_text:
            col_map['dept'] = c - 1
        if "期末预付" in header_text:
            col_map['balance'] = c - 1
        if "性质" in header_text:
            col_map['nature'] = c - 1
        if "预付时间" in header_text:
            col_map['time'] = c - 1
        if "品种" in header_text:
            col_map['type'] = c - 1
        if "收款方" in header_text:
            col_map['payee'] = c - 1
        if "平台" in header_text:
            col_map['platform'] = c - 1
        if "合同约定货物交付时间" in header_text:
            col_map['delivery'] = c - 1

    margin_data = []
    goods_data = []

    sum_all = 0

    for row in all_data_rows:
        try:
            balance_val = float(row[col_map.get('balance', -1)]['value'] or 0)
        except (TypeError, ValueError):
            balance_val = 0

        nature_val = str(row[col_map.get('nature', -1)]['value'] or "")

        if balance_val > 0:
            rounded_balance = balance_val / 10000
            sum_all += rounded_balance

            if "保证金" in nature_val:
                margin_row = [
                    row[col_map.get('supplier', -1)]['value'] if 'supplier' in col_map else None,
                    row[col_map.get('dept', -1)]['value'] if 'dept' in col_map else None,
                    rounded_balance,
                    row[col_map.get('time', -1)]['value'] if 'time' in col_map else None,
                    row[col_map.get('type', -1)]['value'] if 'type' in col_map else None,
                    row[col_map.get('payee', -1)]['value'] if 'payee' in col_map else None,
                    "平台冻结"
                ]
                margin_data.append(margin_row)

            if "货款" in nature_val:
                goods_row = [
                    row[col_map.get('supplier', -1)]['value'] if 'supplier' in col_map else None,
                    row[col_map.get('platform', -1)]['value'] if 'platform' in col_map else None,
                    row[col_map.get('dept', -1)]['value'] if 'dept' in col_map else None,
                    rounded_balance,
                    row[col_map.get('time', -1)]['value'] if 'time' in col_map else None,
                    row[col_map.get('payee', -1)]['value'] if 'payee' in col_map else None,
                    row[col_map.get('delivery', -1)]['value'] if 'delivery' in col_map else None,
                    None,
                    None
                ]
                goods_data.append(goods_row)

    # 排序子表
    def sort_new_sheets_data(row, balance_idx, time_idx):
        val1 = row[balance_idx] if isinstance(row[balance_idx], (int, float)) else -float('inf')
        val2 = row[time_idx]
        if not isinstance(val2, datetime.datetime):
            val2 = datetime.datetime.min
        return (val1, val2)

    margin_data.sort(key=lambda x: sort_new_sheets_data(x, 2, 3), reverse=True)
    goods_data.sort(key=lambda x: sort_new_sheets_data(x, 3, 4), reverse=True)

    # 创建保证金 sheet
    ws_margin = out_wb.create_sheet("保证金")
    headers_margin = ['客商/平台名称', '经营部/区域', '预付余额（万元）', '预付时间', '品种', '收款方', '合同约定退回保证金时间']
    ws_margin.append(headers_margin)
    for row in margin_data:
        ws_margin.append(row)
    sum_margin = sum(r[2] for r in margin_data if isinstance(r[2], (int, float)))
    ws_margin.append(['合计', None, sum_margin, None, None, None, None])
    apply_new_sheet_style(ws_margin, sum_col_idx=3)

    # 创建货款 sheet
    ws_goods = out_wb.create_sheet("货款")
    headers_goods = ['供应商', '平台名称', '经营部/区域', '预付余额（万元）', '付款时间', '收款方', '合同约定的交货时间', '逾期金额（万元）', '备注']
    ws_goods.append(headers_goods)
    for row in goods_data:
        ws_goods.append(row)
    sum_goods = sum(r[3] for r in goods_data if isinstance(r[3], (int, float)))
    ws_goods.append(['合计', None, None, sum_goods, None, None, None, None, None])
    apply_new_sheet_style(ws_goods, sum_col_idx=4)

    # ==========================================
    # 3. 处理周度汇总文件
    # ==========================================
    summary_bytes = None
    if summary_file is not None:
        logs.append("正在更新预付款业务周度明细汇总…")
        summary_bytes = _update_summary_in_memory(summary_file, sum_all, sum_goods, sum_margin, logs)
    else:
        logs.append("未上传汇总文件，跳过周度明细更新。")

    # ==========================================
    # 4. 生成 DOCX 报告
    # ==========================================
    logs.append("正在生成 DOCX 报告…")
    docx_bytes = _generate_docx_in_memory(margin_data, goods_data,
                                          sum_all, sum_goods, sum_margin,
                                          summary_file)

    # ==========================================
    # 5. 输出 Excel
    # ==========================================
    excel_bytes = io.BytesIO()
    out_wb.save(excel_bytes)
    excel_bytes.seek(0)

    logs.append(f"✅ 处理完成！共合并 {len(data_files)} 个数据文件，{len(all_data_rows)} 条记录。")
    return excel_bytes, docx_bytes, summary_bytes, logs


# ==========================================
# 内部辅助：更新周度汇总文件（内存版）
# ==========================================
def _update_summary_in_memory(summary_file, sum_all, sum_goods, sum_margin, logs):
    """将本周合计值写入汇总文件的第一个空白行，返回更新后的 BytesIO"""
    try:
        summary_file.seek(0)
        wb = load_workbook(io.BytesIO(summary_file.read()))
        ws = wb["data"]

        target_cols = {
            3: sum_all,
            7: sum_goods,
            8: sum_margin,
        }
        col_names = {3: "本周", 7: "货款余额", 8: "保证金余额"}

        for col_idx, value in target_cols.items():
            first_empty = None
            for r in range(2, ws.max_row + 2):
                if ws.cell(row=r, column=col_idx).value is None:
                    first_empty = r
                    break
            if first_empty is None:
                first_empty = ws.max_row + 1
            cell = ws.cell(row=first_empty, column=col_idx)
            cell.value = value
            cell.number_format = '#,##0.00'
            logs.append(f"  {col_names[col_idx]}({get_column_letter(col_idx)}{first_empty}) <- {value:,.2f}")

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        wb.close()
        return output
    except Exception as e:
        logs.append(f"⚠️ 汇总文件更新失败: {str(e)}")
        return None


# ==========================================
# 内部辅助：生成 DOCX 报告（内存版）
# ==========================================
def _generate_docx_in_memory(margin_data, goods_data,
                              sum_all, sum_goods, sum_margin,
                              summary_file):
    """生成 DOCX 报告到 BytesIO"""

    # --- 读取历史数据（用于环比） ---
    prev_week = prev_goods = prev_margin = 0.0
    cur_week = sum_all
    cur_goods = sum_goods
    cur_margin = sum_margin

    if summary_file is not None:
        try:
            summary_file.seek(0)
            wb = load_workbook(io.BytesIO(summary_file.read()), data_only=True)
            ws = wb["data"]
            last_row = None
            for r in range(2, ws.max_row + 1):
                if ws.cell(row=r, column=3).value is not None:
                    last_row = r
            if last_row is not None:
                prev_week = float(ws.cell(row=last_row, column=3).value or 0)
                prev_goods = float(ws.cell(row=last_row, column=7).value or 0)
                prev_margin = float(ws.cell(row=last_row, column=8).value or 0)
            wb.close()
        except Exception:
            pass

    # --- 环比计算 ---
    def _change(cur, prev):
        diff = cur - prev
        if diff == 0:
            pct = 0
            direction = "持平"
        else:
            pct = round(abs(diff) / prev * 100) if prev != 0 else 100
            direction = "增加" if diff > 0 else "减少"
        return diff, pct, direction

    week_diff, week_pct, week_dir = _change(cur_week, prev_week)
    goods_diff, goods_pct, goods_dir = _change(cur_goods, prev_goods)
    margin_diff, margin_pct, margin_dir = _change(cur_margin, prev_margin)

    trade_cur = cur_week - cur_margin - cur_goods
    trade_prev = prev_week - prev_margin - prev_goods
    has_trade = trade_cur > 0.005

    report_date = get_last_wednesday()
    date_str = f"{report_date.month}月{report_date.day}日"

    # ========== helpers: 文本段格式 ==========
    def _text_run_font(run, bold, east_asian, color_red=False):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = bold
        if color_red:
            run.font.color.rgb = RGBColor(255, 0, 0)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), east_asian)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')

    def _text_fmt_para(para, line_spacing):
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Pt(28)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _add_text_para(doc, text_data, east_asian="仿宋_GB2312", bold=False, line_spacing=22):
        p = doc.add_paragraph()
        _text_fmt_para(p, line_spacing)

        if isinstance(text_data, str):
            segments = [(text_data, bold, False)]
        else:
            segments = text_data

        for text, seg_bold, seg_red in segments:
            run = p.add_run(text)
            _text_run_font(run, seg_bold, east_asian, color_red=seg_red)
        return p

    # ========== helpers: 表格格式 ==========
    def _tbl_run_font(run, bold=False):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(7.5)
        run.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')

    def _tbl_fmt_para(para):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(10)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    def _write_cell(cell, text, bold=False, shade=False):
        p = cell.paragraphs[0]
        p.clear()
        _tbl_fmt_para(p)
        run = p.add_run(str(text) if text is not None else '')
        _tbl_run_font(run, bold)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if shade:
            tcPr = cell._element.get_or_add_tcPr()
            old = tcPr.find(qn('w:shd'))
            if old is not None:
                tcPr.remove(old)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F2F2F2')
            tcPr.append(shd)

    def _fmt_val(val):
        if val is None:
            return ''
        if isinstance(val, datetime.datetime):
            return val.strftime('%Y/%m/%d')
        if isinstance(val, (int, float)):
            return f'{val:,.2f}'
        return str(val)

    def _tr_height(tr):
        tr.height = Cm(1.2)
        tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    def _tbl_borders(table):
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        old = tblPr.find(qn('w:tblBorders'))
        if old is not None:
            tblPr.remove(old)
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            borders.append(el)
        tblPr.append(borders)

    # ========== 构建文档 ==========
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(3)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)

    # --- 文本段 ---
    _add_text_para(doc, "（二）预付业务情况", east_asian="黑体", bold=False, line_spacing=28)

    intro_segments = [
        (f"截至{date_str}，沿海大区预付款余额", False, False),
        (f"{cur_week:,.2f}万元", False, True),
        (f"，环比上期的{prev_week:,.2f}万元{'持平' if week_dir == '持平' else f'{week_dir}{week_pct}%或{abs(week_diff):,.2f}万元'}。其中，政策粮购销业务预付保证金余额", False, False),
        (f"{cur_margin:,.2f}万元", False, True),
        ("、货款余额", False, False),
        (f"{cur_goods:,.2f}万元", False, True),
        ("；", False, False)
    ]

    if has_trade:
        intro_segments.extend([
            ("贸易粮购销业务预付余额", False, False),
            (f"{trade_cur:,.2f}万元", False, False),
            ("、具体情况如下：", False, False)
        ])
    else:
        intro_segments.append(("无贸易粮购销业务预付。具体情况如下：", False, False))

    _add_text_para(doc, intro_segments)

    _add_text_para(doc, "1、贸易性预付款", bold=True)
    _add_text_para(doc, "无")

    _add_text_para(doc, "2、政策性保证金预付款", bold=True)
    margin_body_segments = [
        (f"截至{date_str}，政策性保证金预付款余额", False, False),
        (f"{cur_margin:,.2f}万元", False, True),
        (f"，环比上期{prev_margin:,.2f}万元{'持平' if margin_dir == '持平' else f'{margin_dir}{abs(margin_diff):,.2f}万元'}。", False, False)
    ]
    p_margin = _add_text_para(doc, margin_body_segments)
    p_margin.paragraph_format.space_after = Pt(14)

    # --- 保证金表 ---
    m_heads = ['客商/平台名称', '经营部/区域', '预付余额（万元）', '预付时间', '品种', '收款方', '合同约定退回保证金时间']
    m_widths = [Cm(3.02), Cm(1.51), Cm(2.56), Cm(2.33), Cm(0.79), Cm(3.01), Cm(2.02)]
    m_sum_idx = 2
    m_total = sum(r[m_sum_idx] for r in margin_data if isinstance(r[m_sum_idx], (int, float)))
    nrows = 1 + len(margin_data) + 1

    m_tbl = doc.add_table(rows=nrows, cols=7)
    m_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(m_widths):
        for row in m_tbl.rows:
            row.cells[i].width = w

    for i, h in enumerate(m_heads):
        _write_cell(m_tbl.rows[0].cells[i], h, bold=True, shade=True)
    _tr_height(m_tbl.rows[0])

    for r, row_data in enumerate(margin_data):
        tr = m_tbl.rows[r + 1]
        for c in range(7):
            val = row_data[c] if c < len(row_data) else None
            _write_cell(tr.cells[c], _fmt_val(val), bold=False)
        _tr_height(tr)

    for c in range(7):
        if c == 0:
            _write_cell(m_tbl.rows[-1].cells[c], '合计', bold=True, shade=True)
        elif c == m_sum_idx:
            _write_cell(m_tbl.rows[-1].cells[c], f'{m_total:,.2f}', bold=True, shade=True)
        else:
            _write_cell(m_tbl.rows[-1].cells[c], '', bold=True, shade=True)
    _tr_height(m_tbl.rows[-1])
    _tbl_borders(m_tbl)

    doc.add_paragraph()

    _add_text_para(doc, "3、政策性货款预付款", bold=True)
    goods_body_segments = [
        (f"截至{date_str}，政策性货款预付款余额", False, False),
        (f"{cur_goods:,.2f}万元", False, True),
        (f"，环比上期{prev_goods:,.2f}万元{'持平' if goods_dir == '持平' else f'{goods_dir}{abs(goods_diff):,.2f}万元'}，无逾期。", False, False)
    ]
    p_goods = _add_text_para(doc, goods_body_segments)
    p_goods.paragraph_format.space_after = Pt(14)

    # --- 货款表 (前7列) ---
    g_heads = ['供应商', '平台名称', '经营部/区域', '预付余额（万元）', '付款时间', '收款方', '合同约定的交货时间']
    g_widths = [Cm(2.66), Cm(1.88), Cm(1.56), Cm(2.20), Cm(2.00), Cm(2.73), Cm(2.21)]
    g_sum_idx = 3
    g_data = [row[:7] for row in goods_data]
    g_total = sum(r[g_sum_idx] for r in g_data if isinstance(r[g_sum_idx], (int, float)))
    nrows = 1 + len(g_data) + 1

    g_tbl = doc.add_table(rows=nrows, cols=7)
    g_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(g_widths):
        for row in g_tbl.rows:
            row.cells[i].width = w

    for i, h in enumerate(g_heads):
        _write_cell(g_tbl.rows[0].cells[i], h, bold=True, shade=True)
    _tr_height(g_tbl.rows[0])

    for r, row_data in enumerate(g_data):
        tr = g_tbl.rows[r + 1]
        for c in range(7):
            val = row_data[c] if c < len(row_data) else None
            _write_cell(tr.cells[c], _fmt_val(val), bold=False)
        _tr_height(tr)

    for c in range(7):
        if c == 0:
            _write_cell(g_tbl.rows[-1].cells[c], '合计', bold=True, shade=True)
        elif c == g_sum_idx:
            _write_cell(g_tbl.rows[-1].cells[c], f'{g_total:,.2f}', bold=True, shade=True)
        else:
            _write_cell(g_tbl.rows[-1].cells[c], '', bold=True, shade=True)
    _tr_height(g_tbl.rows[-1])
    _tbl_borders(g_tbl)

    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes
