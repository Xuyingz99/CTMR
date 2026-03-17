import pandas as pd
import numpy as np
import io
import warnings
import re
import os
import glob
import datetime
from decimal import Decimal, ROUND_HALF_UP

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import openpyxl
from openpyxl.styles import Alignment, Border, Side, Font, PatternFill

warnings.filterwarnings('ignore')

# ================= 配置与常量 =================
COLOR_BATCH = "E6F7FF"  
COLOR_ONCE = "FFF7E6"   

REASON1_ORDER = ['一、客户原因/客户原因为主', '二、我方原因/我方原因为主', '三、既非我方原因也非对方原因']
REASON2_ORDER = {
    '一、客户原因/客户原因为主': ['客户加工计划调整', '客户基差合同点价晚', '客户其他原因（需详细说明）', '客户仓容紧张', '客户资金紧张', '我方正常到货，客户提货能力不足', '客户额度不足', '客户工厂检修', '销价高客户不愿提货'],
    '二、我方原因/我方原因为主': ['一体化协同粮源质量问题', '外采粮源质量问题', '大区责任物流原因', '我方货源不足', '一体化非大区责任物流原因', '外采供应商责任物流原因', '我方其他原因（需详细说明）', '我方到货集中', '我方修路'],
    '三、既非我方原因也非对方原因': ['其他原因（需详细说明）', '天气原因', '政府行为', '自然灾害', '社会异常事件']
}
TIME_ORDER = ['1-10天', '11-20天', '21-30天', '31-60天', '61-90天', '90天以上']

# ================= 辅助函数 =================
def clean_text_for_match(text):
    """强力去污：清除所有可见/不可见空格，并将中文括号转为英文括号，专门用于匹配"""
    if pd.isna(text): return ""
    s = str(text).strip()
    s = re.sub(r'\s+', '', s)
    s = s.replace('（', '(').replace('）', ')')
    return s

def format_num(val, dec=2, is_int=False, is_percent=False):
    if pd.isna(val) or val == "": return ""
    try: d_val = Decimal(str(round(float(val), 6)))
    except: return val
    if is_percent:
        if d_val == 0: return "0%"
        elif abs(d_val) >= 1:
            rounded = d_val.quantize(Decimal('1'), rounding=ROUND_HALF_UP)
            return f"{int(rounded)}%"
        else:
            s_val = f"{float(abs(d_val)):.6f}"
            if '.' in s_val:
                dec_part = s_val.split('.')[1]
                non_zero_idx = next((i + 1 for i, digit in enumerate(dec_part) if digit != '0'), 0)
                q_str = '0.' + '0' * (non_zero_idx - 1) + '1' if non_zero_idx > 0 else '1'
                rounded = d_val.quantize(Decimal(q_str), rounding=ROUND_HALF_UP)
                res = f"{float(rounded)}".rstrip('0').rstrip('.')
                return f"{res}%"
            return f"{float(d_val)}%"
    if is_int:
        rounded = d_val.quantize(Decimal('1'), rounding=ROUND_HALF_UP)
        return f"{int(rounded):,}"
    q = Decimal('1.' + '0' * dec) if dec > 0 else Decimal('1')
    rounded = d_val.quantize(q, rounding=ROUND_HALF_UP)
    return f"{float(rounded):,.{dec}f}"

def format_qty(val):
    if pd.isna(val) or val == "": return ""
    try: d_val = Decimal(str(round(float(val), 6)))
    except: return val
    if d_val == 0: return "0"
    if abs(d_val) >= 1: return format_num(val, 2)
    q2 = Decimal('1.00')
    rounded2 = d_val.quantize(q2, rounding=ROUND_HALF_UP)
    if rounded2 != 0: return f"{float(rounded2):.2f}"
    s_val = f"{float(abs(d_val)):.10f}"
    if '.' in s_val:
        dec_part = s_val.split('.')[1]
        non_zero_idx = next((i + 1 for i, digit in enumerate(dec_part) if digit != '0'), 0)
        if non_zero_idx == 0: return "0"
        q_str = '0.' + '0' * (non_zero_idx - 1) + '1'
        rounded = d_val.quantize(Decimal(q_str), rounding=ROUND_HALF_UP)
        return f"{float(rounded)}".rstrip('0').rstrip('.')
    return f"{float(d_val)}"

def map_reason1(x):
    if pd.isna(x): return x
    if "客户原因" in str(x): return "一、客户原因/客户原因为主"
    if "我方原因" in str(x): return "二、我方原因/我方原因为主"
    if "既非我方" in str(x) or "非对方" in str(x): return "三、既非我方原因也非对方原因"
    return str(x)

# ================= 数据清洗 =================
def locate_header_and_read_stream(file_stream, key_columns):
    try:
        file_stream.seek(0)
        df_raw = pd.read_excel(file_stream, header=None)
        header_row_index = -1
        for i, row in df_raw.iterrows():
            row_values = [str(x).strip().replace('\n', '').replace(' ', '') for x in row.values if pd.notna(x)]
            match_count = sum(1 for key in key_columns if key in row_values)
            if match_count >= len(key_columns) - 1:
                header_row_index = i
                break
        if header_row_index == -1: return None
        file_stream.seek(0)
        df = pd.read_excel(file_stream, header=header_row_index)
        df.columns = df.columns.astype(str).str.replace('\n', '', regex=False).str.strip()
        if '大区' in df.columns:
            col_idx = df.columns.get_loc('大区')
            df = df.iloc[:, col_idx:]
        df.dropna(how='all', inplace=True)
        return df
    except Exception: return None

def process_basic_columns(df, date_cols, float_cols, int_cols=None):
    for col in date_cols:
        if col in df.columns:
            temp_dates = pd.to_datetime(df[col], errors='coerce')
            mask_invalid = (temp_dates.dt.year < 2025) | (temp_dates.dt.year > 2030) | (temp_dates.isna())
            if mask_invalid.any():
                raw_values = df.loc[mask_invalid, col]
                clean_raw = raw_values.astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
                corrected = pd.to_datetime(clean_raw, format='%Y%m%d', errors='coerce')
                temp_dates.loc[mask_invalid] = corrected
            df[col] = temp_dates
    for col in float_cols:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
    if int_cols:
        for col in int_cols:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
    return df

# ================= Word 周报排版模块 =================
def set_font_mixed(run_or_style, size_pt, bold=False, east_asia='仿宋_GB2312', ascii_font='Times New Roman'):
    run_or_style.font.size = Pt(size_pt)
    run_or_style.font.name = ascii_font
    run_or_style._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    run_or_style._element.rPr.rFonts.set(qn('w:ascii'), ascii_font)
    run_or_style._element.rPr.rFonts.set(qn('w:hAnsi'), ascii_font)
    run_or_style.font.bold = bold

def init_styles(doc):
    style_ch = doc.styles.add_style('ChapterTitle', 1)
    set_font_mixed(style_ch, 14.0, False, '黑体', 'Times New Roman')
    style_ch.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_ch.paragraph_format.line_spacing = Pt(28)
    style_ch.paragraph_format.space_before = Pt(0)
    style_ch.paragraph_format.space_after = Pt(0)
    style_ch.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_ch.paragraph_format.first_line_indent = Pt(28)

    style_tb = doc.styles.add_style('TableTitle', 1)
    set_font_mixed(style_tb, 12.0, False, '微软雅黑', 'Times New Roman')
    style_tb.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_tb.paragraph_format.line_spacing = Pt(22)
    style_tb.paragraph_format.space_before = Pt(0)
    style_tb.paragraph_format.space_after = Pt(0)
    style_tb.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    style_nm = doc.styles.add_style('NormalContent', 1)
    set_font_mixed(style_nm, 14.0, False, '仿宋_GB2312', 'Times New Roman')
    style_nm.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_nm.paragraph_format.line_spacing = Pt(22)
    style_nm.paragraph_format.space_before = Pt(0)
    style_nm.paragraph_format.space_after = Pt(0)
    style_nm.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_nm.paragraph_format.first_line_indent = Pt(28) 

    style_app = doc.styles.add_style('AppendixTitle', 1)
    set_font_mixed(style_app, 15.0, False, '黑体', 'Times New Roman')
    style_app.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_app.paragraph_format.line_spacing = Pt(28)
    style_app.paragraph_format.space_before = Pt(0)
    style_app.paragraph_format.space_after = Pt(0)
    style_app.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_page_margins(doc):
    section = doc.sections[0]
    section.page_width = Cm(266710.5 / 12700.0)
    section.page_height = Cm(377194.0 / 12700.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

def build_cell_text(cell, text, align='center', bold=False, is_max=False, is_appendix=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    if is_appendix:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(12)
        font_sz = 7.5
        en_font_sz = 7.5
    else:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(10)
        font_sz = 9.0
        en_font_sz = 10.0

    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left': p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    text_str = str(text) if not pd.isna(text) else ""
    parts = re.split(r'([a-zA-Z0-9.,%+-]+)', text_str)
    for part in parts:
        if not part: continue
        run = p.add_run(part)
        if is_max: run.font.color.rgb = RGBColor(255, 0, 0)
        if re.match(r'^[a-zA-Z0-9.,%+-]+$', part):
            set_font_mixed(run, en_font_sz, bold, '微软雅黑', 'Times New Roman')
        else:
            set_font_mixed(run, font_sz, bold, '微软雅黑', 'Times New Roman')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_cell_background(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def apply_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') 
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tblCellMar = OxmlElement('w:tblCellMar')
    for m in [('top', '28'), ('bottom', '28'), ('left', '57'), ('right', '57')]:
        node = OxmlElement(f'w:{m[0]}')
        node.set(qn('w:w'), m[1])
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)
    tblPr.append(tblCellMar)

def set_fixed_col_widths(table, widths, is_cm=False):
    table.autofit = False
    table.allow_autofit = False
    for i, w_val in enumerate(widths):
        w = Cm(w_val) if is_cm else Cm(w_val / 12700.0)
        table.columns[i].width = w
        for cell in table.columns[i].cells: cell.width = w

def set_table_row_height(row, height_pt):
    row.height = Pt(height_pt)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

# ================= 完整 Word 生成 =================
def generate_word_report(df, df_unique):
    if '原因分类1' in df_unique.columns:
        df_unique['标准原因分类1'] = df_unique['原因分类1'].apply(map_reason1)

    total_amount = df_unique['逾期金额（万元）'].sum()
    safe_total = total_amount if total_amount > 0 else 1e-9
    total_amount_str = format_num(total_amount, 0, True)

    doc = Document()
    set_page_margins(doc)
    init_styles(doc)

    # ----- (三) 逾期销售天数 -----
    doc.add_paragraph('（三）逾期销售天数', style='ChapterTitle')
    total_qty = df_unique['逾期数量（万吨）'].sum()
    safe_total_qty = total_qty if total_qty > 0 else 1e-9
    avg_days_val = (df_unique['逾期数量（万吨）'] * df_unique['逾期天数']).sum() / safe_total_qty if total_qty > 0 else 0
    avg_days = int(Decimal(str(round(avg_days_val, 6))).quantize(Decimal('1'), rounding=ROUND_HALF_UP))
    
    time_stats = df_unique.groupby('逾期天数分类').agg({'逾期金额（万元）': 'sum', '合同编号': 'count', '逾期数量（万吨）': 'sum'}).reindex(TIME_ORDER).fillna(0)
    
    p1 = doc.add_paragraph(style='NormalContent')
    run_avg1 = p1.add_run(f"平均逾期{avg_days}天，")
    set_font_mixed(run_avg1, 14.0, bold=True)
    run_avg2 = p1.add_run("周环比无数据")
    set_font_mixed(run_avg2, 14.0, bold=True)
    run_avg2.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run_avg3 = p1.add_run("。")
    set_font_mixed(run_avg3, 14.0, bold=True)

    max_qty_cat = time_stats['逾期数量（万吨）'].idxmax()
    valid_times = [t for t in TIME_ORDER if time_stats.loc[t, '逾期金额（万元）'] > 0]
    for i, t in enumerate(valid_times):
        amt = time_stats.loc[t, '逾期金额（万元）']
        ratio = amt / safe_total * 100
        is_last = (i == len(valid_times) - 1)
        punctuation = "。" if is_last else "；"
        text_part = f"逾期天数在{t.replace('天', '')}天的，共涉及{format_num(amt, 0, True)}万元，占总逾期金额的{format_num(ratio, is_percent=True)}"
        run_part = p1.add_run(text_part + punctuation)
        is_max = (t == max_qty_cat)
        set_font_mixed(run_part, 14.0, bold=is_max)
        if is_max: run_part.font.color.rgb = RGBColor(255, 0, 0)
    p1.paragraph_format.space_after = Pt(0)

    doc.add_paragraph('逾期销售提货分时间情况表', style='TableTitle')
    table1 = doc.add_table(rows=1, cols=5)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_col_widths(table1, [2.72, 2.95, 2.95, 2.95, 2.95], is_cm=True)
        
    headers1 = ['逾期时间', '逾期金额\n（万元）', '逾期金额\n占比', '合同个数\n（笔）', '逾期数量\n（万吨）']
    for i, h in enumerate(headers1):
        build_cell_text(table1.cell(0, i), h, bold=True)
        set_cell_background(table1.cell(0, i), 'D9D9D9')
    set_table_row_height(table1.rows[0], Cm(0.62).pt)
    set_repeat_table_header(table1.rows[0])

    for t in TIME_ORDER:
        amt = time_stats.loc[t, '逾期金额（万元）']
        if amt > 0:
            row_cells = table1.add_row().cells
            is_max = (t == max_qty_cat)
            build_cell_text(row_cells[0], t, bold=is_max, is_max=is_max)
            build_cell_text(row_cells[1], format_num(amt, 0, True), bold=is_max, is_max=is_max)
            build_cell_text(row_cells[2], format_num(amt / safe_total * 100, is_percent=True), bold=is_max, is_max=is_max)
            build_cell_text(row_cells[3], format_num(time_stats.loc[t, '合同编号'], 0, True), bold=is_max, is_max=is_max)
            build_cell_text(row_cells[4], format_qty(time_stats.loc[t, '逾期数量（万吨）']), bold=is_max, is_max=is_max)
            set_table_row_height(table1.rows[-1], Cm(0.48).pt)

    tot_cells1 = table1.add_row().cells
    build_cell_text(tot_cells1[0], '总计', bold=True)
    build_cell_text(tot_cells1[1], total_amount_str, bold=True)
    build_cell_text(tot_cells1[2], '100%', bold=True)
    build_cell_text(tot_cells1[3], format_num(time_stats['合同编号'].sum(), 0, True), bold=True)
    build_cell_text(tot_cells1[4], format_qty(time_stats['逾期数量（万吨）'].sum()), bold=True)
    set_cell_background(tot_cells1[0], 'D9E1F4')
    for cell in tot_cells1[1:]: set_cell_background(cell, 'DEEBF6')
    set_table_row_height(table1.rows[-1], Cm(0.48).pt)
    apply_table_borders(table1)

    # ----- (四) 逾期销售原因 -----
    doc.add_paragraph('（四）逾期销售原因', style='ChapterTitle')
    if '标准原因分类1' not in df_unique.columns: df_unique['标准原因分类1'] = ''
    r1_stats = df_unique.groupby('标准原因分类1').agg({'逾期金额（万元）': 'sum'})
    max_r1 = r1_stats['逾期金额（万元）'].idxmax() if not r1_stats.empty else None
    
    p2 = doc.add_paragraph(style='NormalContent')
    r1_texts_count = sum(1 for r1 in REASON1_ORDER if r1 in r1_stats.index and r1_stats.loc[r1, '逾期金额（万元）'] > 0)
    current_idx = 0
    for r1 in REASON1_ORDER:
        amt = r1_stats.loc[r1, '逾期金额（万元）'] if r1 in r1_stats.index else 0
        if amt > 0:
            current_idx += 1
            prefix = "主要由客户原因造成的" if "客户原因" in r1 else ("主要由我方原因造成的" if "我方原因" in r1 else "既非我方原因也非对方原因造成的")
            text_part = f"{prefix}逾期金额{format_num(amt, 0, True)}万元，占比{format_num(amt/safe_total*100, is_percent=True)}"
            is_last = (current_idx == r1_texts_count)
            punctuation = "；详情如下：" if is_last else "；"
            run_part = p2.add_run(text_part + punctuation)
            is_max = (r1 == max_r1)
            set_font_mixed(run_part, 14.0, bold=is_max)
            if is_max: run_part.font.color.rgb = RGBColor(255, 0, 0)
    p2.paragraph_format.space_after = Pt(0)

    doc.add_paragraph('逾期销售提货分原因情况表', style='TableTitle')
    table2 = doc.add_table(rows=1, cols=5)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_col_widths(table2, [7.05, 1.99, 2.01, 1.76, 1.99], is_cm=True)
    
    headers2 = ['逾期原因分类', '逾期金额\n（万元）', '逾期金额\n占比', '合同笔数\n（笔）', '逾期数量\n（万吨）']
    for i, h in enumerate(headers2):
        build_cell_text(table2.cell(0, i), h, bold=True)
        set_cell_background(table2.cell(0, i), 'D9D9D9')
    set_table_row_height(table2.rows[0], Cm(0.3).pt)
    set_repeat_table_header(table2.rows[0])

    for r1 in REASON1_ORDER:
        r1_df = df_unique[df_unique['标准原因分类1'] == r1]
        r1_amt = r1_df['逾期金额（万元）'].sum()
        row_cells = table2.add_row().cells
        build_cell_text(row_cells[0], r1, align='left', bold=True)
        build_cell_text(row_cells[1], format_num(r1_amt, 0, True) if r1_amt>0 else '', align='right', bold=True)
        build_cell_text(row_cells[2], format_num(r1_amt/safe_total*100, is_percent=True) if r1_amt>0 else '', align='right', bold=True)
        build_cell_text(row_cells[3], format_num(len(r1_df), 0, True) if r1_amt>0 else '', align='right', bold=True)
        build_cell_text(row_cells[4], format_qty(r1_df['逾期数量（万吨）'].sum()) if r1_amt>0 else '', align='right', bold=True)
        for cell in row_cells: set_cell_background(cell, 'D9E1F4')
        set_table_row_height(table2.rows[-1], Cm(0.46).pt)

        if '原因分类2' not in r1_df.columns: continue
        r2_stats = r1_df.groupby('原因分类2')['逾期金额（万元）'].sum()
        r2_list = REASON2_ORDER.get(r1, [])
        r2_val_map = {r2: r2_stats.get(r2, 0) for r2 in r2_list}
        sorted_r2 = sorted(r2_list, key=lambda x: r2_val_map[x], reverse=True)

        for r2 in sorted_r2:
            r2_df = r1_df[r1_df['原因分类2'] == r2]
            r2_amt = r2_val_map[r2]
            row_cells = table2.add_row().cells
            build_cell_text(row_cells[0], r2, align='left')
            if r2_amt > 0:
                build_cell_text(row_cells[1], format_num(r2_amt, 0, True), align='right')
                build_cell_text(row_cells[2], format_num(r2_amt/safe_total*100, is_percent=True), align='right')
                build_cell_text(row_cells[3], format_num(len(r2_df), 0, True), align='right')
                build_cell_text(row_cells[4], format_qty(r2_df['逾期数量（万吨）'].sum()), align='right')
            else:
                for idx in range(1, 5): build_cell_text(row_cells[idx], '', align='right')
            set_table_row_height(table2.rows[-1], Cm(0.46).pt)

    tot_cells2 = table2.add_row().cells
    tot_cells2[0].merge(tot_cells2[0])
    build_cell_text(tot_cells2[0], '总计', bold=True)
    build_cell_text(tot_cells2[1], total_amount_str, bold=True)
    build_cell_text(tot_cells2[2], '100%', bold=True)
    build_cell_text(tot_cells2[3], format_num(len(df_unique), 0, True), bold=True)
    build_cell_text(tot_cells2[4], format_qty(df_unique['逾期数量（万吨）'].sum()), bold=True)
    for cell in tot_cells2: set_cell_background(cell, 'D9E1F4')
    set_table_row_height(table2.rows[-1], Cm(0.46).pt)
    apply_table_borders(table2)

    # ----- (五) 逾期销售分品种 -----
    doc.add_paragraph('（五）逾期销售分品种', style='ChapterTitle')
    variety_stats = df_unique.groupby('品种').agg({'逾期金额（万元）': 'sum', '合同编号': 'count', '逾期数量（万吨）': 'sum'}).sort_values(by='逾期金额（万元）', ascending=False)
    max_v = variety_stats['逾期金额（万元）'].idxmax() if not variety_stats.empty else None
    
    p3 = doc.add_paragraph(style='NormalContent')
    v_count = len(variety_stats)
    for i, v in enumerate(variety_stats.index):
        v_amt = variety_stats.loc[v, '逾期金额（万元）']
        text_part = f"{v}逾期金额为{format_num(v_amt, 0, True)}万元，占比{format_num(v_amt/safe_total*100, is_percent=True)}"
        is_last = (i == v_count - 1)
        punctuation = "。详情如下：" if is_last else "；"
        run_part = p3.add_run(text_part + punctuation)
        is_max = (v == max_v)
        set_font_mixed(run_part, 14.0, bold=is_max)
        if is_max: run_part.font.color.rgb = RGBColor(255, 0, 0)
    p3.paragraph_format.space_after = Pt(0)

    doc.add_paragraph('逾期销售提货分品种情况表', style='TableTitle')
    table3 = doc.add_table(rows=1, cols=5)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_col_widths(table3, [2.72, 2.95, 2.95, 2.95, 2.95], is_cm=True)
    
    headers3 = ['品种', '合同笔数', '逾期数量\n（万吨）', '逾期金额\n（万元）', '逾期金额\n占比']
    for i, h in enumerate(headers3):
        build_cell_text(table3.cell(0, i), h, bold=True)
        set_cell_background(table3.cell(0, i), 'D9D9D9')
    set_table_row_height(table3.rows[0], Cm(0.62).pt)
    set_repeat_table_header(table3.rows[0])

    for v in variety_stats.index:
        row_cells = table3.add_row().cells
        v_amt = variety_stats.loc[v, '逾期金额（万元）']
        is_max = (v == max_v)
        build_cell_text(row_cells[0], v, bold=is_max, is_max=is_max)
        build_cell_text(row_cells[1], format_num(variety_stats.loc[v, '合同编号'], 0, True), bold=is_max, is_max=is_max)
        build_cell_text(row_cells[2], format_qty(variety_stats.loc[v, '逾期数量（万吨）']), bold=is_max, is_max=is_max)
        build_cell_text(row_cells[3], format_num(v_amt, 0, True), bold=is_max, is_max=is_max)
        build_cell_text(row_cells[4], format_num(v_amt/safe_total*100, is_percent=True), bold=is_max, is_max=is_max)
        set_table_row_height(table3.rows[-1], Cm(0.48).pt)

    tot_cells3 = table3.add_row().cells
    tot_cells3[0].merge(tot_cells3[0])
    build_cell_text(tot_cells3[0], '总计', bold=True)
    build_cell_text(tot_cells3[1], format_num(len(df_unique), 0, True), bold=True)
    build_cell_text(tot_cells3[2], format_qty(df_unique['逾期数量（万吨）'].sum()), bold=True)
    build_cell_text(tot_cells3[3], total_amount_str, bold=True)
    build_cell_text(tot_cells3[4], '100%', bold=True)
    set_cell_background(tot_cells3[0], 'D9E1F4')
    for cell in tot_cells3[1:]: set_cell_background(cell, 'DEEBF6')
    set_table_row_height(table3.rows[-1], Cm(0.48).pt)
    apply_table_borders(table3)

    # ----- (六) 逾期销售分客户 -----
    def get_cust_type(row):
        grp = row.get('所属集团')
        if pd.notna(grp) and str(grp).strip() != '' and '中粮集团' not in str(grp): return '战略大客户', grp
        intr = row.get('集团内部客户')
        if pd.notna(intr) and str(intr).strip() != '': return '集团内部客户', intr
        return '中小客户', row.get('客户名称', '')
        
    df_unique['客户大类'], df_unique['展示客户名'] = zip(*df_unique.apply(get_cust_type, axis=1))
    c_stats = df_unique.groupby('客户大类')['逾期数量（万吨）'].sum().fillna(0)
    strat_total = c_stats.get('战略大客户', 0)
    mid_total = c_stats.get('中小客户', 0)
    int_total = c_stats.get('集团内部客户', 0)
    strat_cnt = df_unique[df_unique['客户大类'] == '战略大客户']['展示客户名'].nunique()
    mid_cnt = df_unique[df_unique['客户大类'] == '中小客户']['展示客户名'].nunique()
    int_cnt = df_unique[df_unique['客户大类'] == '集团内部客户']['展示客户名'].nunique()
    total_customers = strat_cnt + mid_cnt + int_cnt

    doc.add_paragraph('（六）逾期销售分客户', style='ChapterTitle')
    max_c_type = c_stats.idxmax() if not c_stats.empty else None
    
    p4 = doc.add_paragraph(style='NormalContent')
    run_base1 = p4.add_run(f"逾期提货客户共{total_customers}家，")
    set_font_mixed(run_base1, 14.0, bold=True)
    run_base2 = p4.add_run("周环比无数据")
    set_font_mixed(run_base2, 14.0, bold=True)
    run_base2.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run_base3 = p4.add_run("。包括")
    set_font_mixed(run_base3, 14.0, bold=True)
    
    c_parts = [
        ('战略大客户', f"{strat_cnt}家战略大客户共逾期{format_num(strat_total, 2)}万吨", "，"),
        ('中小客户', f"{mid_cnt}家中小客户共逾期{format_num(mid_total, 2)}万吨", "，"),
        ('集团内部客户', f"{int_cnt}家集团内部客户共逾期{format_num(int_total, 2)}万吨", "。具体情况如下表：")
    ]
    for c_type, text_part, punct in c_parts:
        run_part = p4.add_run(text_part + punct)
        is_max = (c_type == max_c_type and c_stats.get(c_type, 0) > 0)
        set_font_mixed(run_part, 14.0, bold=is_max)
        if is_max: run_part.font.color.rgb = RGBColor(255, 0, 0)
    p4.paragraph_format.space_after = Pt(0)

    doc.add_paragraph('逾期销售提货分客户明细表', style='TableTitle')
    p_unit = doc.add_paragraph()
    p_unit.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_unit.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_unit.paragraph_format.line_spacing = Pt(12)
    p_unit.paragraph_format.space_before = Pt(0)
    p_unit.paragraph_format.space_after = Pt(0)
    run_unit = p_unit.add_run('单位：万吨')
    set_font_mixed(run_unit, 9.0, bold=False, east_asia='微软雅黑', ascii_font='Times New Roman')

    table4 = doc.add_table(rows=1, cols=4)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_col_widths(table4, [1.6, 8.2, 2.8, 2.5], is_cm=True)
    
    headers4 = ['序号', '客户名称/所属集团', '品种', '逾期数量']
    for i, h in enumerate(headers4):
        build_cell_text(table4.cell(0, i), h, bold=True)
        set_cell_background(table4.cell(0, i), 'D9D9D9')
    set_table_row_height(table4.rows[0], Cm(0.73).pt)
    set_repeat_table_header(table4.rows[0])

    def add_cust_rows(c_type, subtotal_name, start_idx):
        sub_df = df_unique[df_unique['客户大类'] == c_type]
        if sub_df.empty: return start_idx
        agg_df = sub_df.groupby('展示客户名').agg({'逾期数量（万吨）': 'sum', '品种': lambda x: '、'.join(x.dropna().astype(str).unique())}).sort_values(by='逾期数量（万吨）', ascending=False).reset_index()
        for _, row in agg_df.iterrows():
            cells = table4.add_row().cells
            build_cell_text(cells[0], start_idx)
            build_cell_text(cells[1], row['展示客户名'])
            build_cell_text(cells[2], row.get('品种', ''))
            build_cell_text(cells[3], format_qty(row['逾期数量（万吨）']))
            set_table_row_height(table4.rows[-1], Cm(0.44).pt)
            start_idx += 1
        sub_cells = table4.add_row().cells
        sub_cells[0].merge(sub_cells[2])
        build_cell_text(sub_cells[0], subtotal_name, bold=True)
        build_cell_text(sub_cells[3], format_qty(agg_df['逾期数量（万吨）'].sum()), bold=True)
        for c in [sub_cells[0], sub_cells[3]]: set_cell_background(c, 'D9D9D9')
        set_table_row_height(table4.rows[-1], Cm(0.44).pt)
        return start_idx

    idx = 1
    idx = add_cust_rows('战略大客户', '战略客户小计', idx)
    idx = add_cust_rows('中小客户', '中小客户小计', idx)
    idx = add_cust_rows('集团内部客户', '集团内部客户小计', idx)

    tot_cells4 = table4.add_row().cells
    tot_cells4[0].merge(tot_cells4[2])
    build_cell_text(tot_cells4[0], '汇总', bold=True)
    build_cell_text(tot_cells4[3], format_qty(df_unique['逾期数量（万吨）'].sum()), bold=True)
    for c in [tot_cells4[0], tot_cells4[3]]: set_cell_background(c, 'DEEBF6')
    set_table_row_height(table4.rows[-1], Cm(0.44).pt)
    apply_table_borders(table4)

    # ----- 附表 -----
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Cm(377194.0 / 12700.0)
    new_section.page_height = Cm(266710.5 / 12700.0)
    new_section.left_margin = Cm(40322.44 / 12700.0)
    new_section.right_margin = Cm(40322.44 / 12700.0)
    new_section.top_margin = Cm(32257.95 / 12700.0)
    new_section.bottom_margin = Cm(32257.95 / 12700.0)

    def create_appendix(title, df_subset, table_idx):
        doc.add_paragraph(title, style='AppendixTitle')
        p_app_unit = doc.add_paragraph()
        p_app_unit.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_app_unit.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p_app_unit.paragraph_format.line_spacing = Pt(28)
        p_app_unit.paragraph_format.space_before = Pt(0)
        p_app_unit.paragraph_format.space_after = Pt(0)
        run_app_unit = p_app_unit.add_run('单位：万吨、元/吨、万元')
        set_font_mixed(run_app_unit, 14.0, bold=False, east_asia='仿宋_GB2312', ascii_font='Times New Roman')

        table = doc.add_table(rows=1, cols=15)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        if table_idx == 5: widths_app = [9386.17, 15792.96, 25313.53, 22154.94, 12656.77, 15792.96, 12656.77, 12656.77, 18055.49, 18077.89, 117674.32, 39224.77, 12656.77, 13194.4, 14157.66]
        elif table_idx == 6: widths_app = [9430.97, 15860.16, 25403.14, 19959.61, 12544.76, 16128.98, 13127.19, 12253.54, 17921.08, 17921.08, 118771.98, 38104.7, 12701.57, 12947.98, 12947.98]
        else: widths_app = [11021.47, 16128.98, 24775.9, 20430.04, 11962.32, 15815.36, 13440.81, 11962.32, 18212.3, 18212.3, 111021.11, 45385.14, 13127.19, 12253.54, 14516.08]
        set_fixed_col_widths(table, widths_app)

        headers = ['序号', '经营部', '客户名称', '交货结束日期', '合同\n数量', '合同\n单价', '逾期\n天数', '逾期数量', '原因\n分类1', '原因\n分类2', '具体逾期原因', '解决方案及解决时间', '责任人', '上级领导', '是否为赊销合同']
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            build_cell_text(cell, h, bold=True, is_appendix=True)
            set_cell_background(cell, 'D9D9D9')
        set_table_row_height(table.rows[0], Cm(0.9).pt)
        set_repeat_table_header(table.rows[0])

        if not df_subset.empty:
            df_subset = df_subset.sort_values(by='逾期天数', ascending=False).reset_index(drop=True)
            leader_map = {'珠三角': '黄旭东', '福建': '肖灿', '广西': '丁峰', '海南': '宋永伍', '粤西': '张文韬'}
            for i, row in df_subset.iterrows():
                cells = table.add_row().cells
                dept_val = str(row.get('经营部', '')) if pd.notna(row.get('经营部')) else ""
                dept_clean = dept_val.replace('经营部', '')
                leader_val = leader_map.get(dept_clean, str(row.get('上级领导', '')) if pd.notna(row.get('上级领导')) else "")
                
                reason_detail = str(row.get('具体逾期原因', ''))
                match = re.search(r'(?:业务|责任|负责|联系)(?:人|人员|员)?(?:[:：，,])?([\u4e00-\u9fa5]{2,3})(?:[\d\s,。、，.!?！？]|$)', reason_detail)
                person_val = match.group(1) if match else ''

                build_cell_text(cells[0], i+1, is_appendix=True)
                build_cell_text(cells[1], dept_clean, align='center', is_appendix=True)
                build_cell_text(cells[2], row.get('客户名称', ''), align='center', is_appendix=True)
                dt = row.get('交货结束日期', '')
                build_cell_text(cells[3], str(dt)[:10] if pd.notna(dt) else "", is_appendix=True)
                build_cell_text(cells[4], format_num(row.get('合同数量(万吨)', ''), 2), is_appendix=True)
                build_cell_text(cells[5], format_num(row.get('合同单价', ''), 0, True), is_appendix=True)
                build_cell_text(cells[6], format_num(row.get('逾期天数', ''), 0, True), is_appendix=True)
                build_cell_text(cells[7], format_qty(row.get('逾期数量（万吨）', '')), is_appendix=True)
                build_cell_text(cells[8], row.get('原因分类1', ''), align='left', is_appendix=True)
                build_cell_text(cells[9], row.get('原因分类2', ''), align='left', is_appendix=True)
                build_cell_text(cells[10], reason_detail, align='left', is_appendix=True)
                sol = str(row.get('解决方案', '')) if pd.notna(row.get('解决方案')) else ""
                sol_dt = str(row.get('预计完成日期', ''))[:10] if pd.notna(row.get('预计完成日期')) else ""
                build_cell_text(cells[11], f"{sol}{sol_dt}", align='left', is_appendix=True)
                build_cell_text(cells[12], person_val, is_appendix=True)
                build_cell_text(cells[13], leader_val, is_appendix=True)
                
                credit_val = "是" if pd.notna(row.get('销售类型')) and str(row.get('销售类型')).strip() == '赊销' else ""
                build_cell_text(cells[14], credit_val, is_appendix=True)
                set_table_row_height(table.rows[-1], Cm(0.6).pt)
        else:
            empty_row = table.add_row().cells
            for cell in empty_row: build_cell_text(cell, '', is_appendix=True)
            set_table_row_height(table.rows[-1], Cm(0.6).pt)
        apply_table_borders(table)

    df_app1 = df[df['逾期天数'] >= 50]
    create_appendix('附表1：逾期60天以上的销售合同情况', df_app1, 5)
    
    if '是否重点关注' in df.columns: df_app2 = df[df['是否重点关注'].astype(str).str.contains('重点关注', na=False)]
    else: df_app2 = pd.DataFrame()
    create_appendix('附表2：其他需要重点关注的销售合同情况（尤其是有潜在风险的）', df_app2, 6)
    
    if '是否严重逾期' in df.columns: df_app3 = df[df['是否严重逾期'].astype(str).str.contains('严重逾期', na=False)]
    else: df_app3 = pd.DataFrame()
    create_appendix('附表3：严重逾期销售合同情况', df_app3, 7)

    word_io = io.BytesIO()
    doc.save(word_io)
    word_io.seek(0)
    return word_io

# ================= 业务分析提醒 =================
def generate_reminders(df_unique):
    yesterday = datetime.datetime.now() - datetime.timedelta(days=1)
    date_str = f"{yesterday.month}月{yesterday.day}日"
    reminders = []
    
    has_region = '大区' in df_unique.columns
    regions = df_unique['大区'].dropna().unique() if has_region else []
    
    if len(regions) > 1 or not has_region:
        tot_cnt = len(df_unique)
        tot_qty = df_unique['逾期数量（万吨）'].sum()
        tot_amt = df_unique['逾期金额（万元）'].sum()
        safe_tot_qty = tot_qty if tot_qty > 0 else 1e-9
        avg_days_val = (df_unique['逾期数量（万吨）'] * df_unique['逾期天数']).sum() / safe_tot_qty if tot_qty > 0 else 0
        avg_days = int(Decimal(str(round(avg_days_val, 6))).quantize(Decimal('1'), rounding=ROUND_HALF_UP))
        
        base_str = f"截至{date_str}，中粮贸易逾期销售提货合同合计{tot_cnt}笔，逾期数量{format_qty(tot_qty)}万吨，逾期金额{format_num(tot_amt, 0, True)}万元，平均逾期{avg_days}天。"
        
        v_stats = df_unique.groupby('品种').agg({'合同编号': 'count', '逾期数量（万吨）': 'sum', '逾期金额（万元）': 'sum'}).sort_values(by='逾期金额（万元）', ascending=False)
        v_parts = []
        safe_tot_amt = tot_amt if tot_amt > 0 else 1e-9
        for v in v_stats.index:
            v_amt = v_stats.loc[v, '逾期金额（万元）']
            v_ratio = format_num(v_amt / safe_tot_amt * 100, is_percent=True)
            v_parts.append(f"{v}{v_stats.loc[v, '合同编号']}笔，逾期数量{format_qty(v_stats.loc[v, '逾期数量（万吨）'])}万吨，逾期金额{format_num(v_amt, 0, True)}万元（{v_ratio}）")
        
        if v_parts: base_str += "其中，" + "；".join(v_parts) + "。"
        
        if has_region:
            base_str += "\n分大区情况如下：\n"
            r_stats = df_unique.groupby('大区').agg({'合同编号': 'count', '逾期数量（万吨）': 'sum', '逾期金额（万元）': 'sum'}).sort_values(by='逾期金额（万元）', ascending=False)
            for i, r in enumerate(r_stats.index, 1):
                base_str += f"{i}、{r}，逾期销售提货合同合计{r_stats.loc[r, '合同编号']}笔，逾期数量{format_qty(r_stats.loc[r, '逾期数量（万吨）'])}万吨，逾期金额{format_num(r_stats.loc[r, '逾期金额（万元）'], 0, True)}万元。\n"
        
        reminders.append({"title": "中粮贸易逾期销售概览", "content": base_str.strip()})

    if has_region:
        for region in regions:
            r_df = df_unique[df_unique['大区'] == region]
            r_cnt = len(r_df)
            r_qty = r_df['逾期数量（万吨）'].sum()
            r_amt = r_df['逾期金额（万元）'].sum()
            r_safe_qty = r_qty if r_qty > 0 else 1e-9
            r_avg_days_val = (r_df['逾期数量（万吨）'] * r_df['逾期天数']).sum() / r_safe_qty if r_qty > 0 else 0
            r_avg_days = int(Decimal(str(round(r_avg_days_val, 6))).quantize(Decimal('1'), rounding=ROUND_HALF_UP))
            
            r_content = f"截至{date_str}，{region}逾期销售提货合同合计{r_cnt}笔，逾期数量{format_qty(r_qty)}万吨，逾期金额{format_num(r_amt, 0, True)}万元，平均逾期{r_avg_days}天。\n"
            
            r_content += "分品种情况如下：\n"
            rv_stats = r_df.groupby('品种').agg({'合同编号': 'count', '逾期数量（万吨）': 'sum', '逾期金额（万元）': 'sum'}).sort_values(by='逾期金额（万元）', ascending=False)
            r_safe_amt = r_amt if r_amt > 0 else 1e-9
            for i, v in enumerate(rv_stats.index, 1):
                v_amt = rv_stats.loc[v, '逾期金额（万元）']
                v_ratio = format_num(v_amt / r_safe_amt * 100, is_percent=True)
                r_content += f"{i}、{v}{rv_stats.loc[v, '合同编号']}笔，逾期数量{format_qty(rv_stats.loc[v, '逾期数量（万吨）'])}万吨，逾期金额{format_num(v_amt, 0, True)}万元（{v_ratio}）。\n"

            if '经营部' in r_df.columns:
                r_content += "\n分经营部情况如下：\n"
                d_stats = r_df.groupby('经营部').agg({'合同编号': 'count', '逾期数量（万吨）': 'sum', '逾期金额（万元）': 'sum'}).sort_values(by='逾期金额（万元）', ascending=False)
                for i, d in enumerate(d_stats.index, 1):
                    r_content += f"{i}、{d}，逾期销售提货合同合计{d_stats.loc[d, '合同编号']}笔，逾期数量{format_qty(d_stats.loc[d, '逾期数量（万吨）'])}万吨，逾期金额{format_num(d_stats.loc[d, '逾期金额（万元）'], 0, True)}万元。\n"
                
                has_focus = '是否重点关注' in r_df.columns
                has_severe = '是否严重逾期' in r_df.columns
                
                def get_label(row):
                    if pd.to_numeric(row.get('逾期天数', 0), errors='coerce') >= 60: return "逾期60天以上"
                    if has_severe and '严重逾期' in str(row.get('是否严重逾期', '')): return "严重逾期"
                    if has_focus and '重点关注' in str(row.get('是否重点关注', '')): return "重点关注"
                    return ""
                
                r_df_labeled = r_df.copy()
                r_df_labeled['特殊标签'] = r_df_labeled.apply(get_label, axis=1)
                spec_df = r_df_labeled[r_df_labeled['特殊标签'] != ""]
                
                if not spec_df.empty:
                    r_content += "\n重点关注/严重逾期客户情况如下：\n"
                    spec_df = spec_df.sort_values(by='逾期数量（万吨）', ascending=False)
                    for _, row in spec_df.iterrows():
                        c_name = row.get('客户名称', '')
                        l_tag = row['特殊标签']
                        s_qty = format_qty(row.get('逾期数量（万吨）', 0))
                        s_amt = format_num(row.get('逾期金额（万元）', 0), 0, True)
                        s_days = format_num(row.get('逾期天数', 0), 0, True)
                        r_content += f"{c_name}，{l_tag}，逾期数量{s_qty}万吨，逾期金额{s_amt}万元，逾期{s_days}天。\n"

            reminders.append({"title": f"{region}逾期催收提醒", "content": r_content.strip()})

    return reminders

# ================= 报表生成美化 =================
def beautify_excel_io(df_output):
    output = io.BytesIO()
    df_output.to_excel(output, index=False)
    output.seek(0)
    
    wb = openpyxl.load_workbook(output)
    ws = wb.active
    
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    header_font = Font(name='微软雅黑', bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    body_font = Font(name='微软雅黑', size=10)
    fill_batch = PatternFill(start_color=COLOR_BATCH, end_color=COLOR_BATCH, fill_type="solid")
    fill_once = PatternFill(start_color=COLOR_ONCE, end_color=COLOR_ONCE, fill_type="solid")

    source_col_idx = next((cell.column for cell in ws[1] if cell.value == "_Data_Source"), None)

    for row in ws.iter_rows():
        current_fill = None
        if source_col_idx and row[0].row > 1:
            source_val = row[source_col_idx - 1].value
            if source_val == 'batch': current_fill = fill_batch
            elif source_val == 'once': current_fill = fill_once
        
        ws.row_dimensions[row[0].row].height = 20
        for cell in row:
            cell.border = thin_border
            cell.font = body_font
            cell.alignment = Alignment(vertical='center', wrap_text=False)
            if cell.row > 1 and current_fill: cell.fill = current_fill
            if cell.row == 1:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center', vertical='center')

    if source_col_idx: ws.delete_cols(source_col_idx)

    wide_cols = ["合同编号", "客户名称", "具体逾期原因", "最新进展", "解决方案", "集团内部客户"]
    from openpyxl.utils import get_column_letter
    for col in ws.columns:
        try:
            header_val = str(col[0].value).strip()
            width = 14
            if any(k in header_val for k in wide_cols):
                width = 30
                for cell in col:
                    if cell.row > 1: cell.alignment = Alignment(vertical='center', wrap_text=True)
            elif "日期" in header_val or "品种" in header_val: width = 12
            elif "重点关注" in header_val: width = 15
            ws.column_dimensions[col[0].column_letter].width = width
        except Exception: pass

    final_output = io.BytesIO()
    wb.save(final_output)
    final_output.seek(0)
    return final_output

# ================= 主控制逻辑 =================
def process_overdue_data(batch_files, once_files, mapping_file=None, generate_word=False):
    logs = []
    
    header_keywords = ["大区", "经营部", "合同编号", "客户名称"]
    date_columns = ["合同签订日期", "交货开始日期", "交货结束日期", "预计完成日期"]
    all_numeric_columns = ["合同数量", "合同单价", "合同金额", "调整后逾期销售金额", "逾期天数I", "逾期天数Il", "逾期天数II", "逾期天数IV", "逾期天数V", "逾期天数VI"]
    special_int_columns = ["逾期天数I", "逾期天数Il", "逾期天数II", "逾期天数IV", "逾期天数V", "逾期天数VI"]

    # --- 1. 读取并标记来源 ---
    df_batch_list = []
    for f in batch_files:
        temp = locate_header_and_read_stream(f, header_keywords)
        if temp is not None: df_batch_list.append(temp)
    
    df_batch = pd.DataFrame()
    if df_batch_list:
        temp_combined = pd.concat(df_batch_list, ignore_index=True).drop_duplicates()
        df_batch = process_basic_columns(temp_combined, date_columns, all_numeric_columns, special_int_columns)
        calc_cols = [c for c in special_int_columns if c in df_batch.columns]
        df_batch['逾期天数'] = df_batch[calc_cols].max(axis=1).fillna(0) if calc_cols else 0
        df_batch['_Data_Source'] = 'batch'

    df_once_list = []
    for f in once_files:
        temp = locate_header_and_read_stream(f, header_keywords)
        if temp is not None: df_once_list.append(temp)
    
    df_once = pd.DataFrame()
    if df_once_list:
        temp_combined = pd.concat(df_once_list, ignore_index=True).drop_duplicates()
        df_once = process_basic_columns(temp_combined, date_columns, all_numeric_columns)
        if '逾期天数' not in df_once.columns: df_once['逾期天数'] = 0
        df_once['_Data_Source'] = 'once'

    if df_batch.empty and df_once.empty:
        return None, None, [], ["❌ 错误：未读取到任何有效数据，请检查上传文件格式。"]

    # --- 2. 合并与清洗 ---
    df_merged = pd.concat([df_batch, df_once], ignore_index=True)
    
    if '明细品种' in df_merged.columns and '细分品种' in df_merged.columns:
        df_merged['明细品种'] = df_merged['明细品种'].fillna(df_merged['细分品种'])
        df_merged['明细品种'] = df_merged['明细品种'].astype(str).replace('nan', '')
    elif '细分品种' in df_merged.columns:
        df_merged.rename(columns={'细分品种': '明细品种'}, inplace=True)
        
    if '品种' in df_merged.columns:
        if '明细品种' not in df_merged.columns: df_merged['明细品种'] = ""
        df_merged['品种'] = df_merged['品种'].astype(str).str.strip()
        rice_condition = df_merged['明细品种'].str.contains('稻谷|中晚籼', regex=True, na=False)
        df_merged.loc[rice_condition, '品种'] = '稻谷'
        not_in_whitelist = ~df_merged['品种'].isin(['大豆', '稻谷', '小麦'])
        mask = not_in_whitelist & (df_merged['明细品种'] != '')
        df_merged.loc[mask, '品种'] = df_merged.loc[mask, '明细品种']

    target_col = "逾期分类（业绩考核角度）"
    if target_col in df_merged.columns:
        s = df_merged[target_col].astype(str)
        cond = s.str.contains("A", na=False) & s.str.contains("超过交货结束日期", na=False) & s.str.contains("未完成交提货的", na=False)
        df_final = df_merged[cond].copy()
    else:
        df_final = df_merged.copy()

    # --- 3. 强力匹配逻辑 (支持备用上传 + 全量去空格 + 模糊匹配) ---
    group_map, internal_map = {}, {}
    mapping_source = None
    
    if mapping_file is not None:
        mapping_source = mapping_file
        logs.append("✅ 已使用您在网页【手动上传】的《客户关系清单》。")
    else:
        # 全云端仓库递归搜索
        possible_files = glob.glob("**/*清单*.xlsx", recursive=True)
        if possible_files:
            mapping_source = possible_files[0]
            logs.append(f"✅ 已自动找到 Github 仓库中的清单：{mapping_source}")
        else:
            logs.append("⚠️ 未在云端或上传框找到“清单”，如果需要匹配集团，请尝试【手动上传】！")

    if mapping_source:
        try:
            xl = pd.ExcelFile(mapping_source)
            sheet_names = xl.sheet_names
            
            # 【总】表处理
            total_sheet = next((s for s in sheet_names if '总' in s), None)
            if total_sheet:
                df_t = pd.read_excel(mapping_source, sheet_name=total_sheet)
                col_c = next((c for c in df_t.columns if '客户' in str(c) and '名称' in str(c)), None)
                col_g = next((c for c in df_t.columns if '集团' in str(c)), None)
                if col_c and col_g:
                    for _, row in df_t.iterrows():
                        k = clean_text_for_match(row[col_c])
                        v = str(row[col_g]).strip() if pd.notna(row[col_g]) else ""
                        if k and v and v != 'nan': group_map[k] = v

            # 【内部】表处理
            internal_sheet = next((s for s in sheet_names if '内部' in s), None)
            if internal_sheet:
                df_i = pd.read_excel(mapping_source, sheet_name=internal_sheet)
                col_ci = next((c for c in df_i.columns if '客户' in str(c) and '名称' in str(c)), None)
                col_p = next((c for c in df_i.columns if '公司' in str(c) or '专业' in str(c)), None)
                if col_ci and col_p:
                    for _, row in df_i.iterrows():
                        k = clean_text_for_match(row[col_ci])
                        v = str(row[col_p]).strip() if pd.notna(row[col_p]) else ""
                        if k and v and v != 'nan': internal_map[k] = v
        except Exception as e:
            logs.append(f"⚠️ 关系清单解析失败: {e}")

    for col in ['调整后逾期销售金额', '合同单价', '合同金额', '交货结束日期', '交货开始日期', '合同数量']:
        if col not in df_final.columns:
            df_final[col] = pd.NaT if '日期' in col else 0

    if '客户名称' in df_final.columns:
        def get_group(cust_name):
            c_clean = clean_text_for_match(cust_name)
            if not c_clean: return ""
            if c_clean in group_map: return group_map[c_clean] # 精准匹配
            for k, v in group_map.items(): # 模糊包含匹配
                if k in c_clean or c_clean in k: return v
            return ""

        def get_internal(cust_name, group_name):
            if '中粮' not in str(group_name): return "" 
            c_clean = clean_text_for_match(cust_name)
            if not c_clean: return ""
            if c_clean in internal_map: return internal_map[c_clean]
            for k, v in internal_map.items():
                if k in c_clean or c_clean in k: return v
            return ""

        df_final['所属集团'] = df_final['客户名称'].apply(get_group)
        df_final['集团内部客户'] = df_final.apply(lambda row: get_internal(row['客户名称'], row['所属集团']), axis=1)
    else:
        df_final['所属集团'] = ""
        df_final['集团内部客户'] = ""

    bins = [-float('inf'), 10, 20, 30, 60, 90, float('inf')]
    labels = ["1-10天", "11-20天", "21-30天", "31-60天", "61-90天", "90天以上"]
    if '逾期天数' in df_final.columns:
        df_final['逾期天数分类'] = pd.cut(df_final['逾期天数'], bins=bins, labels=labels)

    df_final['合同单价_safe'] = df_final['合同单价'].replace(0, np.nan)
    df_final['逾期数量'] = df_final['调整后逾期销售金额'] / df_final['合同单价_safe'] / 10000
    df_final['逾期数量'] = df_final['逾期数量'].fillna(0)

    df_final['合同执行期(天数)'] = (df_final['交货结束日期'] - df_final['交货开始日期']).dt.days
    df_final['合同执行期(天数)'] = df_final['合同执行期(天数)'].fillna(1).replace(0, 1)
    
    ratio_days = (df_final['逾期天数'] / df_final['合同执行期(天数)']).fillna(0)
    ratio_amt = (df_final['调整后逾期销售金额'] / df_final['合同金额'].replace(0, np.nan)).fillna(0)
    df_final['是否严重逾期'] = np.where((ratio_days > 0.5) & (ratio_amt > 0.5), "严重逾期", "")

    # --- 4. 换算与排序 ---
    df_final['合同数量'] = (df_final['合同数量'] / 10000).round(4)
    df_final['合同金额'] = (df_final['合同金额'] / 10000).round(2)
    df_final['调整后逾期销售金额'] = (df_final['调整后逾期销售金额'] / 10000).round(2)
    df_final['逾期数量'] = df_final['逾期数量'].round(4)

    cond_focus = (df_final['调整后逾期销售金额'] >= 500) & (df_final['逾期天数'] >= 10)
    df_final['是否重点关注'] = np.where(cond_focus, "重点关注", "")
    if '销售类型' not in df_final.columns: df_final['销售类型'] = ""

    df_final.sort_values(
        by=['逾期数量', '逾期天数', '是否严重逾期', '是否重点关注'],
        ascending=[False, False, False, False],
        inplace=True
    )

    col_rename_map = {
        "合同签订日期": "签订日期", "合同数量": "合同数量(万吨)", "合同金额": "合同金额（万元）", 
        "逾期数量": "逾期数量（万吨）", "调整后逾期销售金额": "逾期金额（万元）",
        "逾期原因分类1（责任划分角度）": "原因分类1", "逾期原因分类2（责任划分角度）": "原因分类2",
        "当日最新进展": "最新进展"
    }
    df_final.rename(columns=col_rename_map, inplace=True)

    for col in ["签订日期", "交货结束日期", "预计完成日期"]:
        if col in df_final.columns: df_final[col] = df_final[col].dt.strftime('%Y-%m-%d')

    final_columns = [
        "大区", "经营部", "合同编号", "客户名称", "签订日期", "交货结束日期", "品种", 
        "合同数量(万吨)", "合同单价", "合同金额（万元）", "逾期天数", "逾期数量（万吨）", 
        "逾期金额（万元）", "原因分类1", "原因分类2", "具体逾期原因", "预计完成日期", 
        "解决方案", "最新进展", "是否严重逾期", "逾期天数分类", "所属集团", 
        "集团内部客户", "是否重点关注", "销售类型"
    ]
    for col in final_columns:
        if col not in df_final.columns: df_final[col] = ""

    df_output = df_final[final_columns + ["_Data_Source"]]
    excel_io = beautify_excel_io(df_output)
    
    df_unique = df_output.drop_duplicates(subset=['合同编号']).copy()
    
    # --- 5. 提醒文本生成 ---
    reminders = generate_reminders(df_unique)
    
    # --- 6. 生成 Word 报告 ---
    word_io = None
    if generate_word:
        word_io = generate_word_report(df_final, df_unique)
    
    logs.append("🎉 数据处理与计算完成！")
    return excel_io, word_io, reminders, logs



