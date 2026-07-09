import pandas as pd
import io
import numpy as np
from datetime import datetime, timedelta
import warnings
import openpyxl
from openpyxl.styles import Alignment, Border, Side, Font, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

warnings.filterwarnings('ignore')

HEADER_COLOR = "4472C4"
SHEET_NAME = "逾期采购监控表-日报"
HEADER_KEY_COLS = ["大区", "区域公司", "合同编号", "供应商名称", "品种", "逾期天数"]

DATE_COLS = ["合同签订日期", "合同收货结束日期", "合同终止日期"]
FLOAT_COLS = ["合同数量", "合同单价", "合同金额", "合同履约率", "已入库数量",
              "已入库金额", "已付款金额", "逾期采购金额", "逾期采购金额调整值", "当前市场价格", "已收保证金"]
INT_COLS = ["逾期天数"]

AMOUNT_COLUMNS_TO_CONVERT = ["合同金额", "已入库金额", "已付款金额", "逾期金额确定", "市场价格金额", "盈亏金额", "已收保证金"]
TON_COLS = ["合同数量", "已入库数量", "逾期数量"]
YUAN_COLS = ["合同单价", "当前市场价格", "盯市盈亏"]

FINAL_OUTPUT_COLS = [
    "大区", "区域公司", "合同编号", "供应商名称", "品种",
    "合同签订日期", "合同收货结束日期", "合同终止日期",
    "合同数量（吨）", "合同单价（元）", "合同金额（万元）", "是否预付款", "合同履约率",
    "已入库数量（吨）", "已入库金额（万元）", "已付款金额（万元）", "逾期天数",
    "逾期数量（吨）", "逾期采购金额（万元）",
    "当前市场价格（元）", "盯市盈亏（元）", "市场价格金额（万元）", "盈亏金额（万元）",
    "已收保证金（万元）", "逾期原因"
]

def locate_header_and_read_io(file_stream, sheet_name, key_columns):
    try:
        file_stream.seek(0)
        df_raw = pd.read_excel(file_stream, sheet_name=sheet_name, header=None)
        header_row_idx = -1
        for i, row in df_raw.iterrows():
            row_values = [str(x).strip().replace('\n', '').replace(' ', '') for x in row.values if pd.notna(x)]
            if sum(1 for key in key_columns if key in row_values) >= len(key_columns) * 0.8:
                header_row_idx = i
                break
        if header_row_idx == -1:
            return None
        file_stream.seek(0)
        df = pd.read_excel(file_stream, sheet_name=sheet_name, header=header_row_idx)
        df.columns = df.columns.astype(str).str.replace('\n', '', regex=False).str.strip()
        df.dropna(how='all', inplace=True)
        return df.reset_index(drop=True)
    except Exception as e:
        return None

def process_basic_columns(df, date_cols, float_cols, int_cols):
    for col in date_cols:
        if col in df.columns:
            temp_dates = pd.to_datetime(df[col], errors='coerce')
            invalid_mask = (temp_dates.dt.year < 2025) | (temp_dates.dt.year > 2030) | temp_dates.isna()
            if invalid_mask.any():
                clean_vals = df.loc[invalid_mask, col].astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
                temp_dates.loc[invalid_mask] = pd.to_datetime(clean_vals, format='%Y%m%d', errors='coerce')
            df[col] = temp_dates
    for col in float_cols:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
    for col in int_cols:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0).astype('Int64')
    return df

def beautify_excel_purchase(wb):
    try:
        ws = wb.active
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        header_font = Font(name='微软雅黑', bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color=HEADER_COLOR, end_color=HEADER_COLOR, fill_type="solid")
        body_font = Font(name='微软雅黑', size=10)
        align_center = Alignment(horizontal='center', vertical='center', wrap_text=False)
        align_left_wrap = Alignment(horizontal='left', vertical='center', wrap_text=True)
        format_col_mapping = {}
        for col_idx, cell in enumerate(next(ws.iter_rows(min_row=1, max_row=1)), 1):
            header_val = str(cell.value).strip()
            if "（万元）" in header_val:
                format_col_mapping[col_idx] = '0'
            elif "履约率" in header_val:
                format_col_mapping[col_idx] = '0.00%'
        for row in ws.iter_rows():
            ws.row_dimensions[row[0].row].height = 20
            for cell in row:
                cell.border = thin_border
                cell.font = body_font
                cell.alignment = align_center
                if cell.row == 1:
                    cell.font = header_font
                    cell.fill = header_fill
                elif cell.column in format_col_mapping and isinstance(cell.value, (int, float)):
                    cell.number_format = format_col_mapping[cell.column]
        for col in ws.columns:
            col_letter = get_column_letter(col[0].column)
            header_val = str(col[0].value).strip()
            if any(key in header_val for key in ["合同编号", "供应商名称", "逾期原因"]):
                ws.column_dimensions[col_letter].width = 32
                for cell in col[1:]:
                    cell.alignment = align_left_wrap
            elif any(key in header_val for key in ["品种", "大区", "区域公司"]):
                ws.column_dimensions[col_letter].width = 14
            elif "日期" in header_val:
                ws.column_dimensions[col_letter].width = 12
            elif any(k in header_val for k in ["（万元）", "（元）", "（吨）", "盈亏"]):
                ws.column_dimensions[col_letter].width = 16
            else:
                ws.column_dimensions[col_letter].width = 10
        ws.freeze_panes = "A2"
    except Exception as e:
        pass

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_font_mixed(style, size_pt, is_bold, east_asia_name, ascii_name):
    style.font.size = Pt(size_pt)
    style.font.bold = is_bold
    style.font.name = ascii_name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_name)

def format_qty(val):
    if val == 0:
        return "0.00"
    if abs(val) < 0.01:
        val_str = f"{abs(val):.15f}".split('.')[1]
        for i, digit in enumerate(val_str):
            if digit != '0':
                fmt = f"{{:.{i+1}f}}"
                return fmt.format(val)
    return f"{val:.2f}"

def get_circled_number(n):
    if 1 <= n <= 20:
        return chr(0x245F + n)
    elif 21 <= n <= 35:
        return chr(0x3251 + n - 21)
    elif 36 <= n <= 50:
        return chr(0x32B1 + n - 36)
    else:
        return f"({n})"

def set_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)
    for cs in trPr.xpath('w:cantSplit'):
        trPr.remove(cs)

def apply_cell_font(cell, is_bold=False, font_color=None, cn_size=10.5, en_size=10):
    full_text = "".join(r.text for p in cell.paragraphs for r in p.runs).replace(" ", "").replace("\u3000", "").replace("\n", "").replace("\r", "").strip()
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = Pt(10)
    para.paragraph_format.first_line_indent = 0
    
    parts = re.split(r'([^\x00-\xff]+)', full_text)
    for part in parts:
        if not part: continue
        run = para.add_run(part)
        run.font.bold = is_bold
        if font_color:
            run.font.color.rgb = font_color
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        if re.match(r'[^\x00-\xff]', part):
            run.font.size = Pt(cn_size) 
        else:
            run.font.size = Pt(en_size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def generate_weekly_report(df):
    web_text_lines = []
    try:
        if df.empty:
            return None, ""

        df = df.fillna("")
        df["逾期数量（吨）"] = pd.to_numeric(df["逾期数量（吨）"], errors="coerce").fillna(0)
        df["逾期采购金额（万元）"] = pd.to_numeric(df["逾期采购金额（万元）"], errors="coerce").fillna(0)
        df["合同单价（元）"] = pd.to_numeric(df["合同单价（元）"], errors="coerce").fillna(0)
        df["当前市场价格（元）"] = pd.to_numeric(df["当前市场价格（元）"], errors="coerce").fillna(0)
        df["盈亏金额（万元）"] = pd.to_numeric(df["盈亏金额（万元）"], errors="coerce").fillna(0)
        df["逾期天数"] = pd.to_numeric(df["逾期天数"], errors="coerce").fillna(0).astype(int)
        
        if "已收保证金（万元）" not in df.columns:
            df["已收保证金（万元）"] = 0
        df["已收保证金（万元）"] = pd.to_numeric(df["已收保证金（万元）"], errors="coerce").fillna(0)

        report_dept = df["大区"].iloc[0] if not df.empty and "大区" in df.columns else "沿海大区"

        today = datetime.now()
        offset = (today.weekday() - 2) % 7 
        last_wed = today - timedelta(days=offset)
        report_date = f"{last_wed.month}月{last_wed.day}日"

        doc = Document()
        
        section = doc.sections[0]
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.orientation = WD_ORIENT.PORTRAIT

        style_tb = doc.styles.add_style('TableTitle', 1)
        set_font_mixed(style_tb, 12.0, False, '微软雅黑', 'Times New Roman')
        style_tb.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style_tb.paragraph_format.line_spacing = Pt(22)
        style_tb.paragraph_format.space_before = Pt(0)
        style_tb.paragraph_format.space_after = Pt(0)
        style_tb.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_tb.paragraph_format.first_line_indent = Pt(0)

        style_nm = doc.styles.add_style('NormalContent', 1)
        set_font_mixed(style_nm, 14.0, False, '仿宋_GB2312', 'Times New Roman')
        style_nm.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style_nm.paragraph_format.line_spacing = Pt(22)
        style_nm.paragraph_format.space_before = Pt(0)
        style_nm.paragraph_format.space_after = Pt(0)
        style_nm.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style_nm.paragraph_format.first_line_indent = Pt(28)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title.paragraph_format.space_before = Pt(0)
        title.paragraph_format.space_after = Pt(0)
        title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        title.paragraph_format.line_spacing = Pt(22)
        title.paragraph_format.first_line_indent = 0
        run = title.add_run("一、逾期采购交货")
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        run.font.size = Pt(14)
        run.font.bold = True

        total_contracts = len(df)
        total_overdue_ton = df["逾期数量（吨）"].sum()
        total_overdue_amount = df["逾期采购金额（万元）"].sum()
        qty_str = format_qty(total_overdue_ton / 10000)

        opening_para1 = doc.add_paragraph(style='NormalContent')
        opening_para1.add_run(f"截至{report_date}，{report_dept}逾期采购交货情况如下：")
        
        opening_para2 = doc.add_paragraph(style='NormalContent')
        opening_para2.add_run(f"{report_dept}本期逾期采购合同共计")
        r1 = opening_para2.add_run(f"{total_contracts}笔")
        r1.bold = True
        r1.font.color.rgb = RGBColor(255, 0, 0)
        opening_para2.add_run("，逾期数量")
        r2 = opening_para2.add_run(f"{qty_str}万吨")
        r2.bold = True
        r2.font.color.rgb = RGBColor(255, 0, 0)
        opening_para2.add_run("，逾期金额")
        r3 = opening_para2.add_run(f"{total_overdue_amount:.0f}万元")
        r3.bold = True
        r3.font.color.rgb = RGBColor(255, 0, 0)
        opening_para2.add_run("。")

        # 同步写入网页文本
        web_text_lines.append(f"截至{report_date}，{report_dept}逾期采购交货情况如下：")
        web_text_lines.append(f"{report_dept}本期逾期采购合同共计{total_contracts}笔，逾期数量{qty_str}万吨，逾期金额{total_overdue_amount:.0f}万元。")

        # ------------------- 业务部门分析 (多于1个才生成) -------------------
        dept_group = df.groupby("区域公司").agg(
            逾期金额=("逾期采购金额（万元）", "sum"),
            合同笔数=("合同编号", "count"),
            逾期数量=("逾期数量（吨）", "sum")
        ).reset_index()
        
        if len(dept_group) > 1:
            dept_group["逾期占总比"] = (dept_group["逾期金额"] / total_overdue_amount * 100).fillna(0).round(0).astype(int)
            dept_group = dept_group.sort_values("逾期金额", ascending=False).reset_index(drop=True)

            dept_para = doc.add_paragraph(style='NormalContent')
            dept_para.add_run("分业务部门看，")
            
            dept_web_parts = []
            for idx, row in dept_group.iterrows():
                text = f"{row['区域公司']}经营部逾期金额{row['逾期金额']:.0f}万元、占{row['逾期占总比']}%"
                dept_web_parts.append(text)
                run_dept = dept_para.add_run(text)
                
                if idx == 0:
                    run_dept.bold = True
                    run_dept.font.color.rgb = RGBColor(255, 0, 0)
                    
                if idx < len(dept_group) - 1:
                    dept_para.add_run("；")
                else:
                    dept_para.add_run("。")
            
            web_text_lines.append("分业务部门看，" + "；".join(dept_web_parts) + "。")

            doc.add_paragraph("逾期采购分业务单位情况表", style='TableTitle')

            dept_table = doc.add_table(rows=len(dept_group)+2, cols=5)
            dept_table.style = 'Table Grid'
            dept_table.autofit = False
            dept_table.allow_autofit = False
            dept_widths = [2.72, 2.95, 2.95, 2.95, 2.95]
            for i, col in enumerate(dept_table.columns):
                col.width = Cm(dept_widths[i])
                
            set_header_row(dept_table.rows[0])
            dept_header = ["业务部门", "逾期金额(万元)", "逾期占总比", "合同个数（笔）", "逾期数量（吨）"]
            for col_idx, header_text in enumerate(dept_header):
                dept_table.rows[0].cells[col_idx].text = header_text

            for row_idx, row in dept_group.iterrows():
                table_row = dept_table.rows[row_idx+1]
                table_row.cells[0].text = f"{row['区域公司']}经营部"
                table_row.cells[1].text = f"{row['逾期金额']:.0f}"
                table_row.cells[2].text = f"{row['逾期占总比']}%"
                table_row.cells[3].text = f"{row['合同笔数']}"
                table_row.cells[4].text = f"{row['逾期数量']:.0f}"

            total_row = dept_table.rows[len(dept_group)+1]
            total_row.cells[0].text = "总计"
            total_row.cells[1].text = f"{total_overdue_amount:.0f}"
            total_row.cells[2].text = "100%"
            total_row.cells[3].text = f"{total_contracts}"
            total_row.cells[4].text = f"{total_overdue_ton:.0f}"

            dept_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, row in enumerate(dept_table.rows):
                is_header = (i == 0)
                row.height = Cm(1.0) if is_header else Cm(0.61)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                is_total = (i == len(dept_table.rows) - 1)
                is_max_row = (i == 1) 
                
                for j, cell in enumerate(row.cells):
                    if j < len(dept_widths):
                        cell.width = Cm(dept_widths[j])
                    cell_font_color = RGBColor(255, 0, 0) if is_max_row else None
                    cell_is_bold = True if (is_header or is_total or is_max_row) else False
                    apply_cell_font(cell, is_bold=cell_is_bold, font_color=cell_font_color)
                    if is_header:
                        set_cell_background(cell, 'D9D9D9')
                    elif is_total:
                        set_cell_background(cell, 'DEEBF6')

        # ------------------- 品种分析 -------------------
        variety_group = df.groupby("品种").agg(
            逾期金额=("逾期采购金额（万元）", "sum"),
            合同笔数=("合同编号", "count"),
            逾期数量=("逾期数量（吨）", "sum")
        ).reset_index()
        variety_group["逾期占总比"] = (variety_group["逾期金额"] / total_overdue_amount * 100).fillna(0).round(0).astype(int)
        variety_group = variety_group.sort_values("逾期金额", ascending=False).reset_index(drop=True)

        variety_para = doc.add_paragraph(style='NormalContent')
        variety_para.add_run("分品种看，")
        
        var_web_parts = []
        for idx, row in variety_group.iterrows():
            text = f"{row['品种']}逾期金额{row['逾期金额']:.0f}万元、占{row['逾期占总比']}%"
            var_web_parts.append(text)
            run_var = variety_para.add_run(text)
            
            if idx == 0:
                run_var.bold = True
                run_var.font.color.rgb = RGBColor(255, 0, 0)
                
            if idx < len(variety_group) - 1:
                variety_para.add_run("；")
            else:
                variety_para.add_run("。")
                
        web_text_lines.append("分品种看，" + "；".join(var_web_parts) + "。")

        doc.add_paragraph("逾期采购分品种情况表", style='TableTitle')

        variety_table = doc.add_table(rows=len(variety_group)+2, cols=5)
        variety_table.style = 'Table Grid'
        variety_table.autofit = False
        variety_table.allow_autofit = False
        variety_widths = [2.72, 2.95, 2.95, 2.95, 2.95]
        for i, col in enumerate(variety_table.columns):
            col.width = Cm(variety_widths[i])
            
        set_header_row(variety_table.rows[0])
        variety_header = ["逾期品种", "逾期金额(万元)", "逾期占总比", "合同个数（笔）", "逾期数量（吨）"]
        for col_idx, header_text in enumerate(variety_header):
            variety_table.rows[0].cells[col_idx].text = header_text

        for row_idx, row in variety_group.iterrows():
            table_row = variety_table.rows[row_idx+1]
            table_row.cells[0].text = str(row["品种"])
            table_row.cells[1].text = f"{row['逾期金额']:.0f}"
            table_row.cells[2].text = f"{row['逾期占总比']}%"
            table_row.cells[3].text = f"{row['合同笔数']}"
            table_row.cells[4].text = f"{row['逾期数量']:.0f}"

        total_row = variety_table.rows[len(variety_group)+1]
        total_row.cells[0].text = "总计"
        total_row.cells[1].text = f"{total_overdue_amount:.0f}"
        total_row.cells[2].text = "100%"
        total_row.cells[3].text = f"{total_contracts}"
        total_row.cells[4].text = f"{total_overdue_ton:.0f}"

        variety_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, row in enumerate(variety_table.rows):
            is_header = (i == 0)
            row.height = Cm(1.0) if is_header else Cm(0.61)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            is_total = (i == len(variety_table.rows) - 1)
            is_max_row = (i == 1) 
            
            for j, cell in enumerate(row.cells):
                if j < len(variety_widths):
                    cell.width = Cm(variety_widths[j])
                cell_font_color = RGBColor(255, 0, 0) if is_max_row else None
                cell_is_bold = True if (is_header or is_total or is_max_row) else False
                apply_cell_font(cell, is_bold=cell_is_bold, font_color=cell_font_color)
                if is_header:
                    set_cell_background(cell, 'D9D9D9')
                elif is_total:
                    set_cell_background(cell, 'DEEBF6')

        # ------------------- 供应商分析 -------------------
        doc.add_paragraph("分采购供应商来看：", style='NormalContent')
        web_text_lines.append("分采购供应商来看：")
        
        supplier_group = df.groupby("供应商名称")
        supplier_summary_list = []

        for supplier_name, supplier_df in supplier_group:
            supplier_df = supplier_df.reset_index(drop=True)
            contract_count = len(supplier_df)
            variety_list = supplier_df["品种"].unique().tolist()
            max_overdue_days = supplier_df["逾期天数"].max()
            total_supplier_ton = supplier_df["逾期数量（吨）"].sum()
            total_supplier_amount = supplier_df["逾期采购金额（万元）"].sum()
            total_supplier_margin = supplier_df["已收保证金（万元）"].sum()
            
            total_profit_loss = supplier_df["盈亏金额（万元）"].sum()
            supplier_df['has_risk'] = supplier_df['盈亏金额（万元）'] > supplier_df['已收保证金（万元）']

            if total_supplier_ton > 0:
                price_diff = total_profit_loss * 10000 / total_supplier_ton
            else:
                price_diff = 0
                
            market_risk = "有" if total_supplier_margin < total_profit_loss else ""

            total_contract_ton = supplier_df["合同数量（吨）"].sum()
            if total_contract_ton > 0:
                weighted_price = (supplier_df["合同单价（元）"] * supplier_df["合同数量（吨）"]).sum() / total_contract_ton
            else:
                weighted_price = 0

            supplier_summary_list.append({
                "经营部": supplier_df["区域公司"].iloc[0],
                "供应商名称": supplier_name,
                "品种": "、".join(variety_list),
                "合同笔数": contract_count,
                "合同数量": supplier_df["合同数量（吨）"].sum(),
                "待执行数量": total_supplier_ton,
                "逾期数量": total_supplier_ton,
                "逾期金额": total_supplier_amount,
                "最长逾期天数": max_overdue_days,
                "当前市价": supplier_df["当前市场价格（元）"].max(),
                "单吨涨跌幅": price_diff,
                "涨跌金额": total_profit_loss,
                "已收保证金": total_supplier_margin,
                "市场风险": market_risk,
                "合同单价": weighted_price,
                "_df": supplier_df
            })

        supplier_summary_list.sort(key=lambda x: (x["逾期金额"], x["合同笔数"]), reverse=True)

        for idx, item in enumerate(supplier_summary_list):
            supplier_name = item["供应商名称"]
            supplier_df = item["_df"]
            contract_count = item["合同笔数"]
            variety_list = supplier_df["品种"].unique().tolist()
            max_overdue_days = item["最长逾期天数"]
            total_supplier_ton = item["逾期数量"]
            total_supplier_amount = item["逾期金额"]
            
            all_no_risk = not supplier_df['has_risk'].any()

            supplier_para = doc.add_paragraph(style='NormalContent')
            circled_num = get_circled_number(idx + 1)
            run_name = supplier_para.add_run(f"{circled_num}{supplier_name}")
            run_name.bold = True
            
            supplier_qty_str = format_qty(total_supplier_ton / 10000)
            summary_text = f"，{contract_count}笔{'、'.join(variety_list)}采购合同，最长逾期{max_overdue_days}天，逾期数量{supplier_qty_str}万吨，逾期金额{total_supplier_amount:.0f}万元"
            
            # --- 优化点：网页端颜色加深与加粗 ---
            web_supplier_part = f"<span style='color: #333;'><span style='font-weight: bold;'>{supplier_name}</span>，{contract_count}笔{'、'.join(variety_list)}采购合同，最长逾期{max_overdue_days}天，逾期数量{supplier_qty_str}万吨，逾期金额{total_supplier_amount:.0f}万元"
            
            if all_no_risk:
                supplier_para.add_run(summary_text + "，")
                run_margin_txt = supplier_para.add_run("已收保证金能覆盖潜在涨幅损失。")
                run_margin_txt.bold = True
                run_margin_txt.font.color.rgb = RGBColor(255, 0, 0)
                web_supplier_part += "，已收保证金能覆盖潜在涨幅损失。"
            else:
                supplier_para.add_run(summary_text + "。")
                web_supplier_part += "。"
            
            web_supplier_part += "</span>"
            web_summary_text = f"{idx+1}、{web_supplier_part}"
            
            current_web_lines = []

            for c_idx, contract_row in supplier_df.iterrows():
                if contract_count == 1:
                    detail_text = f"{contract_row['逾期原因']}。"
                    if contract_row["当前市场价格（元）"] > 0:
                        detail_text += f"现货价{contract_row['当前市场价格（元）']:.0f}元/吨。"
                    supplier_para.add_run(detail_text)
                    
                    # --- 优化点：单笔合同时，原因直接连在后面不换行 ---
                    web_summary_text += detail_text
                else:
                    contract_para = doc.add_paragraph(style='NormalContent')
                    
                    # --- 优化点：多合同时，将编号改为（1）、（2）... ---
                    contract_para.add_run(f"（{c_idx+1}）")
                    run_cno = contract_para.add_run(f"{contract_row['合同编号']}")
                    run_cno.bold = True
                    # 合同号中的英文/数字部分使用 Times New Roman，中文部分保持仿宋_GB2312
                    run_cno.font.name = 'Times New Roman'
                    run_cno._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
                    run_cno._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    
                    cqty_str = format_qty(contract_row['逾期数量（吨）'] / 10000)
                    camt = contract_row['逾期采购金额（万元）']
                    cdays = contract_row['逾期天数']
                    cvar = contract_row['品种']
                    
                    detail_text = f"：{cvar}采购合同，逾期{cdays}天，逾期数量{cqty_str}万吨，逾期金额{camt:.0f}万元。{contract_row['逾期原因']}。"
                    
                    # 网页端特定排版，加几个空格让子合同有一点点缩进效果
                    web_detail_text = f"&nbsp;&nbsp;&nbsp;&nbsp;（{c_idx+1}）{contract_row['合同编号']}：{cvar}采购合同，逾期{cdays}天，逾期数量{cqty_str}万吨，逾期金额{camt:.0f}万元。{contract_row['逾期原因']}。"
                    
                    if contract_row["当前市场价格（元）"] > 0:
                        detail_text += f"现货价{contract_row['当前市场价格（元）']:.0f}元/吨。"
                        web_detail_text += f"现货价{contract_row['当前市场价格（元）']:.0f}元/吨。"
                    
                    contract_para.add_run(detail_text)
                
                    if not all_no_risk and not contract_row['has_risk']:
                        run_c_margin = contract_para.add_run("已收保证金能覆盖潜在涨幅损失。")
                        run_c_margin.bold = True
                        run_c_margin.font.color.rgb = RGBColor(255, 0, 0)
                        web_detail_text += "已收保证金能覆盖潜在涨幅损失。"
                        
                    current_web_lines.append(web_detail_text)
            
            web_text_lines.append(web_summary_text)
            web_text_lines.extend(current_web_lines)

        doc.add_paragraph("逾期采购供应商情况表", style='TableTitle')

        p_unit = doc.add_paragraph("单位：吨,元/吨,万元")
        p_unit.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_unit.paragraph_format.space_before = Pt(0)
        p_unit.paragraph_format.space_after = Pt(0)
        p_unit.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p_unit.paragraph_format.line_spacing = Pt(12)
        p_unit.paragraph_format.first_line_indent = 0
        for r in p_unit.runs:
            r.font.name = 'Times New Roman'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            r.font.size = Pt(9) 
            r.font.bold = False

        supplier_table = doc.add_table(rows=len(supplier_summary_list)+2, cols=16)
        supplier_table.style = 'Table Grid'
        supplier_table.autofit = False
        supplier_table.allow_autofit = False
        supplier_widths = [0.75, 1, 2.5, 1, 1, 1.3, 1.3, 1.4, 1.1, 1, 1, 1, 1, 1, 1.1, 1.25]
        for i, col in enumerate(supplier_table.columns):
            col.width = Cm(supplier_widths[i])
            
        set_header_row(supplier_table.rows[0])
        supplier_header = [
            "序号", "经营部", "供应商名称", "品种", "合同笔数", "合同数量", "待执行数量",
            "逾期数量", "逾期金额", "最长逾期天数", "合同单价", "当前市价", "单吨涨跌幅",
            "涨跌金额", "已收保证金", "市场风险"
        ]
        
        for col_idx, header_text in enumerate(supplier_header):
            supplier_table.rows[0].cells[col_idx].text = header_text

        for row_idx, item in enumerate(supplier_summary_list):
            table_row = supplier_table.rows[row_idx+1]
            table_row.cells[0].text = str(row_idx+1)
            table_row.cells[1].text = str(item["经营部"])
            table_row.cells[2].text = str(item["供应商名称"])
            table_row.cells[3].text = str(item["品种"])
            table_row.cells[4].text = f"{item['合同笔数']}"
            table_row.cells[5].text = f"{item['合同数量']:.0f}"
            table_row.cells[6].text = f"{item['待执行数量']:.0f}"
            table_row.cells[7].text = f"{item['逾期数量']:.0f}"
            table_row.cells[8].text = f"{item['逾期金额']:.0f}"
            table_row.cells[9].text = f"{item['最长逾期天数']}"
            table_row.cells[10].text = f"{item['合同单价']:.0f}"
            table_row.cells[11].text = f"{item['当前市价']:.0f}"
            table_row.cells[12].text = f"{item['单吨涨跌幅']:.0f}"
            table_row.cells[13].text = f"{item['涨跌金额']:.0f}"
            table_row.cells[14].text = f"{item['已收保证金']:.0f}"
            table_row.cells[15].text = str(item["市场风险"])

        total_supplier_row = supplier_table.rows[len(supplier_summary_list)+1]
        total_supplier_row.cells[0].merge(total_supplier_row.cells[3])
        total_supplier_row.cells[0].text = f"{report_dept}汇总"
        total_supplier_row.cells[4].text = f"{total_contracts}"
        total_supplier_row.cells[5].text = f"{df['合同数量（吨）'].sum():.0f}"
        total_supplier_row.cells[6].text = f"{total_overdue_ton:.0f}"
        total_supplier_row.cells[7].text = f"{total_overdue_ton:.0f}"
        total_supplier_row.cells[8].text = f"{total_overdue_amount:.0f}"
        total_supplier_row.cells[9].text = f"{df['逾期天数'].max()}"

        all_profit_loss = df["盈亏金额（万元）"].sum()
        if total_overdue_ton > 0:
            all_price_diff = all_profit_loss * 10000 / total_overdue_ton
        else:
            all_price_diff = 0

        total_all_ton = df["合同数量（吨）"].sum()
        if total_all_ton > 0:
            all_weighted_price = (df["合同单价（元）"] * df["合同数量（吨）"]).sum() / total_all_ton
            all_weighted_market = (df["当前市场价格（元）"] * df["合同数量（吨）"]).sum() / total_all_ton
        else:
            all_weighted_price = 0
            all_weighted_market = 0

        total_supplier_row.cells[10].text = f"{all_weighted_price:.0f}"
        total_supplier_row.cells[11].text = f"{all_weighted_market:.0f}"
        total_supplier_row.cells[12].text = f"{all_price_diff:.0f}"
        total_supplier_row.cells[13].text = f"{all_profit_loss:.0f}"
        total_supplier_row.cells[14].text = f"{df['已收保证金（万元）'].sum():.0f}"
        total_supplier_row.cells[15].text = ""

        supplier_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, row in enumerate(supplier_table.rows):
            is_header = (i == 0)
            is_total = (i == len(supplier_table.rows) - 1)
            if is_header:
                row.height = Cm(1.11)
            else:
                row.height = Cm(0.452)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            
            seen_cells = set()
            for j, cell in enumerate(row.cells):
                if cell in seen_cells:
                    continue
                seen_cells.add(cell)
                if is_total and j == 0:
                    cell.width = Cm(sum(supplier_widths[0:4]))
                elif j < len(supplier_widths):
                    cell.width = Cm(supplier_widths[j])
                
                apply_cell_font(cell, is_bold=(is_header or is_total), cn_size=7.5, en_size=7.5)
                if is_header:
                    set_cell_background(cell, 'D9D9D9')
                elif is_total or j == 7:
                    set_cell_background(cell, 'DEEBF6')

        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io, "\n".join(web_text_lines)
    except Exception as e:
        return None, f"生成周报失败: {str(e)}"

def process_overdue_purchase(uploaded_files):
    logs = []
    all_dfs = []
    for file in uploaded_files:
        df_temp = locate_header_and_read_io(file, SHEET_NAME, HEADER_KEY_COLS)
        if df_temp is not None and not df_temp.empty:
            all_dfs.append(df_temp)

    if not all_dfs:
        return None, None, "", ["❌ 未能从上传的文件中识别出有效数据，请检查。"]

    df = pd.concat(all_dfs, ignore_index=True)
    
    dedup_subset = ["大区", "区域公司", "合同编号", "供应商名称", "品种", "合同数量", "合同单价", "合同金额", "逾期天数"]
    actual_dedup_subset = [col for col in dedup_subset if col in df.columns]
    df.drop_duplicates(subset=actual_dedup_subset, inplace=True)

    # === 【新增逻辑】玉米中心业务部门拆解 ===
    if "区域公司" in df.columns and "业务部门" in df.columns:
        mask = (df["区域公司"] == "玉米中心") & df["业务部门"].astype(str).str.contains("营口|锦州|大连", na=False)
        if mask.any():
            def extract_keyword(text):
                match = re.search(r'(营口|锦州|大连)', str(text))
                return match.group(1) if match else text
            df.loc[mask, "区域公司"] = df.loc[mask, "业务部门"].apply(extract_keyword)
            df.loc[mask, "大区"] = "港口平台"

    df = process_basic_columns(df, DATE_COLS, FLOAT_COLS, INT_COLS)
    if "逾期分类" in df.columns:
        df = df[df["逾期分类"] == "A 实际已经逾期（含进口采购）"].copy()

    adjust_col = "逾期采购金额调整值"
    if adjust_col in df.columns and "逾期采购金额" in df.columns:
        df["逾期金额确定"] = df[adjust_col] + df["逾期采购金额"]
    else:
        df["逾期金额确定"] = df.get("逾期采购金额", 0)

    if "合同单价" in df.columns:
        df["合同单价_safe"] = df["合同单价"].replace(0, np.nan)
        df["逾期数量"] = (df["逾期金额确定"] / df["合同单价_safe"]).fillna(0)
    else:
        df["逾期数量"] = 0

    df = df[df["逾期金额确定"] > 0].copy()
    if "合同编号" in df.columns:
        df = df[~df["合同编号"].astype(str).str.contains("ZLMY231109YMG016", na=False)]
    if "供应商名称" in df.columns:
        df = df[df["供应商名称"] != "中粮零售客商"]

    if "区域公司" in df.columns:
        df["区域公司"] = df["区域公司"].astype(str).str.replace("深圳公司", "沿海大区", regex=False).str.replace("经营部", "", regex=False)
    if "品种" in df.columns:
        df["品种"] = df["品种"].astype(str).str.replace("中晚籼", "稻谷", regex=False).str.replace("其他饲料原料及成品", "麸皮", regex=False)

    if "当前市场价格" in df.columns and "合同单价" in df.columns:
        df["盯市盈亏"] = df["当前市场价格"] - df["合同单价"]
        df["市场价格金额"] = df["逾期数量"] * df["当前市场价格"]
        df["盈亏金额"] = df["盯市盈亏"] * df["逾期数量"]

    for col in AMOUNT_COLUMNS_TO_CONVERT:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0) / 10000.0

    df.drop(columns=["逾期采购金额", "逾期采购金额调整值"], errors="ignore", inplace=True)
    rename_mapping = {}
    for col in AMOUNT_COLUMNS_TO_CONVERT:
        rename_mapping[col] = f"{col}（万元）"
    for col in TON_COLS:
        rename_mapping[col] = f"{col}（吨）"
    for col in YUAN_COLS:
        rename_mapping[col] = f"{col}（元）"
    rename_mapping["逾期金额确定"] = "逾期采购金额（万元）"
    df.rename(columns=rename_mapping, inplace=True)

    for col in DATE_COLS:
        if col in df.columns:
            df[col] = df[col].dt.strftime('%Y-%m-%d')
    for col in FINAL_OUTPUT_COLS:
        if col not in df.columns:
            df[col] = ""

    df_output = df[FINAL_OUTPUT_COLS].copy()
    df_output.drop_duplicates(inplace=True)

    excel_io = io.BytesIO()
    with pd.ExcelWriter(excel_io, engine='openpyxl') as writer:
        df_output.to_excel(writer, index=False)
    excel_io.seek(0)

    wb = openpyxl.load_workbook(excel_io)
    beautify_excel_purchase(wb)
    final_excel_io = io.BytesIO()
    wb.save(final_excel_io)
    final_excel_io.seek(0)

    doc_io, web_text = generate_weekly_report(df_output)

    logs.append("✅ 逾期采购处理成功！")
    return final_excel_io, doc_io, web_text, logs
