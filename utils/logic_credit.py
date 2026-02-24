import os
import io
import datetime
from datetime import timedelta
import tempfile
import platform
import warnings
import textwrap
import openpyxl
from openpyxl.utils import range_boundaries, get_column_letter

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 尝试导入绘图库作为 Linux 环境的降级图片生成方案
try:
    import matplotlib.pyplot as plt
    import matplotlib as mpl
    import matplotlib.patches as patches
    from matplotlib.font_manager import FontProperties
    mpl.rcParams['axes.unicode_minus'] = False
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

warnings.simplefilter("ignore", category=UserWarning)

# ==================== 基础辅助函数 ====================

def kill_excel_processes():
    if platform.system() == "Windows":
        try:
            os.system("taskkill /f /im excel.exe >nul 2>&1")
            os.system("taskkill /f /im et.exe >nul 2>&1")
        except:
            pass

def clean_value(val):
    if val is None: return ""
    return str(val).strip()

def clean_money(val):
    if val is None: return 0.0
    try:
        if isinstance(val, str):
            val = val.replace(',', '').strip()
            if not val: return 0.0
        return float(val)
    except:
        return 0.0

def get_cell_fill_color(cell):
    if cell.fill and cell.fill.start_color:
        color = cell.fill.start_color
        if not color.index or color.index == '00000000':
             return False
        return True
    return False

# ==================== Word 生成逻辑 (保持不变) ====================

def set_font_style(run, font_name='宋体', size=12, bold=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

def generate_word_in_memory(file_stream):
    """内存级生成 Word 报告"""
    logs = []
    report_text_dict = {} 
    
    try:
        file_stream.seek(0)
        wb = openpyxl.load_workbook(file_stream, data_only=True)
    except Exception as e:
        return None, {}, [f"❌ 读取 Excel 文件失败: {e}"]

    target_sheet_name = "每日-各品种线战略客户逾期通报"
    if target_sheet_name not in wb.sheetnames:
        return None, {}, [f"❌ 未找到关键工作表: {target_sheet_name}"]
    
    ws = wb[target_sheet_name]
    
    header_row_idx = None
    col_map = {}
    required_cols = ["品种线", "大区", "经营部", "客户名称", "合同号", "品种", "逾期天数", "逾期金额"]
    
    for r in range(1, 21):
        row_values = [clean_value(c.value) for c in ws[r]]
        matches = sum(1 for k in required_cols if any(k in v for v in row_values))
        if matches >= 4:
            header_row_idx = r
            for idx, val in enumerate(row_values):
                for k in required_cols:
                    if k in val: col_map[k] = idx
            break
            
    if not header_row_idx:
        return None, {}, ["❌ 未找到标题行。"]

    scope_ranges = {}
    target_keywords = ["玉米", "粮谷", "大豆"] 
    p_col_idx = col_map.get("品种线")
    
    for rng in ws.merged_cells.ranges:
        if rng.min_col <= (p_col_idx + 1) <= rng.max_col:
            top_val = clean_value(ws.cell(row=rng.min_row, column=rng.min_col).value)
            for kw in target_keywords:
                if kw in top_val:
                    start_r = max(rng.min_row, header_row_idx + 1)
                    end_r = rng.max_row
                    if start_r <= end_r:
                        scope_ranges[kw] = range(start_r, end_r + 1)
                    break

    data_store = {}
    exclude_regions = ["东北大区", "内陆大区", "沿江大区", "沿海大区", "东北", "内陆", "沿江", "沿海"]

    for kw, row_range in scope_ranges.items():
        if kw not in data_store: data_store[kw] = {}
        current_group = None
        if kw == "大豆":
            current_group = "ALL_SOYBEAN"
            data_store[kw][current_group] = []
        
        for r_idx in row_range:
            row = ws[r_idx]
            def get_val(key):
                idx = col_map.get(key)
                return row[idx].value if idx is not None else None

            region_str = clean_value(get_val("大区"))
            client_name = clean_value(get_val("客户名称"))
            
            is_group = False
            if kw != "大豆":
                cell_region = row[col_map.get("大区")]
                if region_str and (region_str not in exclude_regions) and get_cell_fill_color(cell_region):
                    is_group = True
            
            if is_group:
                current_group = region_str
                if current_group not in data_store[kw]:
                    data_store[kw][current_group] = []
            elif current_group:
                money_val = clean_money(get_val("逾期金额"))
                if client_name and money_val > 0:
                    data_store[kw][current_group].append({
                        '大区': region_str, '经营部': clean_value(get_val("经营部")),
                        '客户名称': client_name, '合同号': clean_value(get_val("合同号")),
                        '品种': clean_value(get_val("品种")),
                        '逾期天数': clean_value(get_val("逾期天数")).replace('.0', ''),
                        '逾期金额': money_val
                    })

    doc = Document()
    yesterday = datetime.datetime.now() - timedelta(days=1)
    date_str = f"{yesterday.year}年{yesterday.month}月{yesterday.day}日"
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font_style(p_title.add_run("信用风险管理日报"), font_name='黑体', size=16, bold=True)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font_style(p_sub.add_run(f"截至日期：{date_str}"), font_name='楷体', size=12)
    doc.add_paragraph()

    chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    has_content = False
    
    for key in target_keywords:
        if key not in data_store or not data_store[key]: continue
        groups = data_store[key]

        total_count = 0
        total_money = 0
        valid_groups_list = []
        all_rows_flat = [] 

        for g_name, rows in groups.items():
            if len(rows) > 0:
                g_total = sum(r['逾期金额'] for r in rows)
                valid_groups_list.append({'name': g_name, 'rows': rows, 'total': g_total})
                all_rows_flat.extend(rows)
                total_count += len(rows)
                total_money += g_total
        
        if total_count == 0: continue
        has_content = True
        
        header_text = (f"按照公司领导要求，请{key}中心充分发挥与战略客户的良好沟通机制，"
                       f"协助大区督促客户及时回款，压降逾期赊销。截至{date_str}，"
                       f"{key}中心战略客户，共计逾期{total_count}笔，逾期金额{int(total_money)}万元。")
        
        p_h = doc.add_paragraph()
        set_font_style(p_h.add_run(f"【{key}中心】"), font_name='黑体', size=14, bold=True)
        p_b = doc.add_paragraph()
        p_b.paragraph_format.first_line_indent = Pt(24)
        set_font_style(p_b.add_run(header_text), font_name='宋体', size=12)
        
        center_text_block = f"【{key}中心】\n{header_text}\n"

        if key == "大豆":
            all_rows_flat.sort(key=lambda r: r['逾期金额'], reverse=True)
            for j, row in enumerate(all_rows_flat):
                line = (f"{j+1}、{row['大区']}，{row['经营部']}，{row['客户名称']}，"
                        f"{row['合同号']}，{row['品种']}，逾期{row['逾期天数']}天，"
                        f"{int(row['逾期金额'])}万元；")
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(24)
                set_font_style(p.add_run(line), font_name='宋体', size=12)
                center_text_block += f"{line}\n"
        else:
            valid_groups_list.sort(key=lambda x: x['total'], reverse=True)
            for i, g_data in enumerate(valid_groups_list):
                idx_str = chinese_nums[i] if i < len(chinese_nums) else str(i+1)
                group_line = (f"{idx_str}、{g_data['name']}，共计逾期{len(g_data['rows'])}笔，"
                              f"逾期金额{int(g_data['total'])}万元。")
                
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(24)
                set_font_style(p.add_run(group_line), font_name='黑体', size=12, bold=True)
                center_text_block += f"{group_line}\n"
                
                sorted_rows = sorted(g_data['rows'], key=lambda r: r['逾期金额'], reverse=True)
                for j, row in enumerate(sorted_rows):
                    line = (f"{j+1}、{row['大区']}，{row['经营部']}，{row['客户名称']}，"
                            f"{row['合同号']}，{row['品种']}，逾期{row['逾期天数']}天，"
                            f"{int(row['逾期金额'])}万元；")
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Pt(24)
                    set_font_style(p.add_run(line), font_name='宋体', size=12)
                    center_text_block += f"{line}\n"
        doc.add_paragraph()
        
        report_text_dict[key] = center_text_block

    if not has_content:
        return None, {}, ["⚠️ 未提取到逾期数据，无 Word 报告生成。"]

    out_stream = io.BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    
    logs.append("✅ Word 报告内存生成成功！")
    return out_stream, report_text_dict, logs


# ==================== 终极防蜷缩：100%纯物理镜像渲染引擎 ====================

def render_sheet_range_to_image_stream(ws, range_str):
    """
    废除一切自主推断！完全依附 Excel 原生格式进行 1:1 像素投射。
    解决：红色字体、加粗丢失、标题/时间剥离丢失、非表头序号底色等问题。
    """
    if not MATPLOTLIB_AVAILABLE:
        return None

    current_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 1. 严格挂载字体库
    regular_path = os.path.join(current_dir, 'msyh.ttc')
    if not os.path.exists(regular_path): regular_path = os.path.join(current_dir, 'msyh.ttf')
    custom_font_regular = FontProperties(fname=regular_path) if os.path.exists(regular_path) else None

    bold_path = os.path.join(current_dir, 'msyhbd.ttc')
    if not os.path.exists(bold_path): bold_path = os.path.join(current_dir, 'msyhbd.ttf')
    custom_font_bold = FontProperties(fname=bold_path) if os.path.exists(bold_path) else custom_font_regular

    # 2. 初始框选范围
    range_str = range_str.replace('$', '')
    min_col, min_row, max_col, max_row = range_boundaries(range_str)

    # 3. 动态截断：物理切除底部的冗余复选框行
    actual_max_row = max_row
    for r in range(min_row, max_row + 1):
        row_vals = [str(ws.cell(row=r, column=c).value or "").strip() for c in range(min_col, max_col + 1)]
        combined = "".join(row_vals)
        if "是否填报" in combined or "填报说明" in combined:
            actual_max_row = r - 1 
            break

    # 4. 动态去除完全空白的列 (彻底删除冗余空白列)
    valid_cols_set = set()
    for r in range(min_row, actual_max_row + 1):
        for c in range(min_col, max_col + 1):
            val = ws.cell(row=r, column=c).value
            if val is not None and str(val).strip() != "":
                valid_cols_set.add(c)
                # 必须将合并单元格覆盖的附属列也加入有效列
                for mr in ws.merged_cells.ranges:
                    if mr.min_row <= r <= mr.max_row and mr.min_col <= c <= mr.max_col:
                        for mc in range(mr.min_col, mr.max_col + 1):
                            valid_cols_set.add(mc)
    
    valid_cols = sorted(list(valid_cols_set))
    if not valid_cols: return None

    # 5. 映射所有合并单元格坐标网络
    merged_dict = {}
    for mr in ws.merged_cells.ranges:
        if mr.min_col <= max_col and mr.max_col >= min_col and mr.min_row <= actual_max_row and mr.max_row >= min_row:
            for r in range(mr.min_row, mr.max_row + 1):
                for c in range(mr.min_col, mr.max_col + 1):
                    merged_dict[(r, c)] = {
                        'top_left': (mr.min_row, mr.min_col),
                        'bottom_right': (mr.max_row, mr.max_col)
                    }

    # 6. 定位表格的主体范围 (用于边框与对齐控制)
    header_start_row = min_row
    for r in range(min_row, actual_max_row + 1):
        combined = "".join([str(ws.cell(row=r, column=c).value or "").strip() for c in valid_cols])
        if "序号" in combined or "业务单位" in combined or "大区" in combined:
            header_start_row = r
            break

    # 7. 计算物理最佳行列比例
    col_widths = {c: 4.0 for c in valid_cols}
    row_heights = {}
    row_types = {}
    
    for r in range(min_row, actual_max_row + 1):
        combined = "".join([str(ws.cell(row=r, column=c).value or "").strip() for c in valid_cols])
        
        # 将表格上方/下方的零散信息定义为 meta（无边框层）
        if r < header_start_row:
            row_types[r] = 'meta'
            row_heights[r] = 4.0 if "表" in combined else 2.0
        elif r > header_start_row + 5 and "单位" in combined and "万" in combined and "合计" not in combined:
            row_types[r] = 'meta'
            row_heights[r] = 1.5
        else:
            row_types[r] = 'grid' # 表格核心数据区
            row_heights[r] = 3.2 if r <= header_start_row + 1 else 2.4

    # 文本长度探测列宽
    for r in range(min_row, actual_max_row + 1):
        if row_types[r] == 'meta': continue
        for c in valid_cols:
            is_spanned = False
            for mr in ws.merged_cells.ranges:
                if mr.min_row <= r <= mr.max_row and mr.min_col <= c <= mr.max_col:
                    if mr.max_col > mr.min_col: is_spanned = True
            if is_spanned: continue 

            val = ws.cell(row=r, column=c).value
            if val:
                text_len = sum(1.8 if ord(ch) > 255 else 1.1 for ch in str(val))
                w = text_len * 0.9 + 1.5 
                if w > col_widths[c]: col_widths[c] = min(w, 25.0)

    col_widths[valid_cols[0]] = max(3.5, col_widths[valid_cols[0]])

    # 百分比探测
    col_is_percent = {c: False for c in valid_cols}
    for c in valid_cols:
        for r in range(header_start_row, header_start_row + 3):
            v = str(ws.cell(row=r, column=c).value or "")
            if "率" in v or "占比" in v or "%" in v:
                col_is_percent[c] = True
                break

    # 8. 建立高清 500 DPI 绝对坐标画板
    W_grid = sum(col_widths.values())
    H_grid = sum(row_heights.values())

    A4_W, A4_H = 8.27, 11.69
    margin_x, margin_y = 0.4, 0.4
    
    max_w_in = A4_W - 2 * margin_x
    S = max_w_in / W_grid 
    H_in = H_grid * S
    
    # 动态适应纸张长度
    Final_H = max(A4_H, H_in + 1.0)
    fig = plt.figure(figsize=(A4_W, Final_H), dpi=500) # 🔥 500DPI解决模糊问题
    fig.patch.set_facecolor('white')

    ax = fig.add_axes([margin_x / A4_W, (Final_H - H_in - margin_y) / Final_H, max_w_in / A4_W, H_in / Final_H])
    ax.set_xlim(0, W_grid)
    ax.set_ylim(H_grid, 0)
    ax.axis('off')

    base_fs = 2.5 * S * 72 * 0.42 

    # 9. 逐像素矩阵渲染
    y_curr = 0
    for r in range(min_row, actual_max_row + 1):
        x_curr = 0
        rh = row_heights[r]
        rtype = row_types[r]
        
        for c in valid_cols:
            cw = col_widths[c]
            
            is_merged_top_left = True
            draw_w, draw_h = cw, rh
            
            if (r, c) in merged_dict:
                info = merged_dict[(r, c)]
                if (r, c) != info['top_left']:
                    is_merged_top_left = False 
                else:
                    draw_w = sum(col_widths.get(mc, 4.0) for mc in range(info['top_left'][1], info['bottom_right'][1] + 1) if mc in valid_cols)
                    draw_h = sum(row_heights.get(mr, 2.4) for mr in range(info['top_left'][0], info['bottom_right'][0] + 1))

            if is_merged_top_left:
                cell = ws.cell(row=r, column=c)
                
                # --- A. 原生底色提取 (绝不越权染色) ---
                bg_color = '#FFFFFF'
                if rtype == 'grid':
                    if cell.fill and cell.fill.patternType == 'solid' and cell.fill.start_color.rgb:
                        rgb = str(cell.fill.start_color.rgb)
                        if len(rgb) == 8 and rgb != '00000000':
                            bg_color = '#' + rgb[2:]
                        elif len(rgb) == 6:
                            bg_color = '#' + rgb
                            
                # 线宽：外围标题/时间/单位层不画线框
                lw = 0.8 if rtype == 'grid' else 0.0
                rect = patches.Rectangle((x_curr, y_curr), draw_w, draw_h, facecolor=bg_color, edgecolor='#000000', linewidth=lw)
                ax.add_patch(rect)
                
                # --- B. 字体属性全息继承 ---
                val = cell.value
                fmt = cell.number_format or "General"
                text = ""
                
                if val is not None and str(val).strip() != "":
                    if isinstance(val, (int, float)):
                        if '%' in fmt or col_is_percent[c]:
                            if '.00' in fmt: text = f"{val:.2%}"
                            elif '.0' in fmt: text = f"{val:.1%}"
                            else: text = f"{val:.0%}"
                        elif ',' in fmt or (isinstance(val, (int, float)) and (val >= 1000 or val <= -1000)):
                            if isinstance(val, float) and not val.is_integer():
                                text = f"{val:,.2f}".rstrip('0').rstrip('.')
                            else:
                                text = f"{val:,.0f}"
                        else:
                            if isinstance(val, float):
                                text = f"{val:.2f}".rstrip('0').rstrip('.')
                            else:
                                text = str(val)
                    elif isinstance(val, datetime.datetime):
                        if "年" in fmt: text = val.strftime('%Y年%m月%d日')
                        else: text = val.strftime('%Y-%m-%d')
                    else:
                        text = str(val).strip()
                
                # 🔥 字体加粗与颜色直接取自原生 Excel
                is_bold = False
                if cell.font and cell.font.bold:
                    is_bold = True
                
                # 预防性补丁：大标题无论如何必须加粗
                if rtype == 'meta' and "表" in text:
                    is_bold = True

                # 🔥 还原原表格的红色警示字
                text_color = '#000000'
                if cell.font and cell.font.color and hasattr(cell.font.color, 'rgb') and cell.font.color.rgb:
                    rgb_val = str(cell.font.color.rgb)
                    if len(rgb_val) == 8 and rgb_val != '00000000':
                        text_color = '#' + rgb_val[2:]
                    elif len(rgb_val) == 6:
                        text_color = '#' + rgb_val
                
                # --- C. 对齐方式定位 ---
                halign = 'center'
                valign = 'center'
                
                if rtype == 'meta':
                    excel_h = cell.alignment.horizontal if cell.alignment else None
                    if excel_h in ['left', 'right', 'center']: 
                        halign = excel_h
                    else:
                        # 靠左的制表单位，靠右的单位与时间兜底
                        if "表" in text: halign = 'center'
                        elif "单位:万" in text or "单位：万" in text: halign = 'right'
                        else: halign = 'left' 
                else:
                    halign = 'center' # 强力锁死：表格主体内细节必须全部居中对齐
                    
                pad_x = 1.0
                if halign == 'left': text_x = x_curr + pad_x
                elif halign == 'right': text_x = x_curr + draw_w - pad_x
                else: text_x = x_curr + draw_w / 2
                
                if rtype == 'meta' and not "表" in text:
                    valign = 'bottom'
                    text_y = y_curr + draw_h - 0.2
                else:
                    valign = 'center'
                    text_y = y_curr + draw_h / 2
                
                # 字号阶梯分配
                fs = base_fs
                if rtype == 'meta':
                    if "表" in text: fs = base_fs * 1.5
                    else: fs = base_fs * 0.95

                if isinstance(text, str) and len(text) > (draw_w / 1.1):
                    wrap_w = max(1, int(draw_w / 1.1))
                    text = '\n'.join(textwrap.wrap(text, width=wrap_w))
                    
                # --- D. 投射渲染 ---
                if text:
                    kwargs = {
                        'ha': halign,
                        'va': valign,
                        'color': text_color,
                        'clip_on': True
                    }
                    
                    if is_bold and custom_font_bold:
                        prop = custom_font_bold.copy()
                        prop.set_size(fs)
                        kwargs['fontproperties'] = prop
                    elif custom_font_regular:
                        prop = custom_font_regular.copy()
                        if is_bold: prop.set_weight('bold')
                        prop.set_size(fs)
                        kwargs['fontproperties'] = prop
                    else:
                        kwargs['weight'] = 'bold' if is_bold else 'normal'
                        kwargs['fontsize'] = fs
                        
                    ax.text(text_x, text_y, text, **kwargs)

            x_curr += cw
        y_curr += rh

    img_stream = io.BytesIO()
    fig.savefig(img_stream, format='png', dpi=500, facecolor=fig.get_facecolor(), edgecolor='none', bbox_inches='tight')
    plt.close(fig)
    img_stream.seek(0)
    return img_stream

# ==================== 导出文件生成逻辑 ====================

def generate_export_files_in_memory(file_stream):
    """根据操作系统，智能生成 PDF 或完美防变形 PNG"""
    results = []
    logs = []
    today_mmdd = datetime.datetime.now().strftime('%m%d')
    sys_name = platform.system()
    
    sheets_info = [
        {"name": "每日-中粮贸易外部赊销限额使用监控表", "range": "$A$1:$G$30", "base_title": "中粮贸易外部赊销限额使用监控表"}
    ]
    
    if True: 
        sheets_info.append({"name": "每周-正大额度使用情况", "range": "$A$1:$L$34", "base_title": "正大额度使用情况"})
        
    if sys_name == 'Windows':
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_in:
            file_stream.seek(0)
            tmp_in.write(file_stream.read())
            temp_excel_path = tmp_in.name
            
        try:
            import win32com.client
            try:
                app = win32com.client.Dispatch("Excel.Application")
            except:
                app = win32com.client.Dispatch("Ket.Application") 
                
            app.Visible = False
            app.DisplayAlerts = False
            wb = app.Workbooks.Open(temp_excel_path, ReadOnly=True)
            
            for s_info in sheets_info:
                try:
                    ws = wb.Sheets(s_info['name'])
                    ws.PageSetup.PrintArea = s_info['range']
                    ws.PageSetup.Orientation = 1
                    ws.PageSetup.PaperSize = 9
                    ws.PageSetup.Zoom = False
                    ws.PageSetup.FitToPagesWide = 1
                    ws.PageSetup.FitToPagesTall = 1
                    ws.PageSetup.CenterHorizontally = True
                    
                    out_name = f"{s_info['base_title']}{today_mmdd}.pdf"
                    temp_pdf_path = os.path.join(tempfile.gettempdir(), out_name)
                    
                    ws.ExportAsFixedFormat(0, temp_pdf_path, IgnorePrintAreas=False)
                    
                    with open(temp_pdf_path, "rb") as f:
                        results.append({"name": out_name, "data": f.read(), "type": "pdf"})
                    os.remove(temp_pdf_path)
                    logs.append(f"   ✅ 成功生成 PDF: {out_name}")
                except Exception as e:
                    logs.append(f"   ⚠️ 跳过 {s_info['name']}: {str(e)}")
                    
            wb.Close(SaveChanges=False)
            app.Quit()
        except Exception as e:
            logs.append(f"❌ PDF 引擎调用失败: {str(e)}")
        finally:
            if os.path.exists(temp_excel_path):
                os.remove(temp_excel_path)
                
    else:
        file_stream.seek(0)
        try:
            wb = openpyxl.load_workbook(file_stream, data_only=True)
            for s_info in sheets_info:
                if s_info['name'] in wb.sheetnames:
                    # 调用 100% 原生物理镜像渲染器
                    img_stream = render_sheet_range_to_image_stream(wb[s_info['name']], s_info['range'])
                    if img_stream:
                        out_name = f"{s_info['base_title']}{today_mmdd}.png"
                        results.append({"name": out_name, "data": img_stream.read(), "type": "png"})
                        logs.append(f"   ✅ 成功生成像素级对齐图片: {out_name}")
        except Exception as e:
            logs.append(f"❌ 跨平台渲染引擎出错: {str(e)}")
            
    return results, logs

# ==================== 主控入口 ====================

def process_credit_report(uploaded_file):
    """处理风险管理日报主入口"""
    logs = []
    sys_name = platform.system()
    env_msg = f"当前环境: {sys_name} " + ("(原生支持 PDF 导出)" if sys_name == 'Windows' else "(云端环境，将生成高清预览图替代 PDF)")
    
    kill_excel_processes()
    file_stream = io.BytesIO(uploaded_file.getvalue())
    
    word_bytes, word_text_dict, word_logs = generate_word_in_memory(file_stream)
    logs.extend(word_logs)
    
    export_files, export_logs = generate_export_files_in_memory(file_stream)
    logs.extend(export_logs)
    
    kill_excel_processes()
    
    return word_bytes, word_text_dict, export_files, logs, env_msg
